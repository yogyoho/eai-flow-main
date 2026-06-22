"""Analyze a sample EIA .docx for structural elements — compact output."""
import docx, re, json, sys

import sys; path = sys.argv[1] if len(sys.argv) > 1 else "/tmp/eia_sample.docx"
doc = docx.Document(path)
ns = {"wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"}

# 1. Chapter structure (H1 + H2)
chapters = []
for p in doc.paragraphs:
    s = (p.style.name or "") if p.style else ""
    if not s.startswith("Heading"):
        continue
    m = re.search(r"(\d+)", s)
    lv = int(m.group(1)) if m else 1
    if lv <= 2:
        chapters.append((lv, p.text.strip()[:200]))

# 2. Counts
tables_n = len(doc.tables)
images_n = 0
for p in doc.paragraphs:
    images_n += len(p._element.findall(".//wp:inline", ns))
    images_n += len(p._element.findall(".//wp:anchor", ns))
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for pp in cell.paragraphs:
                images_n += len(pp._element.findall(".//wp:inline", ns))
                images_n += len(pp._element.findall(".//wp:anchor", ns))

all_text = " ".join(p.text for p in doc.paragraphs)

# 3. Key patterns
checks = {
    "total_pages_approx": len(doc.paragraphs) // 8,
    "tables": tables_n,
    "images": images_n,
    "sections_layout": len(doc.sections),
    "figure_captions": len(re.findall(r"图\s*\d+", all_text)),
    "table_captions": len(re.findall(r"表\s*\d+", all_text)),
    "standards_refs": len(re.findall(r"(GB|HJ|DZ|AQ)\s*\d+", all_text)),
    "formula_refs": len(re.findall(r"（\d+(\.\d+)*[）)]|计算公式|计算模型|预测模式|预测公式", all_text)),
    "monitoring_data": len(re.findall(r"监测.*(?:mg/m³|μg|dB|mg/L|pH)", all_text)),
    "appendices_ref": len(re.findall(r"附录|附件[一二三四五六七八九十]", all_text)),
    "list_paras": sum(1 for p in doc.paragraphs if "List" in (p.style.name or "") if p.style),
}

output = {
    "chapters": chapters,
    "counts": checks,
}

json.dump(output, sys.stdout, ensure_ascii=False, indent=1)
