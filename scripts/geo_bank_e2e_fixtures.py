#!/usr/bin/env python3
"""geo_bank_e2e_fixtures.py — geo-sample-bank Phase 2 T9 端到端演练合成夹具生成器。

一次性演练工具（非 CI）：python-docx 造 4 份 .docx 勘查报告——金 e2e-gold-001/002、
煤 e2e-coal-001/002（stage=exploration），供管理 API 全链演练
（upload→parse→redact→review approve→pipeline/compile→resolve_targets 实证）。

文档结构与转换契约（docx→md 见 backend/app/extensions/geo_samples/parsers.py）：
  - 节标题直接用英文内置样式 "Heading 1"/"Heading 2"——中文「标题 1」样式名的归一化
    已由 parsers._STYLE_ALIASES（Phase 2 T1）保证，本脚本不需要也不应另造中文样式；
  - 「1 总论」「2 地质特征」为 Heading 1 → md `## 1 …`/`## 2 …`（bank_compile.SEC_RE
    的切片锚点），「2.1 地层」为 Heading 2 → md `### 2.1 …`（子节归父片 ch2）；
  - 1 个 2 行×3 列表格 → docx_to_markdown 转 md 管道表（calibrate 计 table_rows）；
  - PII 全部为虚构假值（X 尾号占位），分别命中 redactor 的 exploration_cert /
    coord_pair / phone（auto 档替换为 ****）与 person_field（review 档只记事件）；
  - 地质数值（品位/灰分/发热量/厚度/倾角等）红线必须原样过脱敏——redactor 无任何
    匹配地质数值的规则，编译后 SL3 指纹与深度标定才有意义；演练后应在 clean 正文
    与切片里复核这些数值在场、PII 已 mask。

用法：
    python scripts/geo_bank_e2e_fixtures.py [--out <dir>]
默认输出 %TEMP%/gsb_e2e（tempfile.gettempdir()/gsb_e2e）；结束打印 4 份上传 curl。
依赖：python-docx（host 缺失时 `pip install python-docx`，gateway 容器 venv 已带）。
"""

from __future__ import annotations

import argparse
import io
import sys
import tempfile
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[1]

# 4 份合成报告的元数据与植入值。所有证号/坐标/电话/人名均为虚构占位（X 尾号），
# 不对应任何真实矿业权、地点或个人；四份内容互异 → file_hash 去重（409）不误伤。
SPECS: list[dict] = [
    {
        "report_id": "e2e-gold-001",
        "file_name": "e2e-gold-001.docx",
        "mineral": "gold",
        "mineral_cn": "金",
        "year": 2023,
        "title": "E2E 演练合成金矿普查报告 001（数据全部虚构）",
        "cert": "C5300002023000001",
        "coord": "X 3546123.45, Y 38456789.12",
        "phone": "13812345671",
        "person": "张三丰",
        "area": "12.36",
        "grade_main": "3.86 g/t",
        "grade_edge": "1.20 g/t",
        "grade_by": "0.85%",
        "res_ore": "1250",
        "thickness": "8.6",
        "dip": "25",
        "overburden": "0.5～3.2",
        "cv": None,  # 煤特有：发热量
        "ash": None,  # 煤特有：灰分
    },
    {
        "report_id": "e2e-gold-002",
        "file_name": "e2e-gold-002.docx",
        "mineral": "gold",
        "mineral_cn": "金",
        "year": 2024,
        "title": "E2E 演练合成金矿普查报告 002（数据全部虚构）",
        "cert": "C5300002023000002",
        "coord": "X 3546234.56, Y 38457890.23",
        "phone": "13812345672",
        "person": "李四光",
        "area": "9.87",
        "grade_main": "4.12 g/t",
        "grade_edge": "1.00 g/t",
        "grade_by": "0.85%",
        "res_ore": "980",
        "thickness": "7.2",
        "dip": "31",
        "overburden": "0.8～4.1",
        "cv": None,
        "ash": None,
    },
    {
        "report_id": "e2e-coal-001",
        "file_name": "e2e-coal-001.docx",
        "mineral": "coal",
        "mineral_cn": "煤",
        "year": 2023,
        "title": "E2E 演练合成煤矿详查报告 001（数据全部虚构）",
        "cert": "C5300002023000003",
        "coord": "X 3546345.67, Y 38458901.34",
        "phone": "13812345673",
        "person": "王五",
        "area": "26.50",
        "grade_main": "21.5 MJ/kg",  # 煤：收到基低位发热量占品位位
        "grade_edge": "16.8 MJ/kg",  # 最低可采发热量
        "grade_by": "0.85%",  # 全硫（与金报告同一数值，验证数值类通配原样过脱敏）
        "res_ore": "23800",
        "thickness": "6.2",
        "dip": "12",
        "overburden": "1.2～5.6",
        "cv": "21.5",
        "ash": "18.6%",
    },
    {
        "report_id": "e2e-coal-002",
        "file_name": "e2e-coal-002.docx",
        "mineral": "coal",
        "mineral_cn": "煤",
        "year": 2024,
        "title": "E2E 演练合成煤矿详查报告 002（数据全部虚构）",
        "cert": "C5300002023000004",
        "coord": "X 3546456.78, Y 38459012.45",
        "phone": "13812345674",
        "person": "赵六",
        "area": "31.20",
        "grade_main": "22.8 MJ/kg",
        "grade_edge": "17.1 MJ/kg",
        "grade_by": "0.85%",
        "res_ore": "31500",
        "thickness": "7.8",
        "dip": "16",
        "overburden": "2.0～6.3",
        "cv": "22.8",
        "ash": "19.4%",
    },
]


def _narrative(spec: dict) -> list[tuple[str, list[str]]]:
    """单份报告的三节正文：[(标题, [段落…])]，表格单独插在 2.1 末尾。"""
    if spec["mineral"] == "gold":
        intro = [
            f"本报告为 geo-sample-bank 端到端演练专用合成数据，探矿许可证号{spec['cert']}为虚构占位，不对应任何真实矿业权。矿区位于演练省云岭县，面积{spec['area']}平方公里。",
            f"矿区中心高斯坐标：{spec['coord']}（1980 西安坐标系，6 度带投影）。项目联系电话{spec['phone']}；负责人：{spec['person']}。",
            f"主矿体金平均品位{spec['grade_main']}，边界品位{spec['grade_edge']}，伴生铜品位{spec['grade_by']}，估算推断资源量矿石量{spec['res_ore']}万吨，共圈定矿体 3 条。",
        ]
        geology = [
            "矿区大地构造位置位于演练褶皱带东段，出露地层以二叠系灰岩为主，北东向断裂为主要控矿构造，矿化与硅化、黄铁矿化关系密切，围岩蚀变分带明显。",
        ]
        strata = [
            f"区内出露地层由老至新依次为二叠系下统栖霞组、茅口组灰岩及第四系残坡积物，覆盖层厚{spec['overburden']} m。",
            f"含矿带走向长约 1200 m，平均厚度{spec['thickness']} m，倾角{spec['dip']}°～42°，品位变化系数 136%，属较均匀型。",
        ]
        table_header = ["矿体编号", "平均品位", "平均厚度"]
        table_row = ["I号矿体", spec["grade_main"], f"{spec['thickness']} m"]
    else:  # coal
        intro = [
            f"本报告为 geo-sample-bank 端到端演练专用合成数据，探矿许可证号{spec['cert']}为虚构占位，不对应任何真实矿业权。井田位于演练省黑岱县，面积{spec['area']}平方公里。",
            f"井田中心高斯坐标：{spec['coord']}（1980 西安坐标系，6 度带投影）。项目联系电话{spec['phone']}；负责人：{spec['person']}。",
            f"主要可采煤层为 2 号与 9 号煤层，原煤收到基低位发热量{spec['grade_main']}，最低可采发热量{spec['grade_edge']}，全硫{spec['grade_by']}，估算资源量{spec['res_ore']}万吨，属中灰低硫动力煤。",
        ]
        geology = [
            "井田位于演练向斜北翼，地层总体走向北东、倾角平缓，构造以宽缓褶曲为主，未发现规模性断层，水文地质条件简单，瓦斯等级为低瓦斯。",
        ]
        strata = [
            f"区内地层由老至新为二叠系山西组、下石盒子组及第四系松散沉积物，覆盖层厚{spec['overburden']} m。",
            f"2 号煤层厚{spec['thickness']} m，倾角{spec['dip']}°～18°，结构简单，全区可采，顶板为砂质泥岩，底板为粉砂岩。",
        ]
        table_header = ["煤层编号", "平均厚度", "灰分(A_d)"]
        table_row = ["2号煤层", f"{spec['thickness']} m", spec["ash"] or "18.6%"]
    return [
        ("1 总论", intro),
        ("2 地质特征", geology),
        ("2.1 地层", strata),
    ], (table_header, [table_row])


def build_docx(spec: dict) -> bytes:
    """单份规格 → .docx 字节流（Heading 1/2 + 正文段 + 2×3 表格）。"""
    import docx  # lazy：--help 不依赖 python-docx

    doc = docx.Document()
    doc.add_paragraph(spec["title"])  # 文档题名：普通段落（无节号，不参与切片）
    sections, table_spec = _narrative(spec)
    for title, paras in sections:
        level = 2 if "." in title else 1  # 「2.1 地层」→ Heading 2，其余 → Heading 1
        doc.add_heading(title, level=level)
        for text in paras:
            doc.add_paragraph(text)
        if title == "2.1 地层":
            header, rows = table_spec
            table = doc.add_table(rows=1 + len(rows), cols=len(header))  # 2 行×3 列
            try:
                table.style = "Table Grid"
            except KeyError:  # 模板无该样式仅失边框，文本提取不受影响
                pass
            for col, text in enumerate(header):
                table.cell(0, col).text = text
            for row_idx, row in enumerate(rows, start=1):
                for col, text in enumerate(row):
                    table.cell(row_idx, col).text = text
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def upload_commands(out_dir: Path) -> str:
    """打印 4 份夹具的上传 curl（登录 + multipart upload，upload 即自动触发 parse）。"""
    lines = [
        "# 0) 取会话 cookie（extensions 登录门面字段是 username；此后 POST 均需",
        "#    X-CSRF-Token，值取 cookie jar 里 csrf_token 列，Git Bash 用 awk 取第 7 列）：",
        'curl -s -c cookies.txt -H "Content-Type: application/json" \\',
        '  -d \'{"username":"admin@eai-flow.com","password":"Admin@2026"}\' \\',
        "  http://localhost:2026/api/extensions/auth/login",
        "",
    ]
    for spec in SPECS:
        lines += [
            f"# 上传 {spec['report_id']}（返回 document.id，后续 parse/redact/review 用它；中文 region 可选省略）：",
            'curl -s -b cookies.txt -H "X-CSRF-Token: $(awk \'$6=="csrf_token"{print $7}\' cookies.txt)" \\',
            f'  -F "file=@{(out_dir / spec["file_name"]).as_posix()}" \\',
            f'  -F "report_id={spec["report_id"]}" -F "stage=exploration" -F "mineral={spec["mineral"]}" \\',
            f'  -F "year={spec["year"]}" \\',
            "  http://localhost:2026/api/extensions/geo-samples/documents/upload",
            "",
        ]
    lines += [
        "# 逐份后续管线（<doc_id> 换上传返回的 document.id）：",
        "#   POST /api/extensions/geo-samples/documents/<doc_id>/redact   （upload 已自动 parse）",
        '#   POST /api/extensions/geo-samples/documents/<doc_id>/review   body {"decision":"approve"}',
        "#   POST /api/extensions/geo-samples/pipeline/compile?stage=exploration   （空体）",
        "#   GET  /api/extensions/geo-samples/runs   （轮询至 done，detail 应含 ragflow skipped）",
    ]
    return "\n".join(lines)


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description="geo-sample-bank T9 端到端演练合成夹具（4 份 .docx）")
    ap.add_argument("--out", default=str(Path(tempfile.gettempdir()) / "gsb_e2e"), help="输出目录（默认 %%TEMP%%/gsb_e2e）")
    args = ap.parse_args(argv)

    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)
    for spec in SPECS:
        path = out_dir / spec["file_name"]
        path.write_bytes(build_docx(spec))
        print(f"wrote {path} ({path.stat().st_size} bytes)  mineral={spec['mineral']} year={spec['year']}")
    print()
    print("植入 PII（假值）：探矿权证号 C53000020230000XX、高斯坐标 X/Y 对、手机 1381234567X、负责人字段——")
    print("  预期脱敏：前三类 auto 档替换为 ****，person_field 为 review 档只记事件不替换。")
    print("植入地质数值（必须原样过脱敏）：品位 0.85% / 3.86 g/t / 灰分 18.6% / 发热量 21.5 MJ/kg / 厚度 8.6 m 等。")
    print()
    print(upload_commands(out_dir))
    return 0


if __name__ == "__main__":
    sys.exit(main())
