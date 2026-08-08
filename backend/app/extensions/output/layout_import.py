"""Deterministic .docx → layout-template extraction.

Reads page settings, body/heading styles, table style, header/footer, and
best-effort cover structure from a sample .docx and returns a
LayoutTemplate-shaped dict (snake_case) consumed by the output/docmgr
``import-layout`` endpoints. Pure python-docx, no new dependencies.
"""

from __future__ import annotations

import base64
import itertools
import re
from copy import deepcopy
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree

from app.extensions.output.schemas import CoverElementSchema, CoverPageSchema

_PAPER_DIMS = {
    "A4": (21.0, 29.7),
    "A3": (29.7, 42.0),
    "B5": (17.6, 25.0),
    "letter": (21.59, 27.94),
}
_DRAWML = "http://schemas.openxmlformats.org/drawingml/2006/main"
_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_TOC_TEXT_RE = re.compile(r"^目\s*录$|^contents$", re.I)
_DATE_RE = re.compile(r"\d{4}[-/年]\d{1,2}")


def _to_cm(length) -> float:
    """python-docx Length (EMU) → cm, rounded to 2 decimals."""
    return round(length.cm, 2) if length is not None else 0.0


def _paper_from_dimensions(width_cm: float, height_cm: float) -> tuple[str, str]:
    """Map page width/height (cm) → (paperSize, orientation)."""
    orientation = "landscape" if width_cm > height_cm else "portrait"
    w, h = (height_cm, width_cm) if width_cm > height_cm else (width_cm, height_cm)
    best, best_err = "A4", float("inf")
    for name, (pw, ph) in _PAPER_DIMS.items():
        err = abs(w - pw) + abs(h - ph)
        if err < best_err:
            best, best_err = name, err
    return best, orientation


_LOCALE_RE = re.compile(r"^[a-z]{2}-[A-Z]{2}$")


def _clean_font(name: str | None) -> str | None:
    """A real font name, or None. Filters OOXML locale codes (zh-CN / en-US) that
    appear in w:eastAsia slots but are region tags, not fonts."""
    if not name or _LOCALE_RE.match(name):
        return None
    return name


def _style_eastAsia(style) -> str | None:
    """w:eastAsia CJK font declared on a style (cleaned), or None."""
    try:
        m = re.search(r'w:eastAsia="([^"]+)"', style.element.xml)
        return _clean_font(m.group(1)) if m else None
    except Exception:
        return None


def _doc_defaults_eastAsia(doc) -> str | None:
    """Document default CJK font from w:docDefaults' <w:rFonts w:eastAsia=...>.

    CN docs commonly leave the Normal style and most runs unformatted, so body text
    inherits this default (typically 宋体). Scoped to rFonts so it ignores the
    unrelated <w:lang w:eastAsia="zh-CN"> locale tag."""
    try:
        m = re.search(r'<w:docDefaults>.*?<w:rFonts[^>]*w:eastAsia="([^"]+)"', doc.styles.element.xml, re.S)
        return _clean_font(m.group(1)) if m else None
    except Exception:
        return None


def _style_color(style, default: str) -> str:
    try:
        rgb = style.font.color.rgb
        if rgb is not None:
            return f"#{rgb}"  # editor <input type="color"> requires #RRGGBB
    except Exception:
        pass
    return default


def _run_font(run) -> str | None:
    """w:eastAsia CJK font on a run (cleaned), or None when the run sets none and
    therefore inherits. Deliberately NOT run.font.name — that returns the Western
    (ascii) font (e.g. Times New Roman), which is wrong for CJK text."""
    try:
        m = re.search(r'w:eastAsia="([^"]+)"', run._element.xml)
        return _clean_font(m.group(1)) if m else None
    except Exception:
        return None


def _cjk_font(doc, style, paragraphs) -> str:
    """Resolve the CJK font body/heading text actually renders in.

    CN docs rarely set w:eastAsia on every run — most runs are *silent* and inherit
    the document/style default. So if a MAJORITY of runs declare eastAsia the text
    actively uses that font (take the dominant); otherwise the runs inherit and we
    walk style.eastAsia → docDefaults.eastAsia → style.font.name → 宋体. Counting the
    few runs that DO declare (often emphasis/labels) would mislabel the whole body —
    e.g. 17 黑体 runs among 887 silent 宋体 runs must not win."""
    runs = [r for p in paragraphs for r in p.runs]
    declared: dict[str, int] = {}
    for r in runs:
        fam = _run_font(r)
        if fam:
            declared[fam] = declared.get(fam, 0) + 1
    if runs and declared:
        top_font, top_n = max(declared.items(), key=lambda kv: kv[1])
        if top_n >= len(runs) * 0.5:
            return top_font
    return _style_eastAsia(style) or _doc_defaults_eastAsia(doc) or _clean_font(style.font.name) or "宋体"


def _dominant_run_font(paragraphs) -> tuple[float | None, str | None]:
    """Most common (size_pt, family) across explicitly-formatted runs in `paragraphs`.

    Returns (None, None) when no run sets a size/family — caller then falls back to
    the style. Real docs usually set body/heading size on each run (e.g. 四号 = 14pt)
    while leaving the Normal/Heading style at a template default, so sampling runs is
    far more accurate than reading the style alone.
    """
    sizes: dict[float, int] = {}
    families: dict[str, int] = {}
    for p in paragraphs:
        for run in p.runs:
            try:
                sz = run.font.size
            except Exception:
                sz = None
            if sz is not None:
                sizes[sz.pt] = sizes.get(sz.pt, 0) + 1
            fam = _run_font(run)
            if fam:
                families[fam] = families.get(fam, 0) + 1
    size = max(sizes, key=sizes.get) if sizes else None
    family = max(families, key=families.get) if families else None
    return size, family


def _body_paragraphs(doc) -> list:
    """Non-heading, non-empty body paragraphs (excludes table cells)."""
    return [p for p in doc.paragraphs if p.text.strip() and not p.style.name.startswith("Heading")]


def _heading_paragraphs(doc, level: int) -> list:
    return [p for p in doc.paragraphs if p.style.name == f"Heading {level}"]


def _dominant(paragraphs, getter):
    """Most common non-None getter(paragraph) value across paragraphs, or None."""
    counts: dict = {}
    for p in paragraphs:
        try:
            v = getter(p)
        except Exception:
            v = None
        if v is not None:
            counts[v] = counts.get(v, 0) + 1
    return max(counts, key=counts.get) if counts else None


def _line_spacing_value(ls, body_pt: float | None) -> float | None:
    """line_spacing (float 'multiple' e.g. 1.5, or Length 'exact/atLeast') → editor decimal multiple.

    Exact/atLeast (固定值/最小值, common in CN docs) has no clean multiple form, so we approximate
    as exact_pt ÷ body_font_pt — e.g. 固定值28pt at 14pt ≈ 2.0. Returns None when unset.
    """
    if isinstance(ls, float) and ls:
        return round(ls, 2)
    if ls is not None and body_pt:
        return round(ls.pt / body_pt, 2)
    return None


def _para_space_after_pt(p) -> int | None:
    sa = p.paragraph_format.space_after
    return int(sa.pt) if sa is not None else None


def _para_first_indent_chars(p, body_pt: float | None) -> int | None:
    fi = p.paragraph_format.first_line_indent
    if fi is None or not body_pt:
        return None
    return round(fi.pt / body_pt)  # pt indent ÷ pt body size ≈ 字 (char) indent


def _dominant_run_color(paragraphs) -> str | None:
    """Most common explicit run color (#RRGGBB) across paragraphs, or None."""
    counts: dict = {}
    for p in paragraphs:
        for run in p.runs:
            try:
                rgb = run.font.color.rgb
            except Exception:
                rgb = None
            if rgb is not None:
                key = f"#{rgb}"
                counts[key] = counts.get(key, 0) + 1
    return max(counts, key=counts.get) if counts else None


def _dominant_run_bold(paragraphs) -> bool | None:
    """Most common explicit run bold across paragraphs, or None (no run sets it)."""
    counts: dict = {}
    for p in paragraphs:
        for run in p.runs:
            try:
                b = run.font.bold
            except Exception:
                b = None
            if b is not None:
                counts[b] = counts.get(b, 0) + 1
    return max(counts, key=counts.get) if counts else None


_HEADING_CN_RE = re.compile(r"^第[一二三四五六七八九十百零\d]+章|^[一二三四五六七八九十]+[、.．]")
# 第三支 `^\d{1,3}\s{2,}\S` 覆盖中文报告常见的手写章号 "1  设计依据"（数字+≥2空格），
# 不会误伤 "2024 年" / "100 万元"（无 ≥2 空格分隔）。
_HEADING_DEC_RE = re.compile(r"^\d+[.\)）]|^\d+\.\d+|^\d{1,3}\s{2,}\S")

# Figure caption detection. CN report authors hand-type captions (图1-1 / 图1) rather
# than using Word's Caption style, so we match the text first and the style as a fallback.
_CAPTION_STYLE_RE = re.compile(r"caption|题注", re.I)
_CAPTION_TEXT_RE = re.compile(r"^(图|figure|fig\.?)\s*[\d一二三四五六七八九十]+", re.I)
_CAPTION_CHAPTER_RE = re.compile(r"^(图|figure|fig\.?)\s*\d+[-－—.]\d+", re.I)
_SOURCE_RE = re.compile(r"(数据来源|资料来源|图片来源|来源|source)", re.I)


def _heading_numbering(paragraphs) -> str:
    """Best-effort heading numbering style: 'chinese' | 'decimal' | 'none'.

    CN reports usually type the number into the heading text (第一章 / 一、 / 1. / 1.1) rather than
    using Word's list auto-numbering, so we check the text first, then fall back to a w:numPr
    (auto-numbered) check. ponytail: resolving numId→abstractNum→lvlText to distinguish decimal
    from 天干/字母 auto-formats would need numbering.xml parsing; auto-numbered ⇒ 'decimal'.
    """
    for p in paragraphs:
        text = p.text.strip()
        if _HEADING_CN_RE.search(text):
            return "chinese"
        if _HEADING_DEC_RE.search(text):
            return "decimal"
        try:
            pPr = p._p.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:numPr")) is not None:
                return "decimal"
        except Exception:
            pass
    return "none"


def _tbl_border_color(table) -> str | None:
    """First non-auto border color from w:tblBorders (representative of all edges)."""
    try:
        borders = table._tbl.tblPr.find(qn("w:tblBorders"))
        if borders is None:
            return None
        for edge in ("top", "left", "bottom", "right"):
            e = borders.find(qn(f"w:{edge}"))
            if e is not None:
                color = e.get(qn("w:color"))
                if color and color.lower() != "auto":
                    return f"#{color}"
    except Exception:
        pass
    return None


def _extract_body_styles(doc) -> dict:
    style = doc.styles["Normal"]
    pf = style.paragraph_format
    body = _body_paragraphs(doc)
    run_size, _ = _dominant_run_font(body)
    style_size = style.font.size.pt if style.font.size else None
    size = run_size or style_size
    family = _cjk_font(doc, style, body)

    # Real docs set spacing/indent per-paragraph, not on the Normal style — sample the
    # dominant paragraph value, then fall back to the style / a sensible default.
    line_spacing = _dominant(body, lambda p: _line_spacing_value(p.paragraph_format.line_spacing, size))
    if line_spacing is None:
        line_spacing = _line_spacing_value(pf.line_spacing, size) or 1.5
    space_after = _dominant(body, _para_space_after_pt)
    if space_after is None:
        sa = pf.space_after
        # 样例段落普遍未声明 space_after 且 pPrDefault 为空 → 段后距实为 0（密集正文），
        # 不是任意猜测值。仅当 Normal 样式显式定义时才取它。
        space_after = int(sa.pt) if sa else 0
    indent = _dominant(body, lambda p: _para_first_indent_chars(p, size))
    if indent is None:
        indent = 2  # ponytail: not derivable from style → default 2 字

    return {
        "fontFamily": family,
        "fontSize": int(size) if size else 12,
        "lineHeight": line_spacing,
        "paragraphSpacing": space_after,
        "firstLineIndent": indent,
    }


def _extract_heading_styles(doc) -> list[dict]:
    """Per-level heading styles, ONLY for levels the document actually uses.

    A level with zero Heading paragraphs is skipped — otherwise we'd emit Word's
    template-default style (often an English font like Cambria) for a level the sample
    never uses, polluting the editor with bogus values. ponytail: real docs format
    headings on runs too, so we read runs first and fall back to the Heading style.
    """
    defaults = {1: 16, 2: 14, 3: 12, 4: 12}
    out = []
    for level in range(1, 5):
        hps = _heading_paragraphs(doc, level)
        if not hps:
            continue  # 文档未使用该级标题 → 跳过,避免输出 Word 模板残留(如 Cambria)污染表单
        try:
            style = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        run_size, _ = _dominant_run_font(hps)
        style_size = style.font.size.pt if style.font.size else None
        size = run_size or style_size
        run_bold = _dominant_run_bold(hps)
        bold = run_bold if run_bold is not None else bool(style.font.bold)
        out.append(
            {
                "level": level,
                "fontFamily": _cjk_font(doc, style, hps),
                "fontSize": int(size) if size else defaults[level],
                "fontWeight": 700 if bold else 400,
                # 样例 H1 样式与 run 均未定义 color（Word auto），auto 渲染为黑色 #000000，
                # 不是猜测的深灰 #333333。
                "color": _dominant_run_color(hps) or _style_color(style, "#000000"),
                "numbering": _heading_numbering(hps),
            }
        )
    return out


def _shd_fill(shd) -> str | None:
    """w:shd element → fill hex string, or None when absent / 'auto' (no fill)."""
    if shd is None:
        return None
    fill = shd.get(qn("w:fill"))
    return fill if fill and fill.lower() != "auto" else None


def _cell_shading_fill(tc) -> str | None:
    tc_pr = tc.tcPr
    return _shd_fill(tc_pr.find(qn("w:shd")) if tc_pr is not None else None)


def _style_conditional_shading(table, cond_type: str) -> str | None:
    """Fill from a table style's conditional format (firstRow / band1Row / ...), best-effort.

    Covers the common Word case where picking a built-in table style bands the header
    or zebra-stripes rows without writing per-cell shading. ponytail: ignores w:tblLook
    (assumes the conditional applies); honor tblLook if a sample whose band is disabled
    mismatches.
    """
    try:
        style = table.style
        if style is None:
            return None
        for tsp in style.element.findall(qn("w:tblStylePr")):
            if tsp.get(qn("w:type")) != cond_type:
                continue
            tc_pr = tsp.find(qn("w:tcPr"))
            fill = _shd_fill(tc_pr.find(qn("w:shd")) if tc_pr is not None else None)
            if fill:
                return fill
    except Exception:
        pass
    return None


def _header_text_color(table) -> str | None:
    """First header cell's first explicitly-colored run, best-effort."""
    try:
        for para in table.rows[0].cells[0].paragraphs:
            for run in para.runs:
                try:
                    rgb = run.font.color.rgb
                except Exception:
                    rgb = None
                if rgb is not None:
                    return f"#{rgb}"
    except Exception:
        pass
    return None


def _extract_table_styles(doc) -> dict | None:
    if not doc.tables:
        return None
    table = doc.tables[0]
    header_bg = None
    try:
        header_bg = _cell_shading_fill(table.rows[0].cells[0]._tc)
    except Exception:
        header_bg = None
    if not header_bg:
        header_bg = _style_conditional_shading(table, "firstRow")
    header_color = _header_text_color(table)
    # Zebra striping only when the table style defines row banding; plain tables → off.
    stripe = bool(_style_conditional_shading(table, "band1Row") or _style_conditional_shading(table, "band2Row"))
    return {
        # No fill detected → white (no-fill), never an invented blue. Real fill comes
        # from direct cell shading or the table style's firstRow band above.
        "headerBg": f"#{header_bg}" if header_bg else "#FFFFFF",
        "headerColor": header_color or "#333333",
        "borderColor": _tbl_border_color(table) or "#CCCCCC",
        "stripeRows": stripe,
    }


def _extract_header_footer(doc) -> dict:
    sections = doc.sections

    def _text(part) -> str:
        if part is None:
            return ""
        # EAI-CUSTOM: 链接到上一节/空页眉页脚的 part 访问 .paragraphs 会抛
        # AttributeError: 'Part' object has no attribute 'element'（python-docx 惰性加载
        # _get_or_add_definition()），导致 import-layout 500。读不到就当空文本处理。
        try:
            return " ".join(p.text.strip() for p in part.paragraphs if p.text.strip())
        except Exception:
            return ""

    def _has_page_field(part) -> bool:
        if part is None:
            return False
        try:
            return "PAGE" in part._element.xml
        except Exception:
            return False

    def _has_image(part) -> bool:
        if part is None:
            return False
        try:
            return bool(part._element.findall(f".//{{{_DRAWML}}}blip"))
        except Exception:
            return False

    # A multi-section report commonly leaves the cover/TOC sections blank and only
    # fills the body sections, so scan every section: take the first non-empty header/
    # footer, and turn on page-number/logo if ANY section shows them.
    header_text = next((_text(s.header) for s in sections if _text(s.header)), "")
    footer_text = next((_text(s.footer) for s in sections if _text(s.footer)), "")
    return {
        "headerText": header_text,
        "footerText": footer_text,
        "showPageNumber": any(_has_page_field(s.footer) or _has_page_field(s.header) for s in sections),
        "showLogo": any(_has_image(s.header) for s in sections),
    }


def _para_has_image(para) -> bool:
    return bool(para._p.findall(f".//{{{_DRAWML}}}blip"))


def _para_align(para) -> str:
    if para.alignment in (None, WD_ALIGN_PARAGRAPH.LEFT):
        return "left"
    if para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        return "right"
    return "center"


def _detect_cover(doc) -> dict | None:
    """Best-effort cover detection from the first page.

    Returns None (fallback C) when no cover-like structure is found — the
    caller then leaves the cover section untouched rather than guessing.
    """
    section = doc.sections[0]
    try:
        different_first = bool(section.different_first_page_header_footer)
    except Exception:
        different_first = False

    pre: list = []
    has_toc = False
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            break
        if not p.text.strip():
            continue
        # 目录页(TOC)不是封面:toc 样式条目或"目录"/"contents"标题属于前置目录。
        if re.search(r"toc|目录|contents", p.style.name or "", re.I) or re.match(r"^目\s*录$|^contents$", p.text.strip(), re.I):
            has_toc = True
            continue
        pre.append(p)

    # 第一个 Heading 前是目录页 → 按无封面处理(回退 C,保留编辑器原状,不臆造封面)。
    if has_toc:
        return None

    if not different_first and len(pre) < 3:
        return None

    cover: dict = {
        "showLogo": False,
        "logoPosition": "center",
        "showTitle": False,
        "showClient": False,
        "showDate": False,
        "showProjectNumber": False,
    }
    if not pre:
        return cover

    first = pre[0]
    cover["showLogo"] = _para_has_image(first)
    cover["logoPosition"] = _para_align(first)

    title_para = max(pre, key=lambda p: p.runs[0].font.size.pt if p.runs and p.runs[0].font.size else 0)
    if title_para.runs and title_para.runs[0].font.size and title_para.runs[0].font.size.pt >= 14:
        cover["showTitle"] = True

    for p in pre:
        text = p.text
        if re.search(r"(建设单位|单位|公司|业主|client)", text, re.I):
            cover["showClient"] = True
        if re.search(r"(\d{4}[-/年]\d{1,2}[-/月]\d{0,2}日?|日期)", text):
            cover["showDate"] = True
        if re.search(r"(项目编号|编号|项目号|工程号|number)", text, re.I):
            cover["showProjectNumber"] = True
    return cover


def _style_id_sets(doc) -> tuple[set[str], set[str]]:
    """Precompute (heading_style_ids, toc_style_ids) from doc.styles."""
    heading_ids: set[str] = set()
    toc_ids: set[str] = set()
    try:
        for st in doc.styles:
            name = (getattr(st, "name", "") or "").lower()
            sid = getattr(st, "style_id", None)
            if not sid:
                continue
            # Match canonical "Heading N" and localized "标题 N" (Chinese Word);
            # else a localized first heading misses its boundary and the extractor
            # gathers body content as a false cover.
            if name.startswith("heading") or name.startswith("标题"):
                heading_ids.add(sid)
            if "toc" in name:
                toc_ids.add(sid)
    except (KeyError, AttributeError):
        # Realistic style-access failure modes only — a bare except would silently
        # degrade heading/toc detection to text-only (miss a field-based TOC boundary).
        pass
    return heading_ids, toc_ids


def _max_run_font_pt(p_el) -> float:
    """Largest <w:sz w:val=...> (half-points) among runs in a <w:p> element."""
    best = 0.0
    for r in p_el.findall(f"{{{_W}}}r"):
        rpr = r.find(f"{{{_W}}}rPr")
        if rpr is None:
            continue
        sz = rpr.find(f"{{{_W}}}sz")
        val = sz.get(f"{{{_W}}}val") if sz is not None else None
        if val and val.isdigit():
            best = max(best, int(val) / 2.0)
    return best


def _para_text(p_el) -> str:
    return "".join((t.text or "") for t in p_el.iter(f"{{{_W}}}t"))


# id → display label for colon-detected fields.
_COVER_LABELS = {
    "project_number": "项目编号",
    "archive_no": "档案号",
    "version": "版次",
    "certificate_no": "证书号",
    "client": "建设单位",
    "project_name": "项目名",
    "stage": "设计阶段",
}
# (label alternatives → slot id). A label is matched as its chars joined by \s*,
# which tolerates the sample's inter-character spacing (档 案 号 / 版    次) while
# still requiring the label to be contiguous-ish — 项目名 won't match across
# 项目名…项目编号 because \s* can't skip non-whitespace. Single shared mapping:
# BOTH the structured cover-element extraction (_slot_from_colon) and the legacy
# cover_master slot prefill (_prefill_cover_slots) consume it.
_COVER_COLON_LABELS: list[tuple[tuple[str, ...], str]] = [
    (("项目编号", "工程编号"), "project_number"),
    (("档案号",), "archive_no"),
    (("版次",), "version"),
    (("证书号", "资质证书号"), "certificate_no"),
    (("建设单位", "业主单位", "业主"), "client"),
    (("项目名称", "工程名称"), "project_name"),
    (("设计阶段", "阶段"), "stage"),
    (("日期",), "date"),
]
_COVER_COLON_DEFAULT_FROM: dict[str, str | None] = {
    "project_number": None,
    "archive_no": None,
    "version": None,
    "certificate_no": None,
    "client": "frontmatter:client",
    "project_name": None,
    "stage": None,
    "date": "today",
}
# Date / report-title patterns — named once, referenced by both the legacy
# cover_master prefill and the structured cover-element extraction.
_COVER_DATE_RE = re.compile(r"20\d{2}年\d{1,2}月|20XX年0X月|\d{4}[-/年]\d{1,2}[-/月]\d{0,2}日?")
_COVER_TITLE_CN_RE = re.compile(r"第[\d一二三四五六七八九十百两]+\s*[一-龥A-Za-z0-9 ]{0,20}?(?:专篇|报告书|计算书|设计说明)")
_COVER_TITLE_GEN_RE = re.compile(r"[一-龥A-Za-z0-9 ]{2,24}?(?:专篇|报告书|计算书|设计说明)")
_COVER_TITLE_RE = re.compile(_COVER_TITLE_CN_RE.pattern + "|" + _COVER_TITLE_GEN_RE.pattern)


def _prefill_cover_slots(cover_blocks) -> list[dict]:
    """Detect cover fields as editable slots by scanning cover-region text.

    Colon fields carry a label-inclusive ``target`` (e.g. ``"档 案 号：XX"``) so
    duplicate sample values (XX×3 across 项目编号/档案号/证书号) bind unambiguously
    at generation. camelCase keys."""
    paras: list[tuple] = []  # (p_el, text) incl. paragraphs inside tables
    for b in cover_blocks:
        for p in b.iter(f"{{{_W}}}p"):
            txt = _para_text(p)
            if txt.strip():
                paras.append((p, txt))
    full = "\n".join(t for _, t in paras)

    slots: list[dict] = []
    seen: set[str] = set()

    def add(slot_id: str, label: str, value, *, default_from: str | None = None, target: str | None = None, kind: str = "variable") -> None:
        if not value or slot_id in seen:
            return
        seen.add(slot_id)
        val = str(value).strip()
        slots.append({"id": slot_id, "label": label, "kind": kind, "sampleValue": val, "defaultFrom": default_from, "target": (target or val).strip()})

    # 1) Colon fields: LABEL：VALUE (whitespace-tolerant label, value = next token).
    #    Each slot's first matching label wins (add() dedupes by slot id), matching
    #    the previous flat _COVER_COLON_FIELDS ordering exactly.
    for labels, slot_id in _COVER_COLON_LABELS:
        default_from = _COVER_COLON_DEFAULT_FROM.get(slot_id)
        for label_chars in labels:
            m = re.search(r"\s*".join(label_chars) + r"\s*[：:]\s*(?P<val>[^\s：:\n]+)", full)
            if m:
                add(slot_id, _COVER_LABELS.get(slot_id, label_chars), m.group("val"), default_from=default_from, target=m.group(0))
                break

    # 2) Title: prefer a report-name keyword (第…册…专篇 / …专篇) over largest-font,
    #    so a large-font stage value (基础设计) isn't mistaken for the report title.
    title = None
    m = _COVER_TITLE_CN_RE.search(full)
    if not m:
        m = _COVER_TITLE_GEN_RE.search(full)
    if m:
        title = m.group(0).strip()
    if not title:  # fallback: largest-font paragraph ≥ 2 chars
        best, best_sz = "", 0.0
        for p, txt in paras:
            sz = _max_run_font_pt(p)
            if sz > best_sz and len(txt.strip()) >= 2:
                best_sz, best = sz, txt.strip()
        title = best
    add("title", "报告名称", title, default_from="doc_title")

    # 3) 项目名 (positional, no colon) — only if not already a colon field.
    if "project_name" not in seen:
        m = re.search(r"项目名(?:称)?\s+(?P<val>[^\s：:\n]+)", full)
        if m:
            add("project_name", "项目名", m.group("val"))

    # 4) 设计院 / 设计单位 (label-only in sample, no value) → literal slot so the
    #    field is at least recognized; generation has no source to fill it.
    if "design_unit" not in seen and re.search(r"设计院|设计单位", full):
        add("design_unit", "设计单位", "设计院", kind="literal")

    # 5) Date placeholder (20XX年0X月) or real date — only if not a colon field.
    if "date" not in seen:
        m = _COVER_DATE_RE.search(full)
        if m:
            add("date", "日期", m.group(0), default_from="today")

    return slots


def _extract_cover_master(doc, source_file: str = "") -> dict | None:
    """Extract the cover region (blocks before the TOC marker or first Heading) as
    a reusable OOXML master + prefilled slots. Returns None when no meaningful
    cover exists (degrades to the legacy cover_template fallback)."""
    body = doc.element.body
    heading_ids, toc_ids = _style_id_sets(doc)

    cover_blocks: list = []
    boundary: str | None = None

    for child in body.iterchildren():
        tag = child.tag
        if tag == f"{{{_W}}}sectPr":  # final section properties — body content ended
            break
        if tag == f"{{{_W}}}p":
            style_el = child.find(f"{{{_W}}}pPr/{{{_W}}}pStyle")
            style_val = style_el.get(f"{{{_W}}}val") if style_el is not None else None
            text = _para_text(child).strip()
            if _TOC_TEXT_RE.match(text) or style_val in toc_ids:
                boundary = "before_toc"
                break
            if style_val in heading_ids:
                boundary = "before_first_heading"
                break
            cover_blocks.append(child)
        elif tag == f"{{{_W}}}tbl":
            cover_blocks.append(child)
        # other elements (bookmarkStart, etc.) ignored

    if boundary is None:
        return None

    has_table = any(b.tag == f"{{{_W}}}tbl" for b in cover_blocks)
    has_text = any(_para_text(b).strip() for b in cover_blocks if b.tag == f"{{{_W}}}p")
    if not has_table and not has_text:
        return None

    xml = "".join(etree.tostring(deepcopy(b), encoding="unicode") for b in cover_blocks)

    images: list[dict] = []
    seen: set[str] = set()
    for b in cover_blocks:
        for blip in b.iter(f"{{{_DRAWML}}}blip"):
            rid = blip.get(f"{{{_REL}}}embed")
            if not rid or rid in seen:
                continue
            seen.add(rid)
            part = doc.part.related_parts.get(rid)
            blob = getattr(part, "blob", None)
            if not blob:
                continue
            ext = "png"
            partname = getattr(part, "partname", None)
            if partname and "." in str(partname):
                ext = str(partname).rsplit(".", 1)[-1].lower()
            images.append({"origRid": rid, "ext": ext, "b64": base64.b64encode(blob).decode("ascii")})

    return {
        "mode": "master",
        "xml": xml,
        "images": images,
        "slots": _prefill_cover_slots(cover_blocks),
        "sourceFile": source_file,
        "boundary": boundary,
    }


# ---------------------------------------------------------------------------
# Cover elements extraction — structured multi-page editor model (replaces the
# cover_master OOXML passthrough as the primary extraction path; cover_master
# stays as legacy fallback). Each page = elements (text/table/spacer).
# ---------------------------------------------------------------------------
# Monotonic element-id counter: ids must be unique per extraction (T8 editor uses
# id as React key / patchCoverElementsPage match key). Content-hash ids collide on
# duplicated cover text (banner repeats), so a plain counter is used instead.
# 共享的冒号标签映射与日期/标题正则见上方 _COVER_COLON_LABELS / _COVER_DATE_RE /
# _COVER_TITLE_RE。
_ELEM_COUNTER = itertools.count(1)


def _para_style(p_el) -> dict:
    """Best-effort paragraph style: alignment + first run font props (pt)."""
    jc = p_el.find(f"{{{_W}}}pPr/{{{_W}}}jc")
    alignment = {"left": "left", "center": "center", "right": "right"}.get(jc.get(f"{{{_W}}}val") if jc is not None else None, "left")
    fontFamily, fontSize, bold, color = "宋体", 12, False, "#000000"
    r = p_el.find(f"{{{_W}}}r")
    if r is not None:
        rPr = r.find(f"{{{_W}}}rPr")
        if rPr is not None:
            rf = rPr.find(f"{{{_W}}}rFonts")
            if rf is not None:
                fontFamily = rf.get(f"{{{_W}}}eastAsia") or rf.get(f"{{{_W}}}ascii") or "宋体"
            sz = rPr.find(f"{{{_W}}}sz")
            if sz is not None:
                try:
                    fontSize = int(sz.get(f"{{{_W}}}val", "24")) // 2
                except ValueError:
                    fontSize = 12
            b = rPr.find(f"{{{_W}}}b")
            if b is not None:
                bval = b.get(f"{{{_W}}}val")
                bold = bval not in ("0", "false", "off")  # w:b w:val=0/false/off → 非粗体
            c = rPr.find(f"{{{_W}}}color")
            if c is not None:
                val = c.get(f"{{{_W}}}val")
                if val and val.lower() != "auto":  # w:color w:val=auto → 保持默认黑色
                    color = "#" + val
    # 段落间距 w:spacing (twips → pt): 保留样张的垂直节奏; auto 间距跳过。
    spacing = p_el.find(f"{{{_W}}}pPr/{{{_W}}}spacing")
    space_before = space_after = 0
    if spacing is not None and spacing.get(f"{{{_W}}}beforeAutospacing") != "1" and spacing.get(f"{{{_W}}}afterAutospacing") != "1":
        try:
            space_before = int(spacing.get(f"{{{_W}}}before") or 0) // 20
            space_after = int(spacing.get(f"{{{_W}}}after") or 0) // 20
        except ValueError:
            pass
    return {"fontFamily": fontFamily, "fontSize": fontSize, "bold": bold, "color": color, "alignment": alignment, "spaceBefore": space_before, "spaceAfter": space_after}


def _slot_from_colon(text: str) -> str | None:
    """'项目编号：XX' → project_number; 匹配任一标签 → 对应 slot id.

    冒号字段元素在实践里单段单字段;即使文本含多个冒号,这里绑定取第一个匹配标签,
    生成端按最后一个冒号替换值部分(见 _render_cover_elements)——两者通常一致,无需消歧。
    """
    for labels, sid in _COVER_COLON_LABELS:
        for lab in labels:
            if re.search(r"\s*".join(lab) + r"\s*[：:]", text):
                return sid
    return None


def _image_element(doc, el, images_map: dict | None = None) -> dict | None:
    """If the paragraph embeds a drawing image, return an image element dict (else None).

    Reads the first <a:blip> r:embed → related part blob → base64 + extension.
    Cover logos are standalone (empty-text) paragraphs; a paragraph with both text
    and an image is treated as text elsewhere (image ignored).

    When ``doc`` is None (migration from a stored cover_master), the image blobs are
    recovered from ``images_map`` ({origRid: {"b64", "ext"}}) instead — they were
    collected by ``_extract_cover_master`` and stored in ``master["images"]``.
    """
    blips = list(el.iter(f"{{{_DRAWML}}}blip"))
    if not blips:
        return None
    rid = blips[0].get(f"{{{_REL}}}embed")
    if not rid:
        return None
    if doc is not None:
        part = doc.part.related_parts.get(rid)
        blob = getattr(part, "blob", None) if part is not None else None
        if not blob:
            return None
        ext = "png"
        partname = getattr(part, "partname", None)
        if partname and "." in str(partname):
            ext = str(partname).rsplit(".", 1)[-1].lower()
        return {
            "id": f"img{next(_ELEM_COUNTER)}",
            "type": "image",
            "image": {"b64": base64.b64encode(blob).decode("ascii"), "ext": ext},
        }
    # 迁移: 图片来自存储的 master["images"] (按 origRid), 而非 live doc
    img = (images_map or {}).get(rid)
    if not img or not img.get("b64"):
        return None
    return {
        "id": f"img{next(_ELEM_COUNTER)}",
        "type": "image",
        "image": {"b64": img["b64"], "ext": img.get("ext", "png")},
    }


def _block_to_element(doc, el, images_map: dict | None = None) -> dict:
    """Convert a body block (<w:p>|<w:tbl>) to a CoverElementSchema dict.

    Returns a dict fed to ``CoverElementSchema(**el)`` by the caller. Blocks that
    fail are degraded to a spacer by the caller (spec §8 block-level try/except).

    ``images_map`` is threaded to ``_image_element`` for the migration path where
    ``doc`` is None (logos recovered from stored ``master["images"]``)."""
    if el.tag == f"{{{_W}}}tbl":
        rows_el = el.findall(f"{{{_W}}}tr")
        cells: list[list[str]] = []
        for tr in rows_el:
            row = [" ".join(_para_text(p).strip() for p in tc.iter(f"{{{_W}}}p") if _para_text(p).strip()) for tc in tr.iter(f"{{{_W}}}tc")]
            cells.append(row)
        # 真实列数取自 w:tblGrid 的 gridCol 数(物理行 tc 数受 gridSpan/vMerge 影响,
        # 会签表实为 6 列但首行只有 3 个 tc)。无 tblGrid 时回退到各行最大 cell 数。
        grid = el.find(f"{{{_W}}}tblGrid")
        if grid is not None:
            cols = len(grid.findall(f"{{{_W}}}gridCol"))
        else:
            cols = max((len(r) for r in cells), default=0)
        for row in cells:
            row.extend([""] * (cols - len(row)))
        return {
            "id": f"tbl{next(_ELEM_COUNTER)}",
            "type": "table",
            "rows": len(cells),
            "cols": cols,
            "cells": cells,
            "borderColor": "#000000",
        }
    text = _para_text(el).strip()
    if text:
        style = _para_style(el)
        # 冒号字段标签的装饰性字间空格(工  程  编  号)在编辑/生成替换时是噪声 → 归一化,
        # 同时让"工程编号"这类标签可被生成端按标签精确替换。
        sid = _slot_from_colon(text)
        if sid:
            text = re.sub(r"\s+", "", text)
        el_dict = {"id": f"el{next(_ELEM_COUNTER)}", "type": "text", "text": text, **style}
        if sid:
            el_dict["slotId"] = sid
        elif text.strip() in ("项目名", "项目名称"):
            el_dict["slotId"] = "project_name"
        elif _COVER_TITLE_RE.search(text) and el_dict.get("fontSize", 12) >= 16:
            el_dict["slotId"] = "title"
        elif _COVER_DATE_RE.search(text):
            el_dict["slotId"] = "date"
        return el_dict
    # 空段:独立 logo 图片 → image 元素;否则 spacer。
    img = _image_element(doc, el, images_map)
    if img is not None:
        return img
    return {"id": f"sp{next(_ELEM_COUNTER)}", "type": "spacer", "lines": 1}


def _table_has_cover_fields(tbl) -> bool:
    """A cover-field/layout table (banner + colon fields) vs a real content table.

    Layout tables hold the cover metadata (项目编号：XX / 第…册…专篇 / 20XX年0X月 /
    项目名) inside their cells → decompose them into editable text elements so the
    fields bind as slots. Content tables (会签表 / 名单表) carry no such fields →
    keep as a single table element.
    """
    for tc in tbl.iter(f"{{{_W}}}tc"):
        for p in tc.iter(f"{{{_W}}}p"):
            t = _para_text(p).strip()
            if not t:
                continue
            if _slot_from_colon(t) or _COVER_TITLE_RE.search(t) or _COVER_DATE_RE.search(t) or t in ("项目名", "项目名称"):
                return True
    return False


def _table_cell_elements(doc, tbl, images_map: dict | None = None) -> list[dict]:
    """Decompose a layout table into its cell paragraphs as text/image elements."""
    out: list[dict] = []
    for tc in tbl.iter(f"{{{_W}}}tc"):
        for p in tc.iter(f"{{{_W}}}p"):
            try:
                el = _block_to_element(doc, p, images_map)
            except Exception:
                el = {"id": f"sp{next(_ELEM_COUNTER)}", "type": "spacer", "lines": 1}
            # 保留 spacer: 封面表内的空段(如消防 banner 内部空行)是样张垂直节奏的一部分。
            if el["type"] in ("text", "image", "spacer"):
                out.append(el)
    return out


def _blocks_to_pages(doc, blocks, images_map: dict | None = None) -> list[dict]:
    """body blocks → 页列表(按分节符/分页符切页——段落带不带文本都切, cover-field 表分解为元素).

    供 `_extract_cover_pages`(从 live doc 收集 blocks) 与 `_cover_master_to_elements`
    (从存储的 OOXML 片段解析 blocks) 共用 —— spec §7: 迁移与提取同源。
    ``doc`` 可为 None(迁移): 图片从 ``images_map`` ({origRid: {"b64","ext"}}) 恢复。
    返回 [{elements: [dict...]}] (未包 CoverPageSchema; 全 spacer 页丢弃)。
    """
    pages: list[dict] = [{"elements": []}]
    for child in blocks:
        tag = child.tag
        if tag == f"{{{_W}}}p":
            has_sect = child.find(f"{{{_W}}}pPr/{{{_W}}}sectPr") is not None
            has_pgbr = any(br.get(f"{{{_W}}}type") == "page" for br in child.iter(f"{{{_W}}}br"))
            # 非空分节符/分页符段落 → 新页;空的分节符段落多为封面内布局标记
            # (标题横幅表后)不切页——消防样例借此保持单页封面。
            try:
                el = _block_to_element(doc, child, images_map)
            except Exception:
                el = {"id": f"sp{next(_ELEM_COUNTER)}", "type": "spacer", "lines": 1}
            # 分节符/分页符段落(无论是否带文本)都是页边界: 先开新页, 再把该段落元素
            # (空段通常是 spacer)放进新页, 保留原结构。此前仅"非空段"才切页 → 消防样例
            # 标题横幅表后/目录前的空段分节符与分页符被忽略, banner 与会签表被并进一页。
            if has_sect or has_pgbr:
                pages.append({"elements": []})
            pages[-1]["elements"].append(el)
        elif tag == f"{{{_W}}}tbl":
            if _table_has_cover_fields(child):
                pages[-1]["elements"].extend(_table_cell_elements(doc, child, images_map))
            else:
                try:
                    el = _block_to_element(doc, child, images_map)
                except Exception:
                    el = {"id": f"sp{next(_ELEM_COUNTER)}", "type": "spacer", "lines": 1}
                pages[-1]["elements"].append(el)
    return [p for p in pages if any(e["type"] != "spacer" for e in p["elements"])]


def _extract_cover_pages(doc) -> list[CoverPageSchema]:
    """封面区(目录/首个Heading前) → 按分节符/分页符切页(段落文本可有可无) → 每页元素列表.

    返回 list[CoverPageSchema];T4 持久化逐页 `.model_dump()`,T5
    `_cover_master_to_elements` 产出同一 schema 形状。`_block_to_element` 返回
    dict,由这里包成 CoverElementSchema 实例;单块转换失败降级为 spacer(§8),
    绝不中断导入。块转换/切页/表格分解委托共享 helper `_blocks_to_pages`。
    """
    body = doc.element.body
    heading_ids, toc_ids = _style_id_sets(doc)
    blocks: list = []
    for child in body:
        tag = child.tag
        if tag == f"{{{_W}}}sectPr":
            break
        if tag == f"{{{_W}}}p":
            style_el = child.find(f"{{{_W}}}pPr/{{{_W}}}pStyle")
            style_val = style_el.get(f"{{{_W}}}val") if style_el is not None else None
            text = _para_text(child).strip()
            if _TOC_TEXT_RE.match(text) or style_val in toc_ids or style_val in heading_ids:
                break
            blocks.append(child)
        elif tag == f"{{{_W}}}tbl":
            blocks.append(child)
    pages = _blocks_to_pages(doc, blocks)
    return [CoverPageSchema(elements=[CoverElementSchema(**e) for e in p["elements"]]) for p in pages]


def _cover_master_to_elements(master: dict | None) -> dict | None:
    """旧 cover_master（OOXML 片段）→ 元素模型。失败返回 None（保留旧母版）。

    Task 5 迁移：读取模板时若 cover_master 存在而 cover_elements 为空，把存储的
    封面 OOXML 片段解析回 w:p / w:tbl 块，复用共享 helper ``_blocks_to_pages``
    （与 _extract_cover_pages 同源，spec §7）切页 + 表格分解 + 图片恢复：
    - 多页母版（如环评 3 页）保留页数；
    - banner/字段表按 _table_has_cover_fields 分解为可绑定文本元素；
    - logo 从 ``master["images"]``（按 origRid 匹配 <a:blip r:embed>）恢复 b64，
      不再是 spacer。
    转换/解析失败或无元素 → None，让调用方保留旧母版，绝不破坏已有模板。
    """
    if not master or not master.get("xml"):
        return None
    try:
        from lxml import etree as _et

        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        root = _et.fromstring(f'<root xmlns:w="{W}">{master["xml"]}</root>')
        blocks = [b for b in list(root) if b.tag in (f"{{{W}}}p", f"{{{W}}}tbl")]
        images_map = {img["origRid"]: {"b64": img.get("b64", ""), "ext": img.get("ext", "png")} for img in (master.get("images") or []) if img.get("origRid")}
        pages = _blocks_to_pages(None, blocks, images_map)
        if not pages:
            return None
        return {"mode": "elements", "pages": pages, "sourceFile": master.get("sourceFile", "")}
    except Exception:
        return None


def _is_caption(para) -> bool:
    """A figure-caption paragraph: Caption/题注 style, or text like 图1-1 / Figure 2."""
    if _CAPTION_STYLE_RE.search(para.style.name or ""):
        return True
    return bool(_CAPTION_TEXT_RE.match(para.text.strip()))


def _extract_figure_styles(doc) -> dict | None:
    """Best-effort figure-style detection: caption position / numbering / source line.

    Returns None when the sample has neither image paragraphs nor caption text — the
    caller then leaves the figure section untouched (don't clobber user config).
    ponytail: caption adjacency heuristics (图1-1 → chapter, 图1 → continuous,
    image-then-caption → below) match how CN report authors hand-type captions; honor
    tblLook/abstractNum only if a sample mismatches these defaults.
    """
    paragraphs = doc.paragraphs
    cap_indices = {i for i, p in enumerate(paragraphs) if _is_caption(p)}
    image_indices = [i for i, p in enumerate(paragraphs) if _para_has_image(p)]
    if not cap_indices and not image_indices:
        return None

    # Caption position: a caption right after an image → below; right before → above.
    below = sum(1 for i in image_indices if (i + 1) in cap_indices)
    above = sum(1 for i in image_indices if (i - 1) in cap_indices)
    caption_position = "above" if above > below else "below"

    # Numbering: a chapter-segmented caption (图1-1) → chapter, else continuous (图1).
    numbering = "continuous"
    for i in cap_indices:
        if _CAPTION_CHAPTER_RE.match(paragraphs[i].text.strip()):
            numbering = "chapter"
            break

    show_source = any(_SOURCE_RE.search(p.text) for p in paragraphs)

    return {
        "captionPosition": caption_position,
        "numbering": numbering,
        "showSource": show_source,
    }


def extract_layout_from_docx(data: bytes, source_file: str = "") -> dict:
    """Parse a .docx byte stream → LayoutTemplate data subset (snake_case).

    Raises ValueError for non-.docx / unparseable input.
    """
    try:
        doc = Document(BytesIO(data))
    except Exception as exc:
        raise ValueError("无法解析该文件，请确保为 .docx 格式") from exc

    section = doc.sections[0]
    paper, orientation = _paper_from_dimensions(_to_cm(section.page_width), _to_cm(section.page_height))
    cover = _detect_cover(doc)
    cover_master = _extract_cover_master(doc, source_file=source_file)
    cover_pages = _extract_cover_pages(doc)
    cover_elements = {"mode": "elements", "pages": [p.model_dump() for p in cover_pages], "sourceFile": source_file} if cover_pages else None

    return {
        "page_settings": {
            "paperSize": paper,
            "orientation": orientation,
            "marginTop": _to_cm(section.top_margin),
            "marginBottom": _to_cm(section.bottom_margin),
            "marginLeft": _to_cm(section.left_margin),
            "marginRight": _to_cm(section.right_margin),
        },
        "body_styles": _extract_body_styles(doc),
        "heading_styles": _extract_heading_styles(doc),
        "table_styles": _extract_table_styles(doc),
        "figure_styles": _extract_figure_styles(doc),
        "header_footer": _extract_header_footer(doc),
        "cover_template": cover,
        "cover_master": cover_master,
        "cover_elements": cover_elements,
        "cover_detected": cover_master is not None or cover is not None or cover_elements is not None,
    }


def validate_docx_upload(filename: str | None, data: bytes) -> dict:
    """Validate an .docx upload (extension + ≤10MB) and return extracted layout.

    Raises ValueError with a user-facing Chinese message on invalid input; the
    routers map it to HTTP 400. Kept in this pure module so both the output and
    docmgr import-layout endpoints share one implementation.
    """
    if not filename or not filename.lower().endswith(".docx"):
        raise ValueError("仅支持 .docx 文件")
    if len(data) > 10 * 1024 * 1024:
        raise ValueError("文件不能超过 10MB")
    return extract_layout_from_docx(data, source_file=filename or "")
