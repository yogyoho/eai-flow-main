"""Markdown → DOCX generator using python-docx with layout template styling."""

from __future__ import annotations

import base64
import datetime
import logging
import re
from copy import deepcopy
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from lxml import etree

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Minimal markdown parser — produces a list of typed blocks
# ---------------------------------------------------------------------------

HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
LIST_UL_RE = re.compile(r"^(\s*)[-*+]\s+(.+)$")
LIST_OL_RE = re.compile(r"^(\s*)\d+\.\s+(.+)$")
HR_RE = re.compile(r"^(-{3,}|\*{3,}|_{3,})$")

FRONTMATTER_RE = re.compile(r"\A---\s*\n(.*?)\n---\s*\n?", re.S)


def _split_frontmatter(md: str) -> tuple[dict, str]:
    """Split leading ``---\\nkey: value\\n---\\n`` front-matter from markdown.

    Only flat ``key: value`` lines are supported (no nested YAML, no new dep).
    Returns ``(meta_dict, body_markdown)``. If there is no front-matter, or any
    non-blank/non-comment line lacks a colon (malformed), returns ``({}, md)``
    — i.e. treat the whole input as body so generation never crashes.
    """
    m = FRONTMATTER_RE.match(md)
    if not m:
        return {}, md
    meta: dict[str, str] = {}
    for line in m.group(1).splitlines():
        stripped = line.strip()
        if not stripped or stripped.startswith("#"):
            continue
        if ":" not in line:
            return {}, md  # malformed → degrade to whole-body
        key, _, val = line.partition(":")
        meta[key.strip()] = val.strip().strip('"').strip("'")
    return meta, md[m.end() :]


@dataclass
class Block:
    kind: str  # heading | paragraph | ul_item | ol_item | hr | code_block | table
    level: int = 0
    text: str = ""
    rows: list[list[str]] = field(default_factory=list)


def _compute_heading_numbers(blocks: list[Block], heading_styles: list[dict]) -> dict[int, str]:
    """Compute decimal multilevel numbers for heading blocks.

    Returns ``{block_index: number_string}`` (e.g. ``{0: "1", 1: "1.1", 3: "1.2.1"}``)
    for heading blocks whose level has ``numbering == "decimal"``. Headings on
    levels with ``numbering != "decimal"`` (or unspecified) are omitted.

    Note: meaningful only when every relevant level is "decimal". Mixing "none"
    in the middle produces counter-intuitive numbers for deeper levels —
    acceptable since templates configure all-or-none per the spec.
    """
    numbering_by_level: dict[int, str] = {hs.get("level", 0): hs.get("numbering", "none") for hs in heading_styles}
    counters = [0, 0, 0, 0]  # levels 1..4
    result: dict[int, str] = {}
    for i, b in enumerate(blocks):
        if b.kind != "heading":
            continue
        level = max(1, min(b.level, 4))
        if numbering_by_level.get(level, "none") != "decimal":
            continue
        counters[level - 1] += 1
        for deeper in range(level, 4):
            counters[deeper] = 0
        result[i] = ".".join(str(counters[k]) for k in range(level))
    return result


def _resolve_cover_fields(api_fields: dict, frontmatter: dict, blocks: list[Block]) -> dict:
    """Resolve cover field values by priority: API params > front-matter > fallback.

    - title: api > front-matter > first H1 block text
    - client/project_number: api > front-matter (omit if neither)
    - date: api > front-matter > today (ISO)
    """
    resolved: dict = {}
    title = api_fields.get("title") or frontmatter.get("title")
    if not title:
        for b in blocks:
            if b.kind == "heading" and b.level == 1:
                title = b.text
                break
    if title:
        resolved["title"] = title
    for key in ("client", "project_number"):
        val = api_fields.get(key) or frontmatter.get(key)
        if val:
            resolved[key] = val
    date_val = api_fields.get("date") or frontmatter.get("date")
    resolved["date"] = date_val or datetime.date.today().isoformat()
    return resolved


def _render_cover(doc, cover_template: dict | None, cover_fields: dict) -> None:
    """Render the cover page at the start of the document body.

    ``cover_template`` toggles which fields appear (showLogo/showTitle/showClient/
    showDate/showProjectNumber). ``cover_fields`` carries resolved values
    (title/client/date/project_number); a line is rendered only if its toggle is
    on AND its value is present. No-op content-wise when ``cover_template`` is falsy.
    """
    ct = cover_template or {}

    if ct.get("showLogo"):
        p = doc.add_paragraph()
        # EAI-CUSTOM: honor logoPosition (left/center/right); center is the legacy default.
        p.alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }.get(ct.get("logoPosition"), WD_ALIGN_PARAGRAPH.CENTER)
        run = p.add_run("[编制单位 LOGO]")
        _set_run_font(run, "宋体")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    for _ in range(3):  # vertical spacing before title
        doc.add_paragraph()

    if ct.get("showTitle") and cover_fields.get("title"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(cover_fields["title"]))
        _set_run_font(run, "黑体")
        run.font.size = Pt(22)
        run.bold = True

    for _ in range(4):  # spacing before info lines
        doc.add_paragraph()

    def _info_line(label: str, value) -> None:
        if not value:
            return
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}:{value}")
        _set_run_font(run, "宋体")
        run.font.size = Pt(14)

    if ct.get("showClient"):
        _info_line("建设单位", cover_fields.get("client"))
    if ct.get("showProjectNumber"):
        _info_line("项目编号", cover_fields.get("project_number"))
    if ct.get("showDate"):
        _info_line("日期", cover_fields.get("date"))


_W_GEN = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_DRAWML_GEN = "http://schemas.openxmlformats.org/drawingml/2006/main"
_REL_GEN = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def _replace_target_in_para(p_el, target: str, replacement: str) -> bool:
    """Replace `target` with `replacement` across all <w:r>/<w:t> of a <w:p>,
    keeping the first run's <w:rPr> and collapsing text into the first run.
    ponytail: 段落级替换，段落内混合格式（多字体）会丢失——封面槽位通常是
    整行/整格，影响可忽略；升级路径=按 run 边界细粒度替换。"""
    # <w:t> is a grandchild of <w:p> (nested in <w:r>), so use descendant axis.
    t_els = p_el.findall(f".//{{{_W_GEN}}}t")
    full = "".join((t.text or "") for t in t_els)
    if not target or target not in full:
        return False
    runs = p_el.findall(f"{{{_W_GEN}}}r")
    if not runs:
        return False
    first = runs[0]
    for run in runs:  # clear all w:t (keeps rPr/other children)
        for t in run.findall(f"{{{_W_GEN}}}t"):
            run.remove(t)
    new_text = full.replace(target, replacement)
    new_t = etree.SubElement(first, f"{{{_W_GEN}}}t")
    new_t.text = new_text
    new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return True


# EAI-CUSTOM: the resolvable cover slot ids. Mirrored one-for-one by frontend
# cover-state.ts COVER_RESOLVABLE_SLOT_IDS (the UI locks any slot outside this
# set to literal); backend test test_output_cover.py::test_cover_slot_value_keys
# pins this set — keep it in sync with the slot_value dict below.
COVER_SLOT_VALUE_KEYS = ("title", "client", "project_number", "date", "project_name", "stage")


def _render_cover_master(doc, master: dict, resolved: dict, frontmatter: dict) -> None:
    """Inject the cover-master OOXML fragment at the body start, replacing variable
    slots with resolved project values and re-embedding base64 images."""
    if not master.get("xml"):
        return  # nothing to inject; also avoids persisting dangling image rels (I2)
    root = etree.fromstring(f'<root xmlns:w="{_W_GEN}" xmlns:a="{_DRAWML_GEN}" xmlns:r="{_REL_GEN}">{master.get("xml", "")}</root>')

    # keys == COVER_SLOT_VALUE_KEYS (see module doc above)
    slot_value = {
        "title": resolved.get("title"),
        "client": resolved.get("client"),
        "project_number": resolved.get("project_number"),
        "date": resolved.get("date"),
        "project_name": frontmatter.get("project_name"),
        "stage": frontmatter.get("stage"),
    }
    for slot in master.get("slots", []):
        if slot.get("kind") != "variable":
            continue
        repl = slot_value.get(slot.get("id"))
        sample = slot.get("sampleValue")
        # Label-inclusive target (e.g. "项目编号：XX") uniquely anchors colon fields
        # so duplicate sample values (XX×3) don't collide; absent → sampleValue.
        target = slot.get("target") or sample
        if not repl or not target or not sample or str(repl) == sample:
            continue
        # Swap just the value within the (possibly label-inclusive) target so the
        # label survives and a sibling slot sharing the same sample value is untouched.
        replacement = target.replace(sample, str(repl))
        for p_el in root.iter(f"{{{_W_GEN}}}p"):
            _replace_target_in_para(p_el, target, replacement)

    for img in master.get("images", []):
        orig_rid = img.get("origRid")
        blips = [b for b in root.iter(f"{{{_DRAWML_GEN}}}blip") if b.get(f"{{{_REL_GEN}}}embed") == orig_rid]
        if not blips:
            continue
        try:
            blob = base64.b64decode(img["b64"])
            # get_or_add_image returns (rId, Image) — rId already the relationship id.
            new_rid, _image = doc.part.get_or_add_image(BytesIO(blob))
            for blip in blips:
                blip.set(f"{{{_REL_GEN}}}embed", new_rid)
        except Exception as exc:
            # image must never abort cover generation — strip the orphan <w:drawing>
            # so the emitted doc references no missing rId (would trigger Word repair).
            logger.warning("cover image re-embed failed, stripping drawing: %s", exc)
            for blip in blips:
                drawing = blip.getparent()
                while drawing is not None and drawing.tag != f"{{{_W_GEN}}}drawing":
                    drawing = drawing.getparent()
                if drawing is not None and drawing.getparent() is not None:
                    drawing.getparent().remove(drawing)

    body = doc.element.body
    for child in reversed(list(root)):  # insert at 0 in reverse → original order
        body.insert(0, deepcopy(child))


def _render_cover_preset(doc, preset: dict | None, values: dict | None) -> bool:
    """Render a data-driven cover page from a preset layout.

    ``preset["elements"]`` is a list of:
      - ``{"type": "spacer", "lines": N}`` — N blank paragraphs (always rendered)
      - ``{"type": "text",  "field": X, ...}`` — a standalone value (e.g. the title)
      - ``{"type": "info",  "label": L, "field": X, ...}`` — renders ``"L：value"``
    A text/info element whose field value is missing is skipped entirely
    (no empty label line). Unknown element types are skipped. Returns False
    (no-op) when ``preset`` is falsy; True otherwise. Never raises on a bad preset.
    """
    if not preset:
        return False
    vals = values or {}
    align_map = {"left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT}
    for el in preset.get("elements", []):
        etype = el.get("type")
        if etype == "spacer":
            for _ in range(int(el.get("lines", 1))):
                doc.add_paragraph()
            continue
        if etype not in ("text", "info"):
            continue  # ponytail: unknown element type → skip, never crash
        value = vals.get(el.get("field"))
        if value is None or str(value).strip() == "":
            continue  # missing value → skip entire line
        p = doc.add_paragraph()
        p.alignment = align_map.get(el.get("align", "center"), WD_ALIGN_PARAGRAPH.CENTER)
        text = str(value) if etype == "text" else f"{el.get('label', '')}：{value}"
        run = p.add_run(text)
        _set_run_font(run, _resolve_font(el.get("font", "宋体")))
        run.font.size = Pt(el.get("size", 16 if etype == "text" else 14))
        if el.get("bold"):
            run.bold = True
    return True


def _add_toc_field(paragraph, max_depth: int) -> None:
    """Inject a native Word TOC field (TOC \\o "1-N" \\h \\z \\u) into paragraph XML."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def _fld(char_type: str) -> None:
        run = paragraph.add_run()
        el = OxmlElement("w:fldChar")
        el.set(qn("w:fldCharType"), char_type)
        if char_type == "begin":
            el.set(qn("w:dirty"), "true")
        run._element.append(el)

    _fld("begin")
    run_instr = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' TOC \\o "1-{max_depth}" \\h \\z \\u '
    run_instr._element.append(instr)
    _fld("separate")

    run_placeholder = paragraph.add_run("（打开文档后右键“更新域”生成目录）")
    _set_run_font(run_placeholder, "宋体")
    run_placeholder.font.size = Pt(10)
    run_placeholder.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    _fld("end")


def _render_toc(doc, toc_settings: dict | None) -> bool:
    """Add a 目录 heading + TOC field. Returns True if rendered, False if skipped."""
    if not toc_settings:
        return False
    max_depth = toc_settings.get("maxDepth") or 0
    if max_depth <= 0:
        return False
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("目录")
    _set_run_font(run, "黑体")
    run.font.size = Pt(16)
    run.bold = True
    _add_toc_field(doc.add_paragraph(), max_depth)
    return True


def _set_update_fields(doc) -> None:
    """Write <w:updateFields w:val="true"/> into settings.xml so Word/WPS auto-updates the TOC on open."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        el = OxmlElement("w:updateFields")
        el.set(qn("w:val"), "true")
        settings.append(el)


def _set_section_pagenum(section, fmt: str | None = None, start: int | None = None) -> None:
    """Set pgNumType (page number format + restart) on a section's sectPr. Idempotent."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    sectPr = section._sectPr
    pgNum = sectPr.find(qn("w:pgNumType"))
    if pgNum is None:
        pgNum = OxmlElement("w:pgNumType")
        sectPr.append(pgNum)
    if fmt:
        pgNum.set(qn("w:fmt"), fmt)
    if start is not None:
        pgNum.set(qn("w:start"), str(start))


def _add_page_number_footer(section) -> None:
    """Add a centered PAGE field to the section's footer. Caller must unlink first."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _fld(char_type: str) -> None:
        run = footer_para.add_run()
        el = OxmlElement("w:fldChar")
        el.set(qn("w:fldCharType"), char_type)
        run._element.append(el)

    _fld("begin")
    run_instr = footer_para.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run_instr._element.append(instr)
    _fld("end")

    for run in footer_para.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def parse_markdown(md: str) -> list[Block]:
    """Parse markdown text into a flat list of blocks."""
    blocks: list[Block] = []
    lines = md.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Horizontal rule
        if HR_RE.match(line.strip()):
            blocks.append(Block(kind="hr"))
            i += 1
            continue

        # Heading
        m = HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            blocks.append(Block(kind="heading", level=level, text=m.group(2).strip()))
            i += 1
            continue

        # Fenced code block
        if line.strip().startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            blocks.append(Block(kind="code_block", text="\n".join(code_lines)))
            continue

        # Table
        if "|" in line and i + 1 < len(lines) and re.match(r"^\|?[\s\-:|]+\|?$", lines[i + 1].strip()):
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2  # skip separator
            table_rows = [header]
            while i < len(lines) and "|" in lines[i]:
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                table_rows.append(row)
                i += 1
            blocks.append(Block(kind="table", rows=table_rows))
            continue

        # Unordered list
        m = LIST_UL_RE.match(line)
        if m:
            indent = len(m.group(1))
            blocks.append(Block(kind="ul_item", level=indent, text=m.group(2).strip()))
            i += 1
            continue

        # Ordered list
        m = LIST_OL_RE.match(line)
        if m:
            indent = len(m.group(1))
            blocks.append(Block(kind="ol_item", level=indent, text=m.group(2).strip()))
            i += 1
            continue

        # Paragraph — accumulate until blank line
        para_lines: list[str] = []
        while i < len(lines) and lines[i].strip():
            para_lines.append(lines[i])
            i += 1
        blocks.append(Block(kind="paragraph", text=" ".join(para_lines)))

    return blocks


# ---------------------------------------------------------------------------
# Math: LaTeX → OMML (Word-native equations). lxml only, no new deps.
# Handles the common engineering subset (frac, sqrt, ^_, Greek, operators,
# sum/int). Unhandled LaTeX falls back to verbatim text so export never
# crashes. Mirrors frontend mathMarkdown.ts decode step.
# ---------------------------------------------------------------------------

from lxml import etree as _etree  # noqa: E402

_MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
_etree.register_namespace("m", _MATH_NS)
_etree.register_namespace("w", _WORD_NS)


def _m(tag: str) -> str:
    return f"{{{_MATH_NS}}}{tag}"


def _decode_math_placeholders(text: str) -> str:
    """Decode editor math placeholders back to ``$``/``$$`` (mirrors frontend decodeMath)."""
    from html import unescape

    def _norm(s: str) -> str:
        # collapse over-escaped backslashes from editor markdown round-trip (\\frac → \frac)
        s = unescape(s)
        return re.sub(r"\\{2,}([a-zA-Z])", r"\\\1", s)

    def blk(mo):
        return "$$" + _norm(mo.group(1)) + "$$"

    def inl(mo):
        return "$" + _norm(mo.group(1)) + "$"

    # tolerate data-math-x and data-math-x="", both attribute orders, and inner whitespace/newlines
    text = re.sub(r'<div\b[^>]*?\bdata-math-block\b[^>]*?\bdata-latex="([^"]*)"[^>]*>[\s\S]*?</div>', blk, text)
    text = re.sub(r'<div\b[^>]*?\bdata-latex="([^"]*)"[^>]*?\bdata-math-block\b[^>]*>[\s\S]*?</div>', blk, text)
    text = re.sub(r'<span\b[^>]*?\bdata-math-inline\b[^>]*?\bdata-latex="([^"]*)"[^>]*>[\s\S]*?</span>', inl, text)
    text = re.sub(r'<span\b[^>]*?\bdata-latex="([^"]*)"[^>]*?\bdata-math-inline\b[^>]*>[\s\S]*?</span>', inl, text)
    return text


_LATEX_SYMBOLS = {
    "alpha": "α",
    "beta": "β",
    "gamma": "γ",
    "delta": "δ",
    "epsilon": "ε",
    "varepsilon": "ε",
    "zeta": "ζ",
    "eta": "η",
    "theta": "θ",
    "vartheta": "ϑ",
    "iota": "ι",
    "kappa": "κ",
    "lambda": "λ",
    "mu": "μ",
    "nu": "ν",
    "xi": "ξ",
    "pi": "π",
    "rho": "ρ",
    "varrho": "ϱ",
    "sigma": "σ",
    "varsigma": "ς",
    "tau": "τ",
    "upsilon": "υ",
    "phi": "φ",
    "varphi": "φ",
    "chi": "χ",
    "psi": "ψ",
    "omega": "ω",
    "Gamma": "Γ",
    "Delta": "Δ",
    "Theta": "Θ",
    "Lambda": "Λ",
    "Xi": "Ξ",
    "Pi": "Π",
    "Sigma": "Σ",
    "Upsilon": "Υ",
    "Phi": "Φ",
    "Psi": "Ψ",
    "Omega": "Ω",
    "times": "×",
    "cdot": "·",
    "div": "÷",
    "pm": "±",
    "mp": "∓",
    "ast": "∗",
    "star": "⋆",
    "leq": "≤",
    "le": "≤",
    "geq": "≥",
    "ge": "≥",
    "neq": "≠",
    "ne": "≠",
    "approx": "≈",
    "equiv": "≡",
    "sim": "∼",
    "propto": "∝",
    "infty": "∞",
    "partial": "∂",
    "nabla": "∇",
    "sum": "∑",
    "int": "∫",
    "oint": "∮",
    "prod": "∏",
    "bigcup": "∪",
    "bigcap": "∩",
    "forall": "∀",
    "exists": "∃",
    "in": "∈",
    "notin": "∉",
    "subset": "⊂",
    "supset": "⊃",
    "cup": "∪",
    "cap": "∩",
    "emptyset": "∅",
    "rightarrow": "→",
    "to": "→",
    "Rightarrow": "⇒",
    "leftarrow": "←",
    "Leftarrow": "⇐",
    "leftrightarrow": "↔",
    "mapsto": "↦",
    "angle": "∠",
    "perp": "⊥",
    "parallel": "∥",
    "circ": "∘",
    "bullet": "•",
    "prime": "′",
    "dagger": "†",
    "cdots": "⋯",
    "ldots": "…",
    "dots": "…",
    "vdots": "⋮",
    "ddots": "⋱",
    "degree": "°",
    "hbar": "ℏ",
    "ell": "ℓ",
    "quad": " ",
    "qquad": "  ",
}

_LATEX_TEXTY = frozenset(
    {
        "mathrm",
        "mathbf",
        "mathit",
        "mathcal",
        "mathbb",
        "mathfrak",
        "text",
        "textrm",
        "textbf",
        "textit",
        "operatorname",
        "rm",
        "bf",
        "it",
        "boldsymbol",
        "vec",
        "hat",
        "bar",
        "tilde",
        "dot",
        "ddot",
    }
)

_LATEX_IGNORE = frozenset(
    {
        "left",
        "right",
        "big",
        "Big",
        "bigg",
        "Bigg",
        "displaystyle",
        "textstyle",
        "scriptstyle",
        "limits",
        "nolimits",
        "operatorname",
        "ensuremath",
    }
)


class _LatexToOmml:
    """Recursive-descent LaTeX → OMML element list for a paragraph."""

    def __init__(self, s: str) -> None:
        self.s = s
        self.i = 0
        self.n = len(s)

    def _peek(self) -> str:
        return self.s[self.i] if self.i < self.n else ""

    def _next(self) -> str:
        c = self.s[self.i]
        self.i += 1
        return c

    def _skip_ws(self) -> None:
        while self.i < self.n and self._peek() in " \t\r\n":
            self.i += 1

    def _run(self, text: str):
        r = _etree.Element(_m("r"))
        rpr = _etree.SubElement(r, f"{{{_WORD_NS}}}rPr")
        fonts = _etree.SubElement(rpr, f"{{{_WORD_NS}}}rFonts")
        fonts.set(f"{{{_WORD_NS}}}ascii", "Cambria Math")
        fonts.set(f"{{{_WORD_NS}}}hAnsi", "Cambria Math")
        t = _etree.SubElement(r, _m("t"))
        t.set(_XML_SPACE, "preserve")
        t.text = text
        return r

    def _read_atom(self):
        """Read one atom ( {...} group, \\command, or single char ) → element list."""
        self._skip_ws()
        c = self._peek()
        if c == "{":
            self._next()
            nodes = self._parse(stop="}")
            if self._peek() == "}":
                self._next()
            return nodes
        if c == "\\":
            return [self._command()]
        if c:
            self._next()
            return [self._run(c)]
        return []

    def _wrap(self, base, sup_nodes=None, sub_nodes=None):
        if sup_nodes is not None and sub_nodes is not None:
            s = _etree.Element(_m("sSubSup"))
            _append_children(_etree.SubElement(s, _m("e")), [base])
            _append_children(_etree.SubElement(s, _m("sub")), sub_nodes)
            _append_children(_etree.SubElement(s, _m("sup")), sup_nodes)
            return s
        if sup_nodes is not None:
            s = _etree.Element(_m("sSup"))
            _append_children(_etree.SubElement(s, _m("e")), [base])
            _append_children(_etree.SubElement(s, _m("sup")), sup_nodes)
            return s
        s = _etree.Element(_m("sSub"))
        _append_children(_etree.SubElement(s, _m("e")), [base])
        _append_children(_etree.SubElement(s, _m("sub")), sub_nodes)
        return s

    def _command(self):
        self._next()  # consume '\'
        name = ""
        while self.i < self.n and self._peek().isalpha():
            name += self._next()
        if not name:  # single non-letter command like \\ \, \{ \%
            c = self._next() if self.i < self.n else ""
            mp = {"\\": "\n", ",": " ", ";": "  ", ":": " ", "%": "%", "{": "{", "}": "}", "|": "|", "!": "", " ": " "}
            return self._run(mp.get(c, c))

        if name in {"frac", "dfrac", "tfrac", "cfrac"}:
            num = self._read_atom()
            den = self._read_atom()
            f = _etree.Element(_m("f"))
            _append_children(_etree.SubElement(f, _m("num")), num)
            _append_children(_etree.SubElement(f, _m("den")), den)
            return f
        if name == "sqrt":
            deg_nodes = None
            self._skip_ws()
            if self._peek() == "[":
                self._next()
                deg_src = ""
                depth = 1
                while self.i < self.n and depth > 0:
                    ch = self._next()
                    if ch == "[":
                        depth += 1
                    elif ch == "]":
                        depth -= 1
                        if depth == 0:
                            break
                    deg_src += ch
                deg_nodes = _LatexToOmml(deg_src)._parse()
            base = self._read_atom()
            rad = _etree.Element(_m("rad"))
            deg_el = _etree.SubElement(rad, _m("deg"))
            if deg_nodes:
                _append_children(deg_el, deg_nodes)
            _append_children(_etree.SubElement(rad, _m("e")), base)
            return rad
        if name in {"binom", "dbinom", "tbinom"}:
            a = self._read_atom()
            b = self._read_atom()
            f = _etree.Element(_m("f"))
            fpr = _etree.SubElement(f, _m("fPr"))
            _etree.SubElement(fpr, _m("type")).set(_m("val"), "lin")
            _append_children(_etree.SubElement(f, _m("num")), a)
            _append_children(_etree.SubElement(f, _m("den")), b)
            return f
        if name in {"overline", "bar", "vec", "hat", "tilde", "dot", "ddot"}:
            # accent/overbar: use <m:groupChr> with the combining accent glyph over the base.
            accents = {"overline": "̄", "bar": "̄", "vec": "⃗", "hat": "̂", "tilde": "̃", "dot": "̇", "ddot": "̈"}
            base = self._read_atom()
            el = _etree.Element(_m("groupChr"))
            pr = _etree.SubElement(el, _m("groupChrPr"))
            _etree.SubElement(pr, _m("chr")).set(_m("val"), accents[name])
            _etree.SubElement(pr, _m("pos")).set(_m("val"), "top")
            _append_children(_etree.SubElement(el, _m("e")), base)
            return el
        if name in {"underline"}:
            base = self._read_atom()
            el = _etree.Element(_m("bar"))
            _etree.SubElement(_etree.SubElement(el, _m("barPr")), _m("pos")).set(_m("val"), "bot")
            _append_children(_etree.SubElement(el, _m("e")), base)
            return el
        if name == "boxed":
            return self._box_around(self._read_atom())
        if name in _LATEX_TEXTY:
            # text-style: render contained atoms upright (plain math style)
            return self._plain_group(self._read_atom())
        if name in _LATEX_IGNORE:
            return self._passthrough_arg()
        if name in {"begin", "end"}:
            self._read_atom()  # consume {env} — ponytail: matrix/cases bodies render inline via subsequent parse
            return self._run("")
        if name in _LATEX_SYMBOLS:
            return self._run(_LATEX_SYMBOLS[name])
        # Unknown command: render its name verbatim (readable, no crash)
        return self._run("\\" + name)

    def _passthrough_arg(self):
        self._skip_ws()
        if self._peek() == "{":
            self._next()
            nodes = self._parse(stop="}")
            if self._peek() == "}":
                self._next()
            return self._group(nodes)  # wrapped so ^_ bind to the whole group
        return self._run("")

    def _plain_group(self, nodes):
        # force upright (plain) style on each math run, then wrap in a group
        for r in [nd for nd in nodes if nd.tag == _m("r")]:
            self._force_plain(r)
        return self._group(nodes)

    @staticmethod
    def _force_plain(r_el):
        # m:rPr (math props) must precede w:rPr inside m:r
        rpr = _etree.Element(_m("rPr"))
        _etree.SubElement(rpr, _m("sty")).set(_m("val"), "p")
        r_el.insert(0, rpr)

    def _group(self, nodes):
        g = _etree.Element(_m("e"))
        _append_children(g, nodes)
        return g

    def _box_around(self, nodes):
        b = _etree.Element(_m("borderBox"))
        _append_children(_etree.SubElement(b, _m("e")), nodes)
        return b

    def _parse(self, stop=None):
        nodes = []
        while self.i < self.n:
            c = self._peek()
            if stop and c == stop:
                break
            if c == "\\":
                nodes.append(self._command())
            elif c == "{":
                self._next()
                nodes.extend(self._parse(stop="}"))
                if self._peek() == "}":
                    self._next()
            elif c == "^":
                self._next()
                sup = self._read_atom()
                if nodes:
                    nodes.append(self._wrap(nodes.pop(), sup_nodes=sup))
            elif c == "_":
                self._next()
                sub = self._read_atom()
                if nodes:
                    nodes.append(self._wrap(nodes.pop(), sub_nodes=sub))
            elif c in " \t\r\n":
                self.i += 1
            elif c == "~":
                self.i += 1
                nodes.append(self._run(" "))
            elif c == "&":
                self.i += 1
                nodes.append(self._run(" "))
            else:
                self.i += 1
                nodes.append(self._run(c))
        return nodes


def _append_children(parent, nodes):
    for x in nodes:
        if x.tag == _m("e"):
            # Stray grouping wrapper from \left/\mathrm/\text passthrough — splice its children
            # directly. (Container <m:e> of sSup/sSub/rad/bar/borderBox are built via SubElement
            # and never pass through here, so this only flattens the passthrough wrappers.)
            for child in list(x):
                _append_children(parent, [child])
        else:
            parent.append(x)


def _build_omath(latex: str, display: bool = False):
    """Build an <m:oMath> (inline) or <m:oMathPara> (display) element from LaTeX."""
    latex = re.sub(r"\\{2,}([a-zA-Z])", r"\\\1", latex.strip())  # editor over-escape: \\frac → \frac
    try:
        nodes = _LatexToOmml(latex)._parse()
    except Exception:
        # ponytail: never let a math parse error crash the export — fall back to literal text.
        nodes = []
        for ch in latex:
            nodes.append(_LatexToOmml("")._run(ch))
    omath = _etree.Element(_m("oMath"))
    _append_children(omath, nodes)
    if not display:
        return omath
    para = _etree.Element(_m("oMathPara"))
    pr = _etree.SubElement(para, _m("oMathParaPr"))
    _etree.SubElement(pr, _m("jc")).set(_m("val"), "center")
    para.append(omath)
    return para


def _append_math(paragraph, latex: str, display: bool = False) -> None:
    """Append a Word-native equation to a python-docx paragraph."""
    paragraph._p.append(_build_omath(latex, display=display))


# ---------------------------------------------------------------------------
# Inline formatting helper — handles **bold**, *italic*, `code`, and $math$
# ---------------------------------------------------------------------------


_MATH_DELIM_RE = re.compile(r"(\$\$[\s\S]+?\$\$|\$[^\$\n]+?\$)")


def _add_inline_text(paragraph, text: str) -> None:
    """Add text with inline **bold**, *italic*, `code`, and ``$math$`` formatting."""
    # First split out math segments ($$...$$ display, $...$ inline) → Word-native equations.
    if "$" in text:
        for part in _MATH_DELIM_RE.split(text):
            if not part:
                continue
            if part.startswith("$$") and part.endswith("$$") and len(part) >= 4:
                _append_math(paragraph, part[2:-2], display=False)
            elif part.startswith("$") and part.endswith("$") and len(part) >= 2:
                _append_math(paragraph, part[1:-1], display=False)
            else:
                _add_inline_text_plain(paragraph, part)
        return
    _add_inline_text_plain(paragraph, text)


def _add_inline_text_plain(paragraph, text: str) -> None:
    """Add text with inline **bold**, *italic*, `code` formatting (no math)."""
    # Tokenize: split on bold/italic/code patterns while keeping delimiters
    parts = re.split(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


# ---------------------------------------------------------------------------
# DOCX generation from parsed blocks + layout template
# ---------------------------------------------------------------------------

# Map common Chinese font names to English equivalents python-docx understands
FONT_ALIASES = {
    "宋体": "SimSun",
    "黑体": "SimHei",
    "仿宋": "FangSong",
    "楷体": "KaiTi",
    "微软雅黑": "Microsoft YaHei",
    "等线": "DengXian",
}


def _resolve_font(name: str) -> str:
    return FONT_ALIASES.get(name, name)


def generate_docx(
    markdown_content: str,
    template_data: dict,
    output_path: Path,
    watermark: str | None = None,
    cover_fields: dict | None = None,
) -> str:
    """Generate a DOCX from markdown using layout template styling.

    Renders up to three sections: cover (no page number) → TOC (roman) → body
    (arabic from 1). Cover/TOC are added only when the template declares
    ``cover_template`` / ``toc_settings``. Returns the output file path.
    """
    frontmatter, body_md = _split_frontmatter(markdown_content)
    blocks = parse_markdown(body_md)
    doc = Document()

    # --- Page setup (applies to the first section; subsequent inherit via add_section) ---
    ps = template_data.get("page_settings", {})
    section = doc.sections[0]
    section.page_width = Cm(21.0) if ps.get("paperSize", "A4") == "A4" else Cm(29.7)
    section.page_height = Cm(29.7) if ps.get("paperSize", "A4") == "A4" else Cm(42.0)
    if ps.get("orientation") == "landscape":
        section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(ps.get("marginTop", 2.54))
    section.bottom_margin = Cm(ps.get("marginBottom", 2.54))
    section.left_margin = Cm(ps.get("marginLeft", 3.17))
    section.right_margin = Cm(ps.get("marginRight", 3.17))

    # --- Body style defaults ---
    bs = template_data.get("body_styles", {})
    body_font = _resolve_font(bs.get("fontFamily", "宋体"))
    body_size = Pt(bs.get("fontSize", 12))
    body_line_spacing = bs.get("lineHeight", 1.5)

    # --- Heading styles map ---
    heading_styles: dict[int, dict] = {}
    for hs in template_data.get("heading_styles", []):
        heading_styles[hs.get("level", 0)] = hs

    # --- Helper: style a paragraph ---
    def style_paragraph(para, font_name: str = body_font, font_size=body_size, bold: bool = False, color: str | None = None, alignment=None, first_indent: float | None = None, space_after: int | None = None):
        pf = para.paragraph_format
        pf.line_spacing = body_line_spacing
        if space_after is not None:
            pf.space_after = Pt(space_after)
        else:
            pf.space_after = Pt(bs.get("paragraphSpacing", 6))
        if first_indent is not None:
            pf.first_line_indent = Cm(first_indent)
        if alignment is not None:
            pf.alignment = alignment

    cover_master = template_data.get("cover_master")
    has_cover = bool(cover_master or template_data.get("cover_template"))
    has_toc = bool(template_data.get("toc_settings"))
    resolved_cover = _resolve_cover_fields(cover_fields or {}, frontmatter, blocks) if has_cover else {}
    numbers = _compute_heading_numbers(blocks, template_data.get("heading_styles", []))

    # === Section 0: COVER ===
    cover_rendered = False
    if has_cover:
        try:
            if cover_master and cover_master.get("mode") == "master":
                _render_cover_master(doc, cover_master, resolved_cover, frontmatter)
            else:
                _render_cover(doc, template_data.get("cover_template"), resolved_cover)
            cover_rendered = True
        except Exception as exc:  # cover must never abort generation (M5: observable, not silent)
            logger.warning("cover render failed: %s", exc)
        # only section-break when a cover was actually rendered — a failed render
        # must not leave a blank leading page.
        if cover_rendered:
            doc.add_section(WD_SECTION.NEW_PAGE)

    # === Section 1: TOC ===
    if has_toc and _render_toc(doc, template_data.get("toc_settings")):
        doc.add_section(WD_SECTION.NEW_PAGE)

    # === Section 2: BODY ===
    ol_counters: dict[int, int] = {}

    for i, block in enumerate(blocks):
        if block.kind == "heading":
            level = min(block.level, 4)
            hs = heading_styles.get(level, {})
            heading = doc.add_heading(level=level)
            h_text = f"{numbers[i]} {block.text}" if i in numbers else block.text
            _add_inline_text(heading, h_text)

            # Apply heading font
            for run in heading.runs:
                _set_run_font(run, _resolve_font(hs.get("fontFamily", body_font)))
                run.font.size = Pt(hs.get("fontSize", 16))
                if hs.get("fontWeight", 700) >= 700:
                    run.bold = True
                c = hs.get("color")
                if c:
                    run.font.color.rgb = RGBColor.from_string(c.replace("#", ""))
            style_paragraph(heading, space_after=6)

        elif block.kind == "paragraph":
            para = doc.add_paragraph()
            indent = bs.get("firstLineIndent", 2)
            style_paragraph(para, first_indent=indent * body_size.pt / 28.35 * 0.5 if indent else None)
            _add_inline_text(para, block.text)
            for run in para.runs:
                _set_run_font(run, body_font)
                run.font.size = body_size

        elif block.kind == "ul_item":
            para = doc.add_paragraph()
            indent_cm = block.level * 0.6
            para.paragraph_format.left_indent = Cm(indent_cm)
            para.paragraph_format.first_line_indent = Cm(-0.3)
            _set_run_font(para.add_run("• "), body_font)
            _add_inline_text(para, block.text)
            for run in para.runs:
                _set_run_font(run, body_font)
                run.font.size = body_size

        elif block.kind == "ol_item":
            indent_level = block.level
            ol_counters[indent_level] = ol_counters.get(indent_level, 0) + 1
            counter = ol_counters[indent_level]
            para = doc.add_paragraph()
            indent_cm = indent_level * 0.6
            para.paragraph_format.left_indent = Cm(indent_cm)
            _set_run_font(para.add_run(f"{counter}. "), body_font)
            _add_inline_text(para, block.text)
            for run in para.runs:
                _set_run_font(run, body_font)
                run.font.size = body_size

        elif block.kind == "hr":
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            run = para.add_run("─" * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)

        elif block.kind == "code_block":
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1)
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(4)
            run = para.add_run(block.text)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        elif block.kind == "table":
            if block.rows:
                tstyles = template_data.get("table_styles")
                ncols = max(len(r) for r in block.rows)
                table = doc.add_table(rows=len(block.rows), cols=ncols)
                table.style = "Table Grid"

                for ri, row in enumerate(block.rows):
                    for ci, cell_text in enumerate(row):
                        if ci < ncols:
                            cell = table.rows[ri].cells[ci]
                            cell.text = ""
                            para = cell.paragraphs[0]
                            _add_inline_text(para, cell_text.strip())
                            for run in para.runs:
                                _set_run_font(run, body_font)
                                run.font.size = Pt(body_size.pt - 1)
                                if ri == 0 and tstyles:
                                    run.bold = True
                                    c = tstyles.get("headerColor", "#FFFFFF")
                                    run.font.color.rgb = RGBColor.from_string(c.replace("#", ""))

                # Header row shading
                if tstyles and block.rows:
                    from docx.oxml.ns import qn

                    bg = tstyles.get("headerBg", "#2B579A").replace("#", "")
                    for ci in range(ncols):
                        cell = table.rows[0].cells[ci]
                        shading = cell._element.get_or_add_tcPr()
                        shading_elem = shading.makeelement(
                            qn("w:shd"),
                            {
                                qn("w:fill"): bg,
                                qn("w:val"): "clear",
                            },
                        )
                        shading.append(shading_elem)

    # === Per-section footers + page numbering + header_footer + watermark ===
    _apply_section_chrome(doc, template_data, watermark, has_cover, has_toc)

    # === Auto-update TOC on open ===
    if has_toc:
        _set_update_fields(doc)

    # Save
    doc.save(str(output_path))
    return str(output_path)


def _apply_section_chrome(doc, template_data: dict, watermark: str | None, has_cover: bool, has_toc: bool) -> None:
    """Apply per-section footer page numbers, pgNumType, header_footer text, watermark.

    Sections layout: [cover?][toc?][body...]. Cover: no page number. TOC: roman.
    Body section(s): decimal restart 1 + header_footer template config + watermark.
    """
    hf = template_data.get("header_footer") or {}
    show_pn = hf.get("showPageNumber", True)
    sections = doc.sections
    last_idx = len(sections) - 1

    for idx, sec in enumerate(sections):
        is_cover = has_cover and idx == 0
        is_toc = has_toc and ((idx == 1) if has_cover else (idx == 0))

        if is_cover:
            continue  # cover: no footer page number, no chrome

        sec.footer.is_linked_to_previous = False
        sec.header.is_linked_to_previous = False

        if is_toc:
            _set_section_pagenum(sec, fmt="upperRoman", start=1)
            if show_pn:
                _add_page_number_footer(sec)
        else:
            # body section(s): decimal page numbers, restart at 1
            _set_section_pagenum(sec, fmt="decimal", start=1)
            _apply_header_footer_text(sec, hf)
            if show_pn:
                _add_page_number_footer(sec)

    # watermark on body (last) section header
    if watermark and last_idx >= 0:
        labels = {"draft": "初稿", "review": "送审稿", "final": "正式稿"}
        label = labels.get(watermark, watermark)
        body_sec = sections[last_idx]
        existing = body_sec.header.paragraphs[0].text if body_sec.header.paragraphs else ""
        body_sec.header.paragraphs[0].text = f"【{label}】{chr(10)}{existing}".strip()
        for run in body_sec.header.paragraphs[0].runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def _apply_header_footer_text(section, hf: dict) -> None:
    """Apply headerText/footerText from template header_footer config to a body section."""
    if hf.get("headerText"):
        section.header.paragraphs[0].text = hf["headerText"]
        for run in section.header.paragraphs[0].runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    if hf.get("footerText"):
        section.footer.paragraphs[0].text = hf["footerText"]


# ---------------------------------------------------------------------------
# Simple DOCX generation (no template) — used by document export
# ---------------------------------------------------------------------------

DEFAULT_BODY_FONT = "SimSun"
DEFAULT_BODY_SIZE = 12
HEADING_SIZES = {1: 22, 2: 16, 3: 14, 4: 12}


def _set_run_font(run, font_name: str) -> None:
    """Set both ASCII and eastAsia font on a run.

    python-docx's ``run.font.name`` only sets ``w:ascii`` / ``w:hAnsi``.
    Chinese text in Word uses ``w:eastAsia`` instead, so we must set both
    to guarantee correct font rendering on the user's machine.
    """
    from docx.oxml.ns import qn

    run.font.name = font_name
    rPr = run._element.find(qn("w:rPr"))
    if rPr is not None:
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is not None:
            rFonts.set(qn("w:eastAsia"), font_name)


def generate_docx_simple(
    markdown_content: str,
    buf,
    template_data: dict | None = None,
    watermark: str | None = None,
    toc_settings: dict | None = None,
    cover_preset: dict | None = None,
    cover_values: dict | None = None,
) -> None:
    """Generate a DOCX from markdown into a writable buffer.

    Args:
        markdown_content: Source markdown text.
        buf: Writable buffer (BytesIO) for the output.
        template_data: Optional layout template dict with keys like
            page_settings, body_styles, heading_styles, table_styles,
            header_footer. When None, sensible defaults are used.
        watermark: Optional watermark type — "draft", "review", or "final".
        toc_settings: Optional dict ``{"maxDepth": int}``. When present with
            maxDepth > 0, a native Word TOC field is rendered before the body
            (Word/WPS auto-updates page numbers on open).
        cover_preset: Optional cover preset dict (see cover_presets.py). When
            present, a cover page renders in its own section with no page
            number; the body section restarts page numbering at 1.
        cover_values: Optional dict of field values for the cover preset
            (e.g. {"title": ..., "client": ...}); a line whose value is
            missing is skipped.
    """
    td = template_data or {}
    # ponytail: lxml rejects C0 control chars (form-feed \x0c, vtab \x0b, bell, …) → 500.
    # Math source like "\frac" can arrive as a literal form-feed after JSON/transport
    # escaping; strip the XML-incompatible range (keep \t \n \r) so export never crashes.
    markdown_content = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", markdown_content)
    # Decode editor math placeholders (<div data-math-block>/<span data-math-inline>) → $/$$
    markdown_content = _decode_math_placeholders(markdown_content)
    blocks = parse_markdown(markdown_content)
    doc = Document()

    # --- Page setup ---
    ps = td.get("page_settings", {})
    section = doc.sections[0]
    section.page_width = Cm(21.0) if ps.get("paperSize", "A4") == "A4" else Cm(29.7)
    section.page_height = Cm(29.7) if ps.get("paperSize", "A4") == "A4" else Cm(42.0)
    if ps.get("orientation") == "landscape":
        section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(ps.get("marginTop", 2.54))
    section.bottom_margin = Cm(ps.get("marginBottom", 2.54))
    section.left_margin = Cm(ps.get("marginLeft", 3.17))
    section.right_margin = Cm(ps.get("marginRight", 3.17))

    # --- Body style ---
    bs = td.get("body_styles", {})
    body_font = _resolve_font(bs.get("fontFamily", DEFAULT_BODY_FONT))
    body_size = Pt(bs.get("fontSize", DEFAULT_BODY_SIZE))
    body_line_spacing = bs.get("lineHeight", 1.5)
    body_paragraph_spacing = bs.get("paragraphSpacing", 6)
    body_first_indent = bs.get("firstLineIndent", 2)

    # --- Heading styles ---
    hs_map: dict[int, dict] = {}
    for hs in td.get("heading_styles", []):
        hs_map[hs.get("level", 0)] = hs

    ol_counters: dict[int, int] = {}

    # --- Optional cover page (own section, no page number) ---
    has_cover = False
    if cover_preset:
        try:
            has_cover = _render_cover_preset(doc, cover_preset, cover_values)
        except Exception as exc:  # cover must never abort generation (M5: observable, not silent)
            logger.warning("cover render failed: %s", exc)
            has_cover = False
        if has_cover:
            doc.add_section(WD_SECTION.NEW_PAGE)

    # --- Optional Table of Contents (built from markdown headings) ---
    has_toc = _render_toc(doc, toc_settings)
    if has_toc:
        doc.add_page_break()

    for block in blocks:
        if block.kind == "heading":
            level = min(block.level, 4)
            hs = hs_map.get(level, {})
            heading = doc.add_heading(level=level)
            _add_inline_text(heading, block.text)
            for run in heading.runs:
                _set_run_font(run, _resolve_font(hs.get("fontFamily", body_font)))
                run.font.size = Pt(hs.get("fontSize", HEADING_SIZES.get(level, 12)))
                if hs.get("fontWeight", 700) >= 700:
                    run.bold = True
                c = hs.get("color")
                if c:
                    run.font.color.rgb = RGBColor.from_string(str(c).replace("#", ""))

        elif block.kind == "paragraph":
            m_disp = re.fullmatch(r"\$\$([\s\S]+)\$\$", block.text.strip())
            if m_disp:
                # Display equation on its own paragraph: centered Word-native equation.
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _append_math(para, m_disp.group(1), display=True)
            else:
                para = doc.add_paragraph()
                pf = para.paragraph_format
                pf.line_spacing = body_line_spacing
                pf.space_after = Pt(body_paragraph_spacing)
                if body_first_indent:
                    pf.first_line_indent = Cm(body_first_indent * body_size.pt / 28.35 * 0.5)
                _add_inline_text(para, block.text)
                for run in para.runs:
                    _set_run_font(run, body_font)
                    run.font.size = body_size

        elif block.kind == "ul_item":
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(block.level * 0.6)
            para.paragraph_format.first_line_indent = Cm(-0.3)
            _set_run_font(para.add_run("• "), body_font)
            _add_inline_text(para, block.text)
            for run in para.runs:
                _set_run_font(run, body_font)
                run.font.size = body_size

        elif block.kind == "ol_item":
            indent_level = block.level
            ol_counters[indent_level] = ol_counters.get(indent_level, 0) + 1
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(indent_level * 0.6)
            _set_run_font(para.add_run(f"{ol_counters[indent_level]}. "), body_font)
            _add_inline_text(para, block.text)
            for run in para.runs:
                _set_run_font(run, body_font)
                run.font.size = body_size

        elif block.kind == "hr":
            # Skip horizontal rules — not needed in Word export
            pass

        elif block.kind == "code_block":
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1)
            run = para.add_run(block.text)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        elif block.kind == "table":
            if block.rows:
                tstyles = td.get("table_styles")
                ncols = max(len(r) for r in block.rows)
                table = doc.add_table(rows=len(block.rows), cols=ncols)
                table.style = "Table Grid"
                for ri, row in enumerate(block.rows):
                    for ci, cell_text in enumerate(row):
                        if ci < ncols:
                            cell = table.rows[ri].cells[ci]
                            cell.text = ""
                            para = cell.paragraphs[0]
                            _add_inline_text(para, cell_text.strip())
                            for run in para.runs:
                                _set_run_font(run, body_font)
                                run.font.size = Pt(body_size.pt - 1)
                                if ri == 0 and tstyles:
                                    run.bold = True
                                    c = tstyles.get("headerColor", "#FFFFFF")
                                    run.font.color.rgb = RGBColor.from_string(str(c).replace("#", ""))

                # Header row shading
                if tstyles and block.rows:
                    from docx.oxml.ns import qn

                    bg = str(tstyles.get("headerBg", "#2B579A")).replace("#", "")
                    for ci in range(ncols):
                        cell = table.rows[0].cells[ci]
                        shading = cell._element.get_or_add_tcPr()
                        shading_elem = shading.makeelement(
                            qn("w:shd"),
                            {
                                qn("w:fill"): bg,
                                qn("w:val"): "clear",
                            },
                        )
                        shading.append(shading_elem)

    # Chrome (header/footer/watermark/page-number) targets the BODY section
    # (sections[-1]). When a cover was added, unlink it from the cover section
    # and restart page numbering at 1; the cover section (sections[0]) gets no
    # chrome → no page number. No cover → sections[-1] is the same single
    # section already bound, so this rebind is a no-op there (byte-identical).
    section = doc.sections[-1]
    if has_cover:
        section.footer.is_linked_to_previous = False
        section.header.is_linked_to_previous = False
        _set_section_pagenum(section, fmt="decimal", start=1)

    # --- Header / Footer ---
    hf = td.get("header_footer")
    if hf:
        if hf.get("headerText"):
            section.header.paragraphs[0].text = hf["headerText"]
            for run in section.header.paragraphs[0].runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        if hf.get("footerText") or hf.get("showPageNumber"):
            footer_para = section.footer.paragraphs[0]
            if hf.get("footerText"):
                footer_para.text = hf["footerText"]
            if hf.get("showPageNumber"):
                from docx.oxml.ns import qn

                run = footer_para.add_run()
                fld_char_begin = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
                run._element.append(fld_char_begin)
                run2 = footer_para.add_run()
                instr = run2._element.makeelement(qn("w:instrText"), {})
                instr.text = " PAGE "
                run2._element.append(instr)
                run3 = footer_para.add_run()
                fld_char_end = run3._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
                run3._element.append(fld_char_end)
            for run in footer_para.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # --- Watermark ---
    if watermark:
        labels = {"draft": "初稿", "review": "送审稿", "final": "正式稿"}
        label = labels.get(watermark, watermark)
        existing = section.header.paragraphs[0].text if section.header.paragraphs else ""
        section.header.paragraphs[0].text = f"【{label}】{chr(10)}{existing}".strip()
        for run in section.header.paragraphs[0].runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # --- Auto-update the TOC field when the document opens in Word/WPS ---
    if has_toc:
        _set_update_fields(doc)

    doc.save(buf)
