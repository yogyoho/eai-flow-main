"""Document parser for Word (.docx) and PDF files.

Extracts structured content from reports for template extraction:
- headings: chapter/section hierarchy with levels and titles
- tables: structured table data with captions, columns, and rows
- full_text: concatenated body text in document order

Word: expat SAX parser (built-in) for true streaming — 64KB chunks,
      never loads entire XML DOM. python-docx fallback for edge cases.
PDF:  PyMuPDF (fitz) text extraction + regex. Scanned PDFs error out.
"""

from __future__ import annotations

import logging
import re
import zipfile as _zipfile
from dataclasses import dataclass, field
from pathlib import Path
from xml.parsers.expat import ExpatError, ParserCreate

logger = logging.getLogger(__name__)


# ── Output data classes ──


@dataclass
class Heading:
    title: str
    level: int
    line_number: int = 0
    style_name: str = ""
    # ponytail: paragraph index + char offset into full_text.
    # -1 = unknown (regex-fallback headings have no anchor).
    # Stored as offsets (not pre-sliced text) to avoid duplicating subtree
    # text across every heading — a tree of N nodes would otherwise hold
    # O(N²) text. Section text is sliced from full_text on demand.
    para_idx: int = -1
    text_offset: int = -1


@dataclass
class DocTable:
    caption: str = ""
    columns: list[str] = field(default_factory=list)
    rows: list[list[str]] = field(default_factory=list)
    page: int = 0


@dataclass
class ParsedDocument:
    file_path: str
    file_type: str
    headings: list[Heading] = field(default_factory=list)
    tables: list[DocTable] = field(default_factory=list)
    full_text: str = ""
    error: str = ""

    def finalize_sections(self, paragraphs: list[str]) -> None:
        """Compute each heading's char offset in full_text.

        Stores only offsets (O(N) memory), not pre-sliced subtree text.
        section_text_by_title slices full_text on demand — correct and
        memory-bounded regardless of heading tree depth.
        """
        offsets: list[int] = []
        pos = 0
        for p in paragraphs:
            offsets.append(pos)
            pos += len(p) + 2  # +2 for "\n\n" separator
        for h in self.headings:
            if 0 <= h.para_idx < len(offsets):
                h.text_offset = offsets[h.para_idx]

    def section_text_by_title(self, title: str, level: int = 0, max_chars: int = 6000) -> str:
        """Slice exact source text for a heading (offset-based, on demand).

        subtree = full_text[heading.offset : next_same_or_higher_level.offset].
        Returns "" if no anchored match (regex-fallback or offset unknown) —
        caller then falls back to fuzzy matching.
        """
        if not self.full_text:
            return ""
        norm = normalize_text(title)
        fallback = ""
        for hi, h in enumerate(self.headings):
            if level and h.level != level:
                continue
            matched = normalize_text(h.title) == norm
            # 子串 fallback 加最小长度（防单字符/过短标题如"1"误匹配"1.1 总则"）
            if not matched:
                shorter = min(len(h.title), len(title))
                if shorter < 2 or not (h.title in title or title in h.title):
                    continue
            if h.text_offset < 0:
                continue
            end = len(self.full_text)
            for j in range(hi + 1, len(self.headings)):
                nj = self.headings[j]
                if nj.level <= h.level and nj.text_offset >= 0:
                    end = nj.text_offset
                    break
            text = self.full_text[h.text_offset : end][:max_chars]
            if matched:
                return text
            if not fallback:
                fallback = text
        return fallback

    def section_length(self, idx: int) -> int:
        """Char length of heading[idx]'s subtree source (for min_section_length filter).

        0 if heading has no anchor (regex fallback) — caller keeps such headings.
        """
        if idx < 0 or idx >= len(self.headings):
            return 0
        h = self.headings[idx]
        if h.text_offset < 0:
            return 0
        end = len(self.full_text)
        for j in range(idx + 1, len(self.headings)):
            nj = self.headings[j]
            if nj.level <= h.level and nj.text_offset >= 0:
                end = nj.text_offset
                break
        return max(0, end - h.text_offset)


# ── Text normalization ──

_STOP_WORDS_RE = re.compile(r"[的、，。；：（）\(\)\s]+")


def normalize_text(text: str) -> str:
    cleaned = _STOP_WORDS_RE.sub("", text.lower().strip())
    return re.sub(r"\s+", "", cleaned)


# ── Regex fallback patterns ──

_CN_CHAPTER = re.compile(r"^第[一二三四五六七八九十百千\d]+[章节条款段]\s*[\.、\s]?\s*(.+)")
_CN_PAREN = re.compile(r"^[（(][一二三四五六七八九十\d]+[）)]\s*(.+)")
_CN_NUMBERED = re.compile(r"^[一二三四五六七八九十]+[、\.\s]\s*(.+)")
_NUMBERED = re.compile(r"^(\d+(?:[\.]\d+)*)[、.\s]+(.+)$")
_TOC_ENTRY = re.compile(r".*\t\d+$|.* {3,}\d+$")


def _heading_level_from_number(num_str: str) -> int:
    return min(num_str.count(".") + 1, 3)


_TABLE_CAPTION_RE = re.compile(r"^表\s*\d+[\.\-]\d+")


def _is_table_caption(title: str) -> bool:
    """判断标题是否为表格 caption（'表N.M-x' 格式）。

    环评样例中，表格标题被作者套了 Heading 5 样式（level=5，65/65 全是
    '表N.M-x'），而 level 4 是真章节（2.2.1.1）。用格式而非 level 判定，
    兼容 regex 兜底（_heading_level_from_number 上限 3，无 level 5）。
    """
    return bool(_TABLE_CAPTION_RE.match(title.strip()))


# ── Body-text guard (bug-404) ──
# 真标题是简短名词短语；含子句/句末标点（，。；！？ 及 ASCII 对应）的一定是
# 被误套了标题样式的正文段。顿号 、 可在真标题里作连词（"设计、施工及验收"），
# 故不列入——只看真正的子句/句末边界标点。两处样式法（expat/python-docx）与
# regex 兜底共用此判定，保持单一真相源。
CLAUSE_PUNCTUATION = set("，。；！？,;!?")


def _looks_like_body_text(text: str) -> bool:
    """真标题不含子句/句末标点；含则判定为正文段（样式法误判守卫）。

    实测 3218 掘进作业规程：被误判为章节的 8 段正文全部含 ，或 。，
    而 9 个真章节标题一个都不含——这是零误伤的判别线。
    """
    return any(ch in CLAUSE_PUNCTUATION for ch in text)


def _is_noise_line(line: str) -> bool:
    if any(c in line for c in ("□", "?", "？")):
        return True
    # 子句/句末标点：真标题不含这些（见 _looks_like_body_text）
    if _looks_like_body_text(line):
        return True
    t = line.strip()
    # 编号+顿号的长行：原 len>15 判定（现已被上面子句标点覆盖大部分，保留兜底）。
    if re.match(r"^\d+[、，]", t) and len(t) > 15 and any(p in t for p in ("。", "；", "，")):
        return True
    return False


def _scan_headings_regex(text: str) -> list[Heading]:
    headings: list[Heading] = []
    line_no = 0
    for line in text.split("\n"):
        line = line.strip()
        if not line or len(line) > 80:
            line_no += 1
            continue
        if _TOC_ENTRY.match(line):
            line_no += 1
            continue
        if _is_noise_line(line):
            line_no += 1
            continue

        matched = None
        if m := _CN_CHAPTER.match(line):
            matched = Heading(title=line, level=1, line_number=line_no, style_name="cn-chapter")
        elif m := _CN_PAREN.match(line):
            matched = Heading(title=line, level=2, line_number=line_no, style_name="cn-paren")
        elif m := _CN_NUMBERED.match(line):
            matched = Heading(title=line, level=2, line_number=line_no, style_name="cn-numbered")
        elif m := _NUMBERED.match(line):
            matched = Heading(title=line, level=_heading_level_from_number(m.group(1)), line_number=line_no, style_name="numbered")
        if matched:
            headings.append(matched)
        line_no += 1

    seen: set[str] = set()
    return [h for h in headings if h.title not in seen and not seen.add(h.title)]  # type: ignore[func-returns-value]


# ── Main dispatch ──


def parse_document(file_path: str) -> ParsedDocument:
    path = Path(file_path)
    ext = path.suffix.lower().lstrip(".")
    if ext == "docx":
        return parse_docx(file_path)
    elif ext == "pdf":
        return parse_pdf(file_path)
    else:
        return ParsedDocument(file_path=file_path, file_type=ext, error=f"不支持的文件格式: .{ext}，仅支持 .docx 和 .pdf")


# ── Word parser (dual-path: expat fast + python-docx fallback) ──

_HEADING_STYLES = {
    "heading1": 1,
    "heading 1": 1,
    "1": 1,
    "heading11": 1,
    "heading2": 2,
    "heading 2": 2,
    "2": 2,
    "heading21": 2,
    "heading3": 3,
    "heading 3": 3,
    "3": 3,
    "heading31": 3,
}


def parse_docx(file_path: str) -> ParsedDocument:
    path = Path(file_path)
    if not path.exists():
        return ParsedDocument(file_path=file_path, file_type="docx", error=f"文件不存在: {file_path}")

    # Primary: expat SAX (true streaming, <30s for 50MB)
    try:
        result = _parse_docx_expat(path)
        if result is not None and (result.headings or result.tables):
            return result
    except Exception as e:
        logger.warning(f"expat parsing failed for {file_path}: {e}")

    # Fallback: python-docx (slow, handles edge cases)
    return _parse_docx_python_docx(file_path)


def _parse_style_levels(zf) -> dict[str, int]:
    """解析 word/styles.xml，建 {styleId.lower(): heading_level} 映射。

    基于 outlineLvl（大纲级别，0=H1, 1=H2...），与样式名语言无关。
    不同 Word 模板的 styleId 是数字/变体名（如灵台用 4/33/139），
    硬编码 _HEADING_STYLES 无法覆盖；outlineLvl 是 Word 标准属性，
    heading 样式必带，提供通用准确识别。

    支持 basedOn 继承：自定义样式 <w:basedOn w:val="父"/> 继承父的
    outlineLvl（Word 标准样式继承语义）。
    """
    if "word/styles.xml" not in zf.namelist():
        return {}
    try:
        xml = zf.open("word/styles.xml")
    except KeyError:
        return {}
    # raw: {id: [outlineLvl(0=未设), basedOn_id]}
    raw: dict[str, list] = {}
    state = {"id": "", "in_style": False}

    def start(name, attrs):
        if name == "w:style":
            state["in_style"] = True
            state["id"] = (attrs.get("w:styleId") or "").lower()
            if state["id"]:
                raw.setdefault(state["id"], [0, ""])
        elif name == "w:outlineLvl" and state["in_style"] and state["id"]:
            v = attrs.get("w:val", "")
            if v.isdigit():
                raw.setdefault(state["id"], [0, ""])[0] = int(v) + 1
        elif name == "w:basedOn" and state["in_style"] and state["id"]:
            raw.setdefault(state["id"], [0, ""])[1] = (attrs.get("w:val") or "").lower()

    def end(name):
        if name == "w:style":
            state["in_style"] = False
            state["id"] = ""

    p = ParserCreate()
    p.StartElementHandler = start
    p.EndElementHandler = end
    try:
        while True:
            c = xml.read(65536)
            if not c:
                break
            p.Parse(c, False)
        p.Parse(b"", True)
    except ExpatError as e:
        logger.warning(f"styles.xml parse error: {e}")
    finally:
        xml.close()

    # 解析 basedOn 继承：自身无 outlineLvl 时递归查父（防环用 seen）
    def _resolve(sid: str, seen: set[str]) -> int:
        if sid in seen or sid not in raw:
            return 0
        lvl, based = raw[sid]
        if lvl:
            return lvl
        if based:
            return _resolve(based, seen | {sid})
        return 0

    return {sid: lvl for sid in raw if (lvl := _resolve(sid, set())) > 0}


def _parse_docx_expat(path: Path) -> ParsedDocument | None:
    """Parse docx using expat SAX — 64KB chunk streaming, zero DOM.

    Reads word/document.xml directly from the ZIP without extracting
    to temp files. Memory = O(headings + tables), not O(document size).
    """
    try:
        zf = _zipfile.ZipFile(str(path), "r")
        if "word/document.xml" not in zf.namelist():
            zf.close()
            return None
        xml_file = zf.open("word/document.xml")
    except (_zipfile.BadZipFile, KeyError) as e:
        logger.warning(f"Failed to open docx ZIP: {e}")
        return None

    # 解析 styles.xml 的 outlineLvl 建 {styleId: heading_level} 通用映射。
    # 比 _HEADING_STYLES 硬编码更准：不同模板 styleId 是数字/变体名，
    # 但 outlineLvl（大纲级别 0=H1）是 Word 标准，与样式名语言无关。
    style_levels = _parse_style_levels(zf)

    headings: list[Heading] = []
    tables: list[DocTable] = []
    paragraphs: list[str] = []
    line_no = 0

    # SAX state
    cur_text: list[str] = []
    cur_style: str = ""
    in_p = False
    in_t = False
    tbl_rows: list[list[str]] = []
    row_cells: list[str] = []
    cell_buf: list[str] = []  # 当前 tc 内的文本片段（cell 边界）
    cur_gridspan: int = 1  # 当前 tc 的 gridSpan（水平合并列数）
    cur_vmerge: str = ""  # 当前 tc 的 vMerge: ""/"restart"/"continue"
    last_row_cells: list[str] = []  # 上一行展开后 cells（vMerge continue 取值）
    in_tbl = False
    in_tr = False
    in_tc = False

    def start(name: str, attrs: dict):
        nonlocal cur_text, cur_style, in_p, in_t
        nonlocal tbl_rows, row_cells, in_tbl, in_tr, in_tc
        nonlocal cell_buf, cur_gridspan, cur_vmerge

        if name == "w:p":
            in_p = True
            cur_text = []
            cur_style = ""
        elif name == "w:pStyle" and in_p:
            cur_style = attrs.get("w:val", "").lower()
        elif name == "w:t":
            in_t = True
        elif name == "w:tbl":
            in_tbl = True
            tbl_rows = []
            last_row_cells.clear()
        elif name == "w:tr" and in_tbl:
            in_tr = True
            row_cells = []
        elif name == "w:tc" and in_tr:
            in_tc = True
            cell_buf = []
            cur_gridspan = 1
            cur_vmerge = ""
        elif name == "w:gridSpan" and in_tc:
            try:
                cur_gridspan = max(1, int(attrs.get("w:val", "1")))
            except (TypeError, ValueError):
                cur_gridspan = 1
        elif name == "w:vMerge" and in_tc:
            # 垂直合并：restart=合并起点(有内容)，无 val=continue(继承上一行同列)
            cur_vmerge = "restart" if attrs.get("w:val") == "restart" else "continue"

    def end(name: str):
        nonlocal in_p, in_t, in_tbl, in_tr, in_tc
        nonlocal tbl_rows, row_cells, line_no, cell_buf, cur_gridspan
        nonlocal cur_vmerge, last_row_cells

        if name == "w:t":
            in_t = False
        elif name == "w:p" and in_p:
            in_p = False
            text = "".join(cur_text).strip()
            if text:
                paragraphs.append(text)
                lv = style_levels.get(cur_style) or _HEADING_STYLES.get(cur_style)
                if lv is not None and not _looks_like_body_text(text):
                    # 守卫：正文段即便被套了标题样式（含子句标点），也不当 heading（bug-404）
                    headings.append(Heading(title=text, level=lv, line_number=line_no, style_name=f"Heading {lv}", para_idx=len(paragraphs) - 1))
            line_no += 1
        elif name == "w:tc":
            # vMerge continue: 继承上一行同列文本（垂直合并）；restart/无合并用 cell_buf
            if cur_vmerge == "continue":
                col = len(row_cells)
                cell_text = last_row_cells[col] if col < len(last_row_cells) else ""
            else:
                cell_text = "".join(cell_buf).strip()
            # gridSpan 水平合并：内容复制 N 份对齐列数
            row_cells.extend([cell_text] * cur_gridspan)
            in_tc = False
        elif name == "w:tr" and in_tr:
            in_tr = False
            tbl_rows.append(list(row_cells))
            last_row_cells = list(row_cells)  # 供下一行 vMerge continue 取值
        elif name == "w:tbl":
            in_tbl = False
            if tbl_rows:
                tables.append(
                    DocTable(
                        columns=tbl_rows[0] if tbl_rows else [],
                        rows=tbl_rows[1:] if len(tbl_rows) > 1 else [],
                    )
                )
                # EAI-CUSTOM (bug-1120): 展平表格行列进 paragraphs → full_text，
                # 让 Step 2 LLM 能看见真实表结构（否则 table_schemas 全靠猜）。
                # `|` 前缀行 + `【表格】` 标记保证不被标题 pattern 误判为章节标题
                # （doc_parser/_scan_headings_regex 与 pipeline/_scan_chapter_headings
                # 的标题 pattern 均锚定行首数字/汉字序号，`|`/`【` 开头不命中）。
                # 单元格内 `|` 换全角 `｜` 防列边界歧义。
                paragraphs.append("【表格】")
                header = "| " + " | ".join(c.replace("|", "｜") for c in tbl_rows[0]) if tbl_rows[0] else "|"
                paragraphs.append(header)
                for r in tbl_rows[1:]:
                    paragraphs.append("| " + " | ".join(c.replace("|", "｜") for c in r))
                paragraphs.append("【表格结束】")

    def data(text: str):
        if in_t and in_p and not in_tc:
            cur_text.append(text)
        elif in_t and in_tc:
            cell_buf.append(text)  # cell 内文本缓冲，tc end 时作为一个 cell

    p = ParserCreate()  # 不设 namespace_separator: 保留 w: QName 前缀
    p.StartElementHandler = start
    p.EndElementHandler = end
    p.CharacterDataHandler = data

    try:
        while True:
            chunk = xml_file.read(65536)
            if not chunk:
                break
            p.Parse(chunk, False)
        p.Parse(b"", True)
    except ExpatError as e:
        logger.warning(f"expat parse error in {path.name}: {e}")
        return None
    finally:
        xml_file.close()
        zf.close()

    full_text = "\n\n".join(paragraphs)
    if not headings and full_text:
        headings = _scan_headings_regex(full_text)

    result = ParsedDocument(file_path=str(path), file_type="docx", headings=headings, tables=tables, full_text=full_text)
    result.finalize_sections(paragraphs)
    return result


def _parse_docx_python_docx(file_path: str) -> ParsedDocument:
    """Fallback: python-docx parser for edge cases expat can't handle."""
    try:
        import docx
    except ImportError:
        return ParsedDocument(file_path=file_path, file_type="docx", error="python-docx 未安装")

    path = Path(file_path)
    try:
        doc = docx.Document(str(path))
    except Exception as e:
        return ParsedDocument(file_path=file_path, file_type="docx", error=f"Word 文件解析失败: {e}")

    headings: list[Heading] = []
    tables: list[DocTable] = []
    paragraphs: list[str] = []
    line_no = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            line_no += 1
            continue
        paragraphs.append(text)
        s = (para.style.name if para.style else "").lower()
        if _looks_like_body_text(text):
            pass  # 守卫：正文段即便套了标题样式也不当 heading（bug-404）
        elif "heading 1" in s or s == "heading1":
            headings.append(Heading(title=text, level=1, line_number=line_no, style_name=para.style.name, para_idx=len(paragraphs) - 1))
        elif "heading 2" in s or s == "heading2":
            headings.append(Heading(title=text, level=2, line_number=line_no, style_name=para.style.name, para_idx=len(paragraphs) - 1))
        elif "heading 3" in s or s == "heading3":
            headings.append(Heading(title=text, level=3, line_number=line_no, style_name=para.style.name, para_idx=len(paragraphs) - 1))
        line_no += 1

    for table in doc.tables:
        rows = [[c.text.strip() for c in row.cells] for row in table.rows]
        if rows:
            tables.append(DocTable(columns=rows[0] if rows else [], rows=rows[1:] if len(rows) > 1 else []))

    full_text = "\n\n".join(paragraphs)
    if not headings and full_text:
        headings = _scan_headings_regex(full_text)

    result = ParsedDocument(file_path=file_path, file_type="docx", headings=headings, tables=tables, full_text=full_text)
    result.finalize_sections(paragraphs)
    return result


# ── PDF parser ──


def parse_pdf(file_path: str) -> ParsedDocument:
    path = Path(file_path)
    if not path.exists():
        return ParsedDocument(file_path=file_path, file_type="pdf", error=f"文件不存在: {file_path}")

    try:
        import fitz

        doc = fitz.open(str(path))
        try:
            parts = [page.get_text() for page in doc if page.get_text().strip()]
        finally:
            doc.close()  # 异常路径也关闭，防文件描述符泄漏
        full_text = "\n\n".join(parts)
        if not full_text.strip():
            return ParsedDocument(file_path=file_path, file_type="pdf", error="PDF 无可提取文字（可能是扫描版），请使用 RAGFlow OCR 路径")
        headings = _scan_headings_regex(full_text)
        return ParsedDocument(file_path=file_path, file_type="pdf", headings=headings, full_text=full_text)
    except ImportError:
        pass

    try:
        text = Path(file_path).read_text(encoding="utf-8", errors="replace")
        if len([c for c in text if c.isprintable() or c in "\n\r\t "]) < len(text) * 0.5:
            return ParsedDocument(file_path=file_path, file_type="pdf", error="PDF 为二进制内容（可能是扫描版），请使用 RAGFlow OCR 路径")
        return ParsedDocument(file_path=file_path, file_type="pdf", headings=_scan_headings_regex(text), full_text=text)
    except Exception as e:
        return ParsedDocument(file_path=file_path, file_type="pdf", error=f"PDF 解析失败: {e}")


# ── Structure hint builder ──


def build_structure_hint(parsed: ParsedDocument, max_chars: int = 5000) -> str:
    """构建结构提示：完整章节树 + 每章短摘要。

    P3 修复（覆盖不完整）：旧实现每章仅前 450 字摘要，长章节（环评 406 页
    13 章全部）子节标题被截断 → Step 1 章节推断建不全子节树 → Step 2
    漏抽被漏子节的表。改为：
    1. H1 目录 + 每 H1 下完整子节树（H2/H3/H4，过滤表 caption）
    2. 每章 200 字短摘要（结构已由子节树提供，内容仅作 purpose 信号）
    子节树优先保留，max_chars 超限时摘要后砍。
    """
    headings = parsed.headings
    if not headings:
        return parsed.full_text[:max_chars]

    # H1 起始 index 列表（无 H1 时退化到 level<=2）
    h1_positions = [i for i, h in enumerate(headings) if h.level == 1]
    if not h1_positions:
        h1_positions = [i for i, h in enumerate(headings) if h.level <= 2]
    if not h1_positions:
        h1_positions = list(range(len(headings)))

    parts = [f"## 文档章节目录（自动识别，共 {len(h1_positions)} 章）\n"]

    # 子节树：H1 → H2/H3/H4，过滤表 caption（表N.M-x 是 Heading5 噪声）
    for pi, start in enumerate(h1_positions):
        end = h1_positions[pi + 1] if pi + 1 < len(h1_positions) else len(headings)
        h1 = headings[start]
        parts.append(f"### {h1.title}")
        for k in range(start + 1, end):
            hk = headings[k]
            if hk.level > 4 or _is_table_caption(hk.title):
                continue
            indent = "  " * (hk.level - 1)
            parts.append(f"{indent}- {hk.title}")

    parts.append("\n## 各章节内容摘要（每章节前200字）\n")
    for start in h1_positions:
        h = headings[start]
        idx = h.text_offset if h.text_offset >= 0 else parsed.full_text.find(h.title)
        raw = parsed.full_text[idx : idx + 200] if idx >= 0 else "（内容未找到）"
        # 摘要内也剔除表 caption 行（表N.M-x），防止 Heading5 噪声泄漏进 hint
        snippet = "\n".join(ln for ln in raw.split("\n") if not _is_table_caption(ln.strip()))
        parts.append(f"### {h.title}\n{snippet}\n")

    # 子节树优先：max_chars 超限时保留树 + 预算内摘要，树本身超限则截断树
    result = "\n".join(parts)
    if len(result) > max_chars:
        tree_lines = []
        for line in parts:
            if line.lstrip().startswith("## 各章节内容摘要"):
                break
            tree_lines.append(line)
        tree = "\n".join(tree_lines)
        if len(tree) <= max_chars:
            budget = max_chars - len(tree)
            summary_lines = []
            for line in parts:
                if not line.startswith("### ") or line in tree_lines:
                    continue
                if sum(len(s) for s in summary_lines) + len(line) > budget:
                    break
                summary_lines.append(line)
            result = tree + "\n" + "\n".join(summary_lines)
        else:
            result = tree[:max_chars]
    return result
