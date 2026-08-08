"""Tests for the structured cover element model (replaces cover_master passthrough)."""

import tempfile
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn
from pydantic import ValidationError

from app.extensions.output.generator import _render_cover_elements, generate_docx
from app.extensions.output.layout_import import _extract_cover_pages
from app.extensions.output.schemas import (
    CoverElementSchema,
    CoverSchema,
    LayoutTemplateCreate,
)

SAMPLE1 = Path("C:/Temp/eai-cover1-消防.docx")  # 基地项目-消防设计专篇
SAMPLE2 = Path("C:/Temp/eai-cover2-环评.docx")  # 横城矿区环评报告

# 金标准样例只在本地开发机存在;CI(Linux) 无 C:/Temp → 跳过而非 FileNotFoundError。
skip_missing = pytest.mark.skipif(not (SAMPLE1.exists() and SAMPLE2.exists()), reason="金标准样例缺失（本地开发机才有）")


def _cover_of(path):
    return _extract_cover_pages(Document(str(path)))


def test_cover_schema_roundtrip():
    cover = CoverSchema(
        sourceFile="x.docx",
        pages=[
            {
                "elements": [
                    {"id": "e1", "type": "text", "text": "项目名", "alignment": "center", "fontSize": 22, "bold": True},
                    {"id": "e2", "type": "table", "rows": 2, "cols": 2, "cells": [["专业名称", "编制"], ["总图", ""]]},
                ]
            }
        ],
    )
    d = cover.model_dump()
    assert d["mode"] == "elements"
    assert d["pages"][0]["elements"][0]["text"] == "项目名"
    assert d["pages"][0]["elements"][1]["cells"][0][0] == "专业名称"


def test_cover_element_type_validated():
    with pytest.raises(ValidationError):
        CoverElementSchema(id="x", type="canvas", text="t")


def test_layout_template_create_accepts_cover_elements():
    tpl = LayoutTemplateCreate(
        name="T",
        report_type="g",
        page_settings={"paperSize": "A4", "orientation": "portrait", "marginTop": 2.54, "marginBottom": 2.54, "marginLeft": 3.17, "marginRight": 3.17},
        body_styles={"fontFamily": "宋体", "fontSize": 12, "lineHeight": 1.5, "paragraphSpacing": 0, "firstLineIndent": 2},
        heading_styles=[],
        cover_elements={"sourceFile": "x.docx", "pages": [{"elements": [{"id": "e1", "type": "text", "text": "报告标题"}]}]},
    )
    assert tpl.cover_elements.pages[0].elements[0].text == "报告标题"


@skip_missing
def test_fire_sample_two_pages_banner_and_qianbiao_table():
    """消防设计专篇: 2 页 — 第1页 banner(标题横幅/项目编号/日期), 第2页 会签表.

    分节符/分页符段落为空时也切页: 标题横幅表后的空分节符段落把会签表顶到第2页;
    目录前的空分页符段落生成的幽灵页(全 spacer)被过滤。
    """
    pages = _cover_of(SAMPLE1)
    assert len(pages) == 2, f"消防样例应为 2 页, got {len(pages)}"
    # 第1页 = banner 文本元素(项目名/册标题/编号/日期绑定)
    p0_texts = [e.text for e in pages[0].elements if e.type == "text"]
    assert any("第三册 消防设计专篇" in t for t in p0_texts), "报告名称文本元素缺失(banner页)"
    assert any(t.strip() == "项目名" for t in p0_texts), "独立项目名占位元素缺失(banner页)"
    p0_bound = {e.slotId: e for e in pages[0].elements if e.slotId}
    assert p0_bound.get("project_number"), "项目编号:XX 应绑 project_number(banner页)"
    assert p0_bound.get("date"), "20XX年0X月 应绑 date(banner页)"
    # 第2页 = 会签表表格元素
    p1_tables = [e for e in pages[1].elements if e.type == "table"]
    assert p1_tables, "会签表应为表格元素(第2页)"
    assert p1_tables[0].rows >= 10, f"会签表 rows 应 >=10, got {p1_tables[0].rows}"
    assert p1_tables[0].cols >= 5, f"会签表 cols 应 >=5(实为6), got {p1_tables[0].cols}"
    ids = [e.id for p in pages for e in p.elements]
    assert len(ids) == len(set(ids)), f"元素 id 应唯一, got {len(ids)} els / {len(set(ids))} unique"


@skip_missing
def test_huanping_sample_three_pages():
    """环评报告: 3 页 (封面/批准页/名单页), 含名单表格."""
    pages = _cover_of(SAMPLE2)
    assert len(pages) == 3, f"环评样例应为 3 页, got {len(pages)}"
    p1_texts = [e.text for e in pages[0].elements if e.type == "text"]
    assert any("环境影响报告书" in t for t in p1_texts)
    p2_texts = " ".join(e.text for e in pages[1].elements)
    assert "工程" in p2_texts and "H7367Z" in p2_texts
    imgs = [e for p in pages for e in p.elements if e.type == "image"]
    assert len(imgs) >= 1, f"环评封面应有 ≥1 个 logo 图片元素, got {len(imgs)}"
    p3_tables = [e for e in pages[2].elements if e.type == "table"]
    assert len(p3_tables) == 2, f"名单页应 2 张表, got {len(p3_tables)}"
    assert p3_tables[1].rows >= 16
    assert p3_tables[0].cols >= 3, f"名单页首表 cols 应 >=3(实为3), got {p3_tables[0].cols}"
    ids = [e.id for p in pages for e in p.elements]
    assert len(ids) == len(set(ids)), f"元素 id 应唯一, got {len(ids)} els / {len(set(ids))} unique"


@skip_missing
def test_import_layout_includes_cover_elements():
    from app.extensions.output.layout_import import extract_layout_from_docx

    data = SAMPLE1.read_bytes()
    r = extract_layout_from_docx(data, source_file="消防.docx")
    ce = r.get("cover_elements")
    assert ce and ce["mode"] == "elements", "import 响应应含 cover_elements"
    assert ce["pages"], "应有页面"
    assert ce["sourceFile"] == "消防.docx", "sourceFile 应透传"
    texts = [e["text"] for p in ce["pages"] for e in p["elements"] if e["type"] == "text"]
    assert any("第三册 消防设计专篇" in t for t in texts), "封面文本元素应提取"
    assert r["cover_detected"] is True


def _sample_cover():
    return {
        "mode": "elements",
        "sourceFile": "x.docx",
        "pages": [
            {"elements": [
                {"id": "e1", "type": "text", "text": "项目名", "fontSize": 22, "bold": True, "alignment": "center", "slotId": "project_name"},
                {"id": "e2", "type": "text", "text": "项目编号：XX", "fontSize": 14, "alignment": "center", "slotId": "project_number"},
                {"id": "e3", "type": "text", "text": "环境影响报告书", "fontSize": 22, "alignment": "center"},
                {"id": "e4", "type": "table", "rows": 2, "cols": 2, "cells": [["专业名称", "编制"], ["总图", ""]], "headerBg": "#D9D9D9"},
                {"id": "e5", "type": "spacer", "lines": 2},
            ]},
            {"elements": [
                {"id": "e6", "type": "text", "text": "审定、审查人员名单", "fontSize": 16, "alignment": "center"},
            ]},
        ],
    }


def test_render_cover_element_divider_is_page_break():
    """分隔线元素重定义为分页符：渲染为 <w:br w:type="page"/>，非空段兜底。"""
    from docx.oxml.ns import qn

    cover = {
        "mode": "elements",
        "sourceFile": "x.docx",
        "pages": [{"elements": [
            {"id": "e1", "type": "text", "text": "封面一", "fontSize": 14, "alignment": "center"},
            {"id": "e2", "type": "divider"},
            {"id": "e3", "type": "text", "text": "封面二", "fontSize": 14, "alignment": "center"},
        ]}],
    }
    doc = Document()
    _render_cover_elements(doc, cover, {}, {})
    brs = doc.element.body.findall(f".//{qn('w:br')}")
    assert any(br.get(qn("w:type")) == "page" for br in brs), "divider 应渲染为分页符"
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert "封面一" in texts and "封面二" in texts, "分页符前后文本应都保留"


def test_render_cover_element_page_break():
    """pageBreak 元素：在文本间插一页断 → 首文本段后的下一段含 <w:br w:type="page"/>."""
    cover = {
        "mode": "elements",
        "sourceFile": "x.docx",
        "pages": [{"elements": [
            {"id": "e1", "type": "text", "text": "封面一", "fontSize": 14, "alignment": "center"},
            {"id": "e2", "type": "pageBreak"},
            {"id": "e3", "type": "text", "text": "封面二", "fontSize": 14, "alignment": "center"},
        ]}],
    }
    doc = Document()
    _render_cover_elements(doc, cover, {}, {})
    paras = doc.paragraphs
    # 找到第一段文本段落，断言其后一段含分页符 run
    first_text_idx = next(i for i, p in enumerate(paras) if p.text.strip())
    nxt = paras[first_text_idx + 1]
    brs = nxt._element.findall(f".//{qn('w:br')}")
    assert any(br.get(qn("w:type")) == "page" for br in brs), "pageBreak 应渲染为 <w:br w:type='page'/>"
    texts = [p.text for p in paras if p.text.strip()]
    assert "封面一" in texts and "封面二" in texts, "分页符前后文本应都保留"


def test_render_cover_elements_slot_replacement_and_pages():
    resolved = {"project_name": "基地项目", "project_number": "P001"}
    doc = Document()
    _render_cover_elements(doc, _sample_cover(), resolved, {})
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert "基地项目" in texts, "项目名绑定应替换为 基地项目"
    assert any("项目编号：P001" in t for t in texts), "冒号字段应保留标签替换值"
    assert any("环境影响报告书" in t for t in texts), "未绑定元素保留原文"
    assert len(doc.tables) == 1, "表格元素应生成 1 张 docx 表"
    assert doc.tables[0].rows[0].cells[0].text.strip() == "专业名称"
    tc_pr = doc.tables[0].rows[0].cells[0]._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    assert shd is not None, "表头单元格应写 headerBg 底纹 w:shd"
    assert shd.get(qn("w:fill")) == "D9D9D9", "headerBg #D9D9D9 应落到 w:fill"
    assert len(doc.sections) >= 2, "多页元素应产生分节"


def test_cover_master_to_elements_converts_old_master():
    from app.extensions.output.layout_import import _cover_master_to_elements
    xml = ('<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
           '<w:r><w:t>项目名</w:t></w:r></w:p>')
    master = {"mode": "master", "xml": xml, "images": [], "slots": [], "sourceFile": "old.docx", "boundary": "before_toc"}
    cover = _cover_master_to_elements(master)
    assert cover["mode"] == "elements"
    assert cover["pages"], "应至少 1 页"
    texts = [e["text"] for e in cover["pages"][0]["elements"] if e["type"] == "text"]
    assert "项目名" in texts


def test_cover_master_to_elements_handles_bad_xml():
    from app.extensions.output.layout_import import _cover_master_to_elements
    master = {"mode": "master", "xml": "<w:p>", "images": [], "slots": [], "sourceFile": "bad.docx", "boundary": "before_toc"}
    assert _cover_master_to_elements(master) is None  # 坏 xml → None, 保留旧母版


def test_cover_master_read_migration_does_not_persist_to_db():
    """spec-review 回归: 读时迁移只改内存, get_db 请求尾 commit 不得落库.

    ``_migrate_cover_master`` 设置 ``cover_elements`` 后必须 expunge 出 session,
    否则 get_db 的无条件 commit 会把读时迁移写成新模板, 永久改写旧母版 (图片→spacer)。
    """
    import asyncio

    from sqlalchemy import JSON, Column, Integer
    from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine
    from sqlalchemy.orm import declarative_base

    from app.extensions.output.service import _migrate_cover_master

    TestBase = declarative_base()

    class LegacyTpl(TestBase):
        __tablename__ = "legacy_tpl"
        id = Column(Integer, primary_key=True)
        cover_master = Column(JSON, nullable=True)
        cover_elements = Column(JSON, nullable=True)

    xml = ('<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
           '<w:r><w:t>项目名</w:t></w:r></w:p>')

    async def scenario():
        engine = create_async_engine("sqlite+aiosqlite:///:memory:")
        async with engine.begin() as conn:
            await conn.run_sync(TestBase.metadata.create_all)
        Factory = async_sessionmaker(engine, class_=AsyncSession, expire_on_commit=False)

        async with Factory() as db:
            tpl = LegacyTpl(
                cover_master={"mode": "master", "xml": xml, "images": [], "slots": [], "sourceFile": "old.docx", "boundary": "before_toc"},
                cover_elements=None,
            )
            db.add(tpl)
            await db.commit()
            await db.refresh(tpl)
            # 读路径: 迁移只改内存
            await _migrate_cover_master(tpl, db)
            assert tpl.cover_elements, "内存已迁移"
            # 模拟 get_db 请求尾的无条件 commit
            await db.commit()

        # 新 session 复查: cover_elements 必须仍未落库
        async with Factory() as db2:
            fresh = await db2.get(LegacyTpl, tpl.id)
            assert fresh.cover_elements is None, "GET 读时迁移不得写库"

    asyncio.run(scenario())


def test_cover_master_to_elements_restores_logo_from_images():
    """Fix #1: 迁移从 master["images"] (按 origRid) 恢复 logo, 而非降级为 spacer."""
    from app.extensions.output.layout_import import _cover_master_to_elements

    xml = ('<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
           'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
           'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
           '<w:r><w:drawing><a:blip r:embed="rId1"/></w:drawing></w:r></w:p>')
    master = {
        "mode": "master",
        "xml": xml,
        "images": [{"origRid": "rId1", "ext": "png", "b64": "QUJD"}],
        "slots": [],
        "sourceFile": "old.docx",
        "boundary": "before_toc",
    }
    cover = _cover_master_to_elements(master)
    imgs = [e for e in cover["pages"][0]["elements"] if e["type"] == "image"]
    assert imgs, "logo 应恢复为 image 元素"
    assert imgs[0]["image"]["b64"] == "QUJD"
    assert imgs[0]["image"]["ext"] == "png"


def test_cover_master_to_elements_preserves_multipage():
    """Fix #3: 迁移与 _extract_cover_pages 同源 → 多页母版按分页符保留页数."""
    from app.extensions.output.layout_import import _cover_master_to_elements

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    xml = (
        f'<w:p xmlns:w="{W}"><w:r><w:t>封面</w:t></w:r></w:p>'
        f'<w:p xmlns:w="{W}"><w:r><w:t>批准页</w:t></w:r><w:r><w:br w:type="page"/></w:r></w:p>'
        f'<w:p xmlns:w="{W}"><w:r><w:t>名单页</w:t></w:r></w:p>'
    )
    master = {"mode": "master", "xml": xml, "images": [], "slots": [], "sourceFile": "old.docx", "boundary": "before_toc"}
    cover = _cover_master_to_elements(master)
    assert len(cover["pages"]) == 2, f"多页母版应切为 2 页, got {len(cover['pages'])}"
    p1_texts = [e["text"] for e in cover["pages"][0]["elements"] if e["type"] == "text"]
    p2_texts = [e["text"] for e in cover["pages"][1]["elements"] if e["type"] == "text"]
    assert "封面" in p1_texts
    assert "名单页" in p2_texts


def test_cover_master_migration_update_delete_after_expunge():
    """spec-review #5: expunge 后 update/delete 走真实 re-attach 路径仍能落库."""
    import asyncio

    from sqlalchemy import JSON, Column, Integer, String
    from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine
    from sqlalchemy.orm import declarative_base

    from app.extensions.output.schemas import LayoutTemplateUpdate
    from app.extensions.output.service import LayoutTemplateService, _migrate_cover_master

    TestBase = declarative_base()

    class LegacyTpl(TestBase):
        __tablename__ = "legacy_tpl"
        id = Column(Integer, primary_key=True)
        name = Column(String, nullable=False, default="")
        cover_master = Column(JSON, nullable=True)
        cover_elements = Column(JSON, nullable=True)

    xml = ('<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
           '<w:r><w:t>项目名</w:t></w:r></w:p>')

    async def scenario():
        engine = create_async_engine("sqlite+aiosqlite:///:memory:")
        async with engine.begin() as conn:
            await conn.run_sync(TestBase.metadata.create_all)
        Factory = async_sessionmaker(engine, class_=AsyncSession, expire_on_commit=False)

        async with Factory() as db:
            tpl = LegacyTpl(
                name="旧名",
                cover_master={"mode": "master", "xml": xml, "images": [], "slots": [], "sourceFile": "old.docx", "boundary": "before_toc"},
                cover_elements=None,
            )
            db.add(tpl)
            await db.commit()
            await db.refresh(tpl)
            tpl_id = tpl.id
            # 读路径迁移 → expunge
            await _migrate_cover_master(tpl, db)
            assert tpl.cover_elements, "内存已迁移"
            assert tpl not in db, "expunge 后模板应已脱离 session"
            # UPDATE: expunged → re-attach → set name → commit/refresh
            updated = await LayoutTemplateService.update_template(db, tpl, LayoutTemplateUpdate(name="新名"))
            assert updated.name == "新名"

        async with Factory() as db2:
            fresh = await db2.get(LegacyTpl, tpl_id)
            assert fresh.name == "新名", "update 应落库"
            # DELETE: expunged → re-attach → delete
            await _migrate_cover_master(fresh, db2)
            await LayoutTemplateService.delete_template(db2, fresh)

        async with Factory() as db3:
            assert await db3.get(LegacyTpl, tpl_id) is None, "delete 应落库"

    asyncio.run(scenario())


def test_generate_docx_uses_cover_elements_priority():
    tpl = {
        "page_settings": {"paperSize": "A4", "orientation": "portrait", "marginTop": 2.54, "marginBottom": 2.54, "marginLeft": 3.17, "marginRight": 3.17},
        "body_styles": {"fontFamily": "宋体", "fontSize": 12, "lineHeight": 1.5, "paragraphSpacing": 0, "firstLineIndent": 2},
        "heading_styles": [],
        "cover_elements": _sample_cover(),
        "cover_master": {"mode": "master", "xml": "<w:p/>", "images": [], "slots": [], "sourceFile": "old", "boundary": "before_toc"},
        "cover_template": None,
        "toc_settings": None,
    }
    with tempfile.TemporaryDirectory() as d:
        out = Path(d) / "r.docx"
        generate_docx("# 正文\n", tpl, out, cover_fields={"project_name": "基地项目", "project_number": "P001"})
        doc = Document(str(out))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "基地项目" in texts, "cover_elements 应优先于 cover_master 渲染"
        # 负向断言：cover_master(<w:p/> 空段) 未被渲染 → 首非空段应为封面标题，而非正文/空注入
        assert texts and texts[0] == "基地项目", "cover_elements 应最先渲染（首非空段为封面标题），非 cover_master 空段注入"
        assert "正文" in texts, "封面后正文应正常渲染"
