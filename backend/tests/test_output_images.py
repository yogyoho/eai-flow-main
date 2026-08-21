"""Word export image embedding: markdown ``![](url)`` → inline picture (EAI-CUSTOM)."""

import struct
import zlib
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.extensions.output.generator import generate_docx_simple, parse_markdown


def _png(w: int, h: int) -> bytes:
    """Minimal real PNG (8-bit grayscale), stdlib only — python-docx reads IHDR for dims."""

    def chunk(tag: bytes, data: bytes) -> bytes:
        return struct.pack(">I", len(data)) + tag + data + struct.pack(">I", zlib.crc32(tag + data))

    ihdr = struct.pack(">IIBBBBB", w, h, 8, 0, 0, 0, 0)
    raw = b"".join(b"\x00" + b"\x80" * w for _ in range(h))
    return b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", zlib.compress(raw)) + chunk(b"IEND", b"")


def test_parse_markdown_standalone_image_is_image_block():
    blocks = parse_markdown("前言\n\n![logo](/api/extensions/docmgr/images/aaaaaaaaaaaa.png)\n\n后记")
    assert [b.kind for b in blocks] == ["paragraph", "image", "paragraph"]
    assert blocks[1].text == "/api/extensions/docmgr/images/aaaaaaaaaaaa.png"


def test_export_embeds_image_bytes():
    buf = BytesIO()
    generate_docx_simple("![x](/api/extensions/docmgr/images/aaaaaaaaaaaa.png)", buf, image_fetcher=lambda url: _png(10, 5))
    doc = Document(buf)
    assert len(doc.inline_shapes) == 1
    assert doc.paragraphs[-1].alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_export_image_fallback_to_text_when_fetcher_misses():
    buf = BytesIO()
    generate_docx_simple("![x](/api/extensions/docmgr/images/bbbbbbbbbbbb.png)", buf, image_fetcher=lambda url: None)
    doc = Document(buf)
    assert len(doc.inline_shapes) == 0
    assert "/api/extensions/docmgr/images/bbbbbbbbbbbb.png" in doc.paragraphs[-1].text


def test_export_image_width_capped_to_content_width():
    buf = BytesIO()
    generate_docx_simple("![](u)", buf, image_fetcher=lambda url: _png(5000, 100))
    doc = Document(buf)
    sec = doc.sections[-1]
    avail = sec.page_width - sec.left_margin - sec.right_margin
    shape = doc.inline_shapes[0]
    assert 0 < shape.width <= avail


def test_export_image_fetcher_exception_does_not_crash():
    def boom(url: str) -> bytes:
        raise RuntimeError("disk gone")

    buf = BytesIO()
    generate_docx_simple("![](u)", buf, image_fetcher=boom)
    doc = Document(buf)
    assert len(doc.inline_shapes) == 0


def _jpg_with_orientation(w: int, h: int, orientation: int) -> bytes:
    from PIL import Image

    img = Image.new("L", (w, h), 0x80)
    exif = Image.Exif()
    exif[274] = orientation
    out = BytesIO()
    img.save(out, format="JPEG", exif=exif)
    return out.getvalue()


def test_export_applies_exif_orientation():
    """手机照片 EXIF 方向：浏览器自动旋转、Word 不认 —— 导出必须把旋转烤进像素。"""
    import pytest

    pytest.importorskip("PIL")
    # 原始像素横图 100×50，orientation=8 → 正确显示应为竖图
    buf = BytesIO()
    generate_docx_simple("![](u)", buf, image_fetcher=lambda url: _jpg_with_orientation(100, 50, 8))
    doc = Document(buf)
    shape = doc.inline_shapes[0]
    assert shape.height > shape.width
