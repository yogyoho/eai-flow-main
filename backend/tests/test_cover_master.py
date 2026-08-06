"""Tests for cover-master OOXML passthrough + slot binding (B1)."""

import base64
import struct
import zlib
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace

import pytest  # noqa: F401  # used in later tasks (B2+ fixtures/markers)
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: F401  # used in later tasks (B2+ OOXML asserts)
from docx.oxml import OxmlElement  # noqa: F401  # used in later tasks (B2+ OOXML construction)
from docx.oxml.ns import qn  # noqa: F401  # used in later tasks (B2+ OOXML asserts)
from docx.shared import Cm, Pt  # noqa: F401  # used in later tasks (B2+ layout asserts)

from app.extensions.output.layout_import import extract_layout_from_docx  # noqa: F401  # used in later tasks (B2+)
from app.extensions.output.schemas import CoverMasterSchema, CoverSlotSchema

# __file__-relative so the test runs under the canonical `make test` (pytest from
# backend/ cwd), not just from the repo root. parents[1] of backend/tests/*.py = backend/.
SAMPLE = Path(__file__).resolve().parents[1] / "data/users/f8766d55-2b1b-422e-a945-5fcf268a8a39/knowledge/8376f624-95de-47b1-b871-0bb000b5a934/基地项目-消防设计专篇.docx"


def _docx_bytes(doc: Document) -> bytes:
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Task 1: schema + wiring ────────────────────────────────────────────────


def test_cover_master_schema_camelcase_and_defaults():
    m = CoverMasterSchema(
        xml="<w:p/>",
        slots=[CoverSlotSchema(id="client", label="建设单位", sampleValue="甲公司", defaultFrom="frontmatter:client")],
    )
    d = m.model_dump()
    assert d["mode"] == "master"
    assert d["sourceFile"] == ""
    assert d["boundary"] == "before_toc"
    assert d["images"] == []
    assert d["slots"][0]["sampleValue"] == "甲公司"
    assert d["slots"][0]["defaultFrom"] == "frontmatter:client"
    assert d["slots"][0]["kind"] == "variable"


def test_build_template_data_includes_cover_master():
    from app.extensions.output.routers import _build_template_data

    tpl = SimpleNamespace(
        page_settings={},
        body_styles={},
        heading_styles=[],
        table_styles=None,
        figure_styles=None,
        header_footer=None,
        reference_style="gb7714",
        appendix_rules=None,
        cover_template=None,
        toc_settings=None,
        cover_master={"mode": "master", "xml": "<w:p/>", "images": [], "slots": [], "sourceFile": "x.docx", "boundary": "before_toc"},
    )
    td = _build_template_data(tpl)
    assert td["cover_master"]["mode"] == "master"


# ── Task 2: extraction ────────────────────────────────────────────────────


def _make_table_cover_docx() -> bytes:
    """Synthetic replica of the real sample's cover: empty spacer, title-banner
    table, a 建设单位 line, a 会签 table, a 目录 marker, then the body heading."""
    doc = Document()
    s = doc.sections[0]
    s.page_width = Cm(21.0)
    s.page_height = Cm(29.7)

    doc.add_paragraph()  # leading blank spacer

    # table[0]: title banner (single cell, big font)
    t1 = doc.add_table(rows=1, cols=1)
    banner = t1.rows[0].cells[0].paragraphs[0]
    run = banner.add_run("基地项目 消防设计专篇")
    run.font.size = Pt(22)
    run.bold = True

    doc.add_paragraph("建设单位：甲公司")  # client line (body-level paragraph)

    # table[1]: 编制会签表
    doc.add_table(rows=2, cols=2)

    doc.add_paragraph("目录")  # TOC marker (Normal style, text-based)
    doc.add_heading("第一章 概述", level=1)  # body
    return _docx_bytes(doc)


def test_extract_cover_master_table_cover():
    data = extract_layout_from_docx(_make_table_cover_docx(), source_file="sample.docx")
    cm = data["cover_master"]
    assert cm is not None
    assert cm["boundary"] == "before_toc"
    assert cm["xml"].count("<w:tbl ") == 2  # banner + 会签 (trailing space excludes <w:tblPr/<w:tblGrid/etc.)
    assert cm["images"] == []
    assert cm["sourceFile"] == "sample.docx"
    ids = {s["id"] for s in cm["slots"]}
    assert "client" in ids
    assert "title" in ids
    assert data["cover_detected"] is True


def test_extract_cover_master_none_for_plain_doc():
    """Body starts with a Heading, nothing before → no cover master."""
    doc = Document()
    doc.add_heading("第一章 概述", level=1)
    doc.add_paragraph("正文。")
    assert extract_layout_from_docx(_docx_bytes(doc))["cover_master"] is None


def test_extract_cover_master_none_for_toc_pre_region():
    """TOC-styled entries before the heading → no cover (regression guard)."""
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document()
    try:
        toc = doc.styles.add_style("toc 1", WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        toc = doc.styles["toc 1"]
    doc.add_paragraph("第一章 概述 .......... 1", style=toc)
    doc.add_heading("第一章 概述", level=1)
    assert extract_layout_from_docx(_docx_bytes(doc))["cover_master"] is None


@pytest.mark.skipif(not SAMPLE.exists(), reason="real sample not checked into repo")
def test_extract_real_sample_cover_master():
    data = extract_layout_from_docx(SAMPLE.read_bytes(), source_file="基地项目-消防设计专篇.docx")
    cm = data["cover_master"]
    assert cm is not None, "real sample should yield a cover master"
    assert cm["xml"].count("<w:tbl") >= 2
    assert cm["images"] == []
    assert cm["boundary"] in ("before_toc", "before_first_heading")
    assert cm["sourceFile"] == "基地项目-消防设计专篇.docx"
    assert "title" in {s["id"] for s in cm["slots"]}


def _png_bytes(w: int = 2, h: int = 2) -> bytes:
    """A minimal valid grayscale PNG built with stdlib (python-docx rejects some
    hand-encoded 1x1 strings; this is guaranteed parseable). Exercises the
    image-extraction (a:blip → r:embed → blob → b64) chain."""

    def chunk(tag: bytes, data: bytes) -> bytes:
        return struct.pack(">I", len(data)) + tag + data + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)

    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", w, h, 8, 0, 0, 0, 0)  # 8-bit grayscale
    raw = b"".join(b"\x00" + b"\xff" * w for _ in range(h))  # filter 0 + w white px per row
    return sig + chunk(b"IHDR", ihdr) + chunk(b"IDAT", zlib.compress(raw)) + chunk(b"IEND", b"")


_PNG_BYTES = _png_bytes()


def _make_image_cover_docx() -> bytes:
    """Cover region with an inline image + client line, then 目录 marker + body."""
    doc = Document()
    doc.add_picture(BytesIO(_PNG_BYTES), width=Cm(3))  # cover image paragraph
    doc.add_paragraph("建设单位：甲公司")
    doc.add_paragraph("目录")  # TOC marker → before_toc boundary
    doc.add_heading("第一章 概述", level=1)
    return _docx_bytes(doc)


def test_extract_cover_master_captures_image():
    """The most novel extraction path (a:blip → related_parts → base64) must be covered."""
    data = extract_layout_from_docx(_make_image_cover_docx(), source_file="logo.docx")
    cm = data["cover_master"]
    assert cm is not None
    assert cm["boundary"] == "before_toc"
    assert len(cm["images"]) == 1
    img = cm["images"][0]
    assert img["origRid"]  # non-empty rId
    assert img["ext"] == "png"
    assert base64.b64decode(img["b64"]) == _PNG_BYTES


# ── Task 3: generation ────────────────────────────────────────────────────


def test_render_cover_master_round_trip_replaces_variable():
    """Extract a master from a synthetic table-cover doc, render it into a fresh
    doc with a different client value → client slot replaced, tables present."""
    from app.extensions.output.generator import _render_cover_master

    master = extract_layout_from_docx(_make_table_cover_docx())["cover_master"]
    doc = Document()
    _render_cover_master(doc, master, {"client": "乙公司"}, {})

    assert len(doc.tables) == 2  # banner + 会签 both carried over
    all_text = "\n".join(p.text for p in doc.paragraphs) + "\n" + "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    assert "乙公司" in all_text  # replaced
    assert "甲公司" not in all_text  # old value gone
    assert "消防设计专篇" in all_text  # banner title (literal) preserved


def test_render_cover_master_literal_slot_not_replaced():
    from app.extensions.output.generator import _render_cover_master

    w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    master = {
        "mode": "master",
        "images": [],
        "sourceFile": "x",
        "boundary": "before_toc",
        "xml": f'<w:p xmlns:w="{w}"><w:r><w:t>建设单位：甲公司</w:t></w:r></w:p>',
        "slots": [{"id": "client", "label": "建设单位", "kind": "literal", "sampleValue": "甲公司", "defaultFrom": "frontmatter:client"}],
    }
    doc = Document()
    _render_cover_master(doc, master, {"client": "乙公司"}, {})
    assert "甲公司" in "\n".join(p.text for p in doc.paragraphs)  # literal kept
    assert "乙公司" not in "\n".join(p.text for p in doc.paragraphs)


def test_render_cover_master_rewrites_image_embed():
    """A master carrying one base64 image: render re-embeds it and rewrites r:embed."""
    from app.extensions.output.generator import _render_cover_master

    png_b64 = base64.b64encode(_PNG_BYTES).decode("ascii")  # valid PNG (bug-1113)
    w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    a = "http://schemas.openxmlformats.org/drawingml/2006/main"
    r = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    master = {
        "mode": "master",
        "slots": [],
        "sourceFile": "x",
        "boundary": "before_toc",
        "xml": (f'<w:p xmlns:w="{w}" xmlns:a="{a}" xmlns:r="{r}"><w:r><w:drawing><a:blip r:embed="rIdOld"/></w:drawing></w:r></w:p>'),
        "images": [{"origRid": "rIdOld", "ext": "png", "b64": png_b64}],
    }
    doc = Document()
    _render_cover_master(doc, master, {}, {})
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    image_rids = [rid for rid, rel in doc.part.rels.items() if rel.reltype == RT.IMAGE]
    assert image_rids, "image should be re-embedded"
    # M1: assert r:embed was actually rewritten to the new rId, not left stale.
    blips = doc.element.body.findall(f".//{{{a}}}blip")
    assert blips, "blip should remain in body"
    assert blips[0].get(f"{{{r}}}embed") in image_rids, "r:embed should point to the re-embedded rId"
    assert blips[0].get(f"{{{r}}}embed") != "rIdOld", "r:embed should have been rewritten"


def test_render_cover_master_strips_orphan_image_on_failure():
    """A master with an undecodable image b64: render strips the orphan <w:drawing>
    so the doc carries no missing-relationship reference (Word-repair guard, I1)."""
    from app.extensions.output.generator import _render_cover_master

    w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    a = "http://schemas.openxmlformats.org/drawingml/2006/main"
    r = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    master = {
        "mode": "master",
        "slots": [],
        "sourceFile": "x",
        "boundary": "before_toc",
        "xml": (f'<w:p xmlns:w="{w}" xmlns:a="{a}" xmlns:r="{r}"><w:r><w:drawing><a:blip r:embed="rIdOld"/></w:drawing></w:r></w:p>'),
        "images": [{"origRid": "rIdOld", "ext": "png", "b64": "AAAAA"}],  # bad padding → binascii.Error
    }
    doc = Document()
    _render_cover_master(doc, master, {}, {})
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    image_rids = [rid for rid, rel in doc.part.rels.items() if rel.reltype == RT.IMAGE]
    assert not image_rids, "no image relationship should be created for a bad b64"
    blips = doc.element.body.findall(f".//{{{a}}}blip")
    assert not blips, "orphan blip (stale rIdOld) must be stripped, not left as a broken ref"
