"""Regression: generate_docx must set w:eastAsia font so Chinese renders correctly."""
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from app.extensions.output.generator import generate_docx


def _eastasia(run) -> str | None:
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        return None
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        return None
    return rFonts.get(qn("w:eastAsia"))


def test_heading_and_body_runs_have_eastasia(tmp_path: Path):
    md = "# 总论标题\n\n这是正文段落。\n"
    tpl = {
        "body_styles": {"fontFamily": "宋体", "fontSize": 12},
        "heading_styles": [{"level": 1, "fontFamily": "黑体", "numbering": "none"}],
    }
    out = tmp_path / "t.docx"
    generate_docx(md, tpl, out)

    doc = Document(str(out))
    heading_run = next(
        (r for p in doc.paragraphs if "标题" in p.text for r in p.runs if r.text),
        None,
    )
    assert heading_run is not None, "heading run not found"
    assert _eastasia(heading_run) == "SimHei"  # _resolve_font("黑体") → SimHei

    body_run = next(
        (r for p in doc.paragraphs if "正文" in p.text for r in p.runs if r.text),
        None,
    )
    assert body_run is not None, "body run not found"
    assert _eastasia(body_run) == "SimSun"  # _resolve_font("宋体") → SimSun
