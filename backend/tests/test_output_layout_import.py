"""Tests for deterministic .docx → layout-template extraction."""

from io import BytesIO

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.extensions.output.layout_import import extract_layout_from_docx


def _docx_bytes(doc: Document) -> bytes:
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx(*, body_font: str | None = None, heading_font: str | None = None, with_table: bool = False) -> bytes:
    doc = Document()
    s = doc.sections[0]
    s.page_width = Cm(21.0)
    s.page_height = Cm(29.7)
    s.top_margin = Cm(2.0)
    s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(3.0)
    s.right_margin = Cm(3.0)
    if body_font:
        normal = doc.styles["Normal"]
        normal.font.name = body_font
        normal.font.size = Pt(12)
    if heading_font:
        h1 = doc.styles["Heading 1"]
        h1.font.name = heading_font
        h1.font.size = Pt(16)
    if with_table:
        doc.add_table(rows=2, cols=2)
    doc.add_heading("第一章 概述", level=1)
    doc.add_paragraph("这是正文内容，用于测试提取。")
    return _docx_bytes(doc)


def data_for(**kw):
    return extract_layout_from_docx(_make_docx(**kw))


def test_extracts_page_settings():
    data = extract_layout_from_docx(_make_docx())
    ps = data["page_settings"]
    assert ps["paperSize"] == "A4"
    assert ps["orientation"] == "portrait"
    assert ps["marginTop"] == 2.0
    assert ps["marginBottom"] == 2.0
    assert ps["marginLeft"] == 3.0
    assert ps["marginRight"] == 3.0


def test_extracts_body_and_heading_fonts():
    data = extract_layout_from_docx(_make_docx(body_font="SimSun", heading_font="SimHei"))
    assert data["body_styles"]["fontSize"] == 12
    assert data["body_styles"]["fontFamily"] == "SimSun"
    h1 = data["heading_styles"][0]
    assert h1["level"] == 1
    assert h1["fontFamily"] == "SimHei"
    assert h1["fontSize"] == 16
    assert h1["numbering"] == "chinese"  # heading text "第一章 …" → chinese numbering


def test_heading_color_has_hash_prefix():
    """Extracted heading color must be #RRGGBB (editor <input type=color> rejects bare hex)."""
    doc = Document()
    from docx.shared import RGBColor

    doc.styles["Heading 1"].font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
    doc.add_heading("第一章 概述", level=1)  # must actually use the level, else it's skipped
    h1 = extract_layout_from_docx(_docx_bytes(doc))["heading_styles"][0]
    assert h1["color"] == "#2B579A"


def test_body_size_read_from_runs_not_style():
    """Regression: body size set on runs (real docs, e.g. 四号=14pt) must be read, not the Normal default."""
    doc = Document()
    # Normal style left at template default; body runs explicitly at 四号 = 14pt
    for txt in ("本项目位于某工业园区。", "厂区总平面布置符合防火间距要求。"):
        run = doc.add_paragraph().add_run(txt)
        run.font.size = Pt(14)
    doc.add_heading("第一章 概述", level=1)
    assert extract_layout_from_docx(_docx_bytes(doc))["body_styles"]["fontSize"] == 14


def test_heading_size_read_from_runs_not_style():
    """Heading size overridden on the run must be read, not only the Heading style."""
    doc = Document()
    h = doc.add_heading("第一章 概述", level=1)
    h.runs[0].font.size = Pt(18)
    doc.add_paragraph("正文内容。")
    assert extract_layout_from_docx(_docx_bytes(doc))["heading_styles"][0]["fontSize"] == 18


def test_heading_styles_only_emit_levels_actually_used():
    """Regression: levels with zero Heading paragraphs must be skipped, not emitted with
    Word template-default values (e.g. an English font like Cambria) that pollute the form.
    _make_docx() has one Heading 1 paragraph and no Heading 2-4 → only level 1 may appear."""
    levels = [h["level"] for h in data_for()["heading_styles"]]
    assert levels == [1]


def test_heading_styles_emit_each_used_level():
    """A doc using Heading 1 and 2 emits exactly [1, 2] (not 1-4)."""
    doc = Document()
    doc.add_heading("第一章 概述", level=1)
    doc.add_heading("1.1 项目位置", level=2)
    doc.add_paragraph("正文。")
    levels = [h["level"] for h in extract_layout_from_docx(_docx_bytes(doc))["heading_styles"]]
    assert levels == [1, 2]


def test_table_style_present_only_when_table_exists():
    assert data_for(with_table=True)["table_styles"] is not None
    # plain table has no row banding → stripeRows must be False (not invented True)
    assert data_for(with_table=True)["table_styles"]["stripeRows"] is False
    assert data_for()["table_styles"] is None


def _shaded_cell_table(*, direct_fill: str | None = None, style_firstrow_fill: str | None = None) -> bytes:
    doc = Document()
    doc.add_heading("标题", level=1)
    table = doc.add_table(rows=2, cols=2)
    if direct_fill:
        tc_pr = table.rows[0].cells[0]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), direct_fill)
        tc_pr.append(shd)
    if style_firstrow_fill:
        table.style = "Table Grid"
        tsp = OxmlElement("w:tblStylePr")
        tsp.set(qn("w:type"), "firstRow")
        tc_pr = OxmlElement("w:tcPr")
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), style_firstrow_fill)
        tc_pr.append(shd)
        tsp.append(tc_pr)
        table.style.element.append(tsp)
    return _docx_bytes(doc)


def test_table_plain_header_is_white_not_blue():
    """Regression: a plain table (no header shading) must NOT invent a blue header."""
    ts = data_for(with_table=True)["table_styles"]
    assert ts["headerBg"] == "#FFFFFF"  # no fill detected → white, never the old #2B579A blue
    assert ts["headerColor"] == "#333333"


def test_table_direct_cell_shading_extracted():
    """A directly-shaded header cell must yield its real fill color."""
    ts = extract_layout_from_docx(_shaded_cell_table(direct_fill="D9E2F3"))["table_styles"]
    assert ts["headerBg"] == "#D9E2F3"


def test_table_style_firstrow_shading_extracted():
    """A table style with a firstRow band must yield the band fill (common Word case)."""
    ts = extract_layout_from_docx(_shaded_cell_table(style_firstrow_fill="C6E0B4"))["table_styles"]
    assert ts["headerBg"] == "#C6E0B4"


def test_table_style_banding_enables_striping():
    """stripeRows is True only when the table style defines row banding; plain → False."""
    # plain table: no banding
    assert data_for(with_table=True)["table_styles"]["stripeRows"] is False
    # banded style: defines band1Row shading
    doc = Document()
    doc.add_heading("标题", level=1)
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    tsp = OxmlElement("w:tblStylePr")
    tsp.set(qn("w:type"), "band1Row")
    tc_pr = OxmlElement("w:tcPr")
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    tc_pr.append(shd)
    tsp.append(tc_pr)
    table.style.element.append(tsp)
    assert extract_layout_from_docx(_docx_bytes(doc))["table_styles"]["stripeRows"] is True


def test_body_line_spacing_read_from_paragraphs_not_style():
    """Regression: line spacing set per-paragraph (real docs) must be read, not the Normal default."""
    doc = Document()
    for txt in ("正文段落一。", "正文段落二。"):
        p = doc.add_paragraph(txt)
        p.paragraph_format.line_spacing = 1.73
    doc.add_heading("第一章 概述", level=1)
    assert extract_layout_from_docx(_docx_bytes(doc))["body_styles"]["lineHeight"] == 1.73


def test_body_paragraph_spacing_read_from_paragraphs_not_style():
    """Regression: space_after set per-paragraph must be read, not the Normal default."""
    doc = Document()
    for txt in ("正文段落一。", "正文段落二。"):
        p = doc.add_paragraph(txt)
        p.paragraph_format.space_after = Pt(10)
    doc.add_heading("第一章 概述", level=1)
    assert extract_layout_from_docx(_docx_bytes(doc))["body_styles"]["paragraphSpacing"] == 10


def test_body_first_line_indent_derived_from_paragraphs():
    """Regression: first-line indent (Pt) is converted to 字 chars using the body size."""
    doc = Document()
    for txt in ("正文段落一。", "正文段落二。"):
        p = doc.add_paragraph(txt)
        p.runs[0].font.size = Pt(14)  # 四号
        p.paragraph_format.first_line_indent = Pt(28)  # 2 chars at 14pt
    doc.add_heading("第一章 概述", level=1)
    assert extract_layout_from_docx(_docx_bytes(doc))["body_styles"]["firstLineIndent"] == 2


def test_heading_color_read_from_runs_not_style():
    """Regression: heading color set on the run must be read, not the Heading style's color."""
    doc = Document()
    h = doc.add_heading("第一章 概述", level=1)
    h.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # red on run, style untouched
    doc.add_paragraph("正文内容。")
    assert extract_layout_from_docx(_docx_bytes(doc))["heading_styles"][0]["color"] == "#FF0000"


def test_table_border_color_extracted_from_tblBorders():
    """Regression: table border color must be read from w:tblBorders, not hardcoded gray."""
    doc = Document()
    doc.add_heading("标题", level=1)
    table = doc.add_table(rows=2, cols=2)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "000000")
        borders.append(e)
    table._tbl.tblPr.append(borders)
    ts = extract_layout_from_docx(_docx_bytes(doc))["table_styles"]
    assert ts["borderColor"] == "#000000"


def test_body_exact_line_spacing_mapped_to_multiple():
    """Exact (固定值) line spacing Length is mapped to a multiple via exact_pt ÷ body_font_pt."""
    doc = Document()
    for txt in ("正文段落一。", "正文段落二。"):
        p = doc.add_paragraph(txt)
        p.runs[0].font.size = Pt(14)  # 四号
        p.paragraph_format.line_spacing = Pt(28)  # 固定值 28pt → 28/14 = 2.0
    doc.add_heading("第一章 概述", level=1)
    assert extract_layout_from_docx(_docx_bytes(doc))["body_styles"]["lineHeight"] == 2.0


def test_heading_weight_read_from_runs_not_style():
    """Run-level bold must win over the (bold-by-default) Heading style."""
    doc = Document()
    h = doc.add_heading("第一章 概述", level=1)
    h.runs[0].font.bold = False  # explicitly not-bold on the run; style is still bold
    doc.add_paragraph("正文内容。")
    assert extract_layout_from_docx(_docx_bytes(doc))["heading_styles"][0]["fontWeight"] == 400


def test_heading_numbering_chinese_from_text():
    doc = Document()
    doc.add_heading("第一章 项目概述", level=1)
    doc.add_paragraph("正文。")
    assert extract_layout_from_docx(_docx_bytes(doc))["heading_styles"][0]["numbering"] == "chinese"


def test_heading_numbering_decimal_from_text():
    doc = Document()
    doc.add_heading("1.1 项目概述", level=1)
    doc.add_paragraph("正文。")
    assert extract_layout_from_docx(_docx_bytes(doc))["heading_styles"][0]["numbering"] == "decimal"


def test_heading_numbering_none_when_plain():
    doc = Document()
    doc.add_heading("项目概述", level=1)  # no number pattern, no numPr
    doc.add_paragraph("正文。")
    assert extract_layout_from_docx(_docx_bytes(doc))["heading_styles"][0]["numbering"] == "none"


def test_header_logo_detected_from_image():
    """An image (drawingml blip) in the header must set showLogo True (not hardcoded False)."""
    doc = Document()
    doc.add_heading("第一章 概述", level=1)
    run = doc.sections[0].header.paragraphs[0].add_run()
    run._r.append(OxmlElement("a:blip"))  # minimal drawingml image marker
    assert extract_layout_from_docx(_docx_bytes(doc))["header_footer"]["showLogo"] is True


def test_header_footer_text_extracted_from_real_parts():
    """Regression: header/footer text must be read from the real part paragraphs, not left empty."""
    doc = Document()
    doc.add_heading("第一章 概述", level=1)
    doc.sections[0].header.paragraphs[0].add_run("某公司消防设计专篇")
    doc.sections[0].footer.paragraphs[0].add_run("第 X 页")
    hf = extract_layout_from_docx(_docx_bytes(doc))["header_footer"]
    assert hf["headerText"] == "某公司消防设计专篇"
    assert hf["footerText"] == "第 X 页"


def test_header_footer_page_number_detected_from_field():
    """A PAGE field (w:fldSimple instr=PAGE) in the footer → showPageNumber True."""
    doc = Document()
    doc.add_heading("第一章 概述", level=1)
    fp = doc.sections[0].footer.paragraphs[0]
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    fp._p.append(fld)
    assert extract_layout_from_docx(_docx_bytes(doc))["header_footer"]["showPageNumber"] is True


def test_header_footer_page_number_false_when_no_field():
    """No PAGE field anywhere → showPageNumber False (not a hardcoded True)."""
    doc = Document()
    doc.add_heading("第一章 概述", level=1)
    doc.sections[0].footer.paragraphs[0].add_run("固定文字页脚")  # plain text, no field
    assert extract_layout_from_docx(_docx_bytes(doc))["header_footer"]["showPageNumber"] is False


def test_figure_styles_null():
    assert data_for()["figure_styles"] is None


def _blip_run(para) -> None:
    """Append a minimal drawingml image marker to a paragraph (mimics an embedded figure)."""
    run = para.add_run()
    run._r.append(OxmlElement("a:blip"))


def _doc_with_figure(*, caption: str | None = None, caption_above: bool = False, source: str | None = None) -> bytes:
    """A doc with one figure (image paragraph) and an optional caption / source line."""
    doc = Document()
    doc.add_heading("第一章 概述", level=1)
    doc.add_paragraph("正文内容。")
    if caption and caption_above:
        doc.add_paragraph(caption)
    _blip_run(doc.add_paragraph())  # the figure (image paragraph)
    if caption and not caption_above:
        doc.add_paragraph(caption)
    if source:
        doc.add_paragraph(source)
    return _docx_bytes(doc)


def test_figure_styles_none_when_no_figures_or_captions():
    """No image and no caption text → figure_styles None (leave the form's figure config alone)."""
    assert data_for()["figure_styles"] is None


def test_figure_caption_position_below():
    """A caption paragraph right after an image → captionPosition 'below'."""
    fs = extract_layout_from_docx(_doc_with_figure(caption="图1 厂区平面图"))["figure_styles"]
    assert fs["captionPosition"] == "below"


def test_figure_caption_position_above():
    """A caption paragraph right before an image → captionPosition 'above'."""
    fs = extract_layout_from_docx(_doc_with_figure(caption="图1 厂区平面图", caption_above=True))["figure_styles"]
    assert fs["captionPosition"] == "above"


def test_figure_numbering_chapter():
    """A chapter-segmented caption (图1-1) → numbering 'chapter'."""
    fs = extract_layout_from_docx(_doc_with_figure(caption="图1-1 厂区平面图"))["figure_styles"]
    assert fs["numbering"] == "chapter"


def test_figure_numbering_continuous():
    """A plain sequential caption (图1) → numbering 'continuous'."""
    fs = extract_layout_from_docx(_doc_with_figure(caption="图1 厂区平面图"))["figure_styles"]
    assert fs["numbering"] == "continuous"


def test_figure_show_source_true_when_source_line_present():
    """A 数据来源 line near the figure → showSource True."""
    fs = extract_layout_from_docx(_doc_with_figure(caption="图1 厂区平面图", source="数据来源：自绘"))["figure_styles"]
    assert fs["showSource"] is True


def test_figure_show_source_false_when_no_source_line():
    """A figure with no source line → showSource False (not a hardcoded True)."""
    fs = extract_layout_from_docx(_doc_with_figure(caption="图1 厂区平面图"))["figure_styles"]
    assert fs["showSource"] is False


def test_no_cover_detected_for_plain_document():
    assert data_for()["cover_detected"] is False


def test_cover_logo_position_left_for_default_left_alignment():
    doc = Document()
    s = doc.sections[0]
    s.page_width = Cm(21.0)
    s.page_height = Cm(29.7)
    doc.add_paragraph("某化工项目消防设计专篇")  # no explicit alignment → None → left
    doc.add_paragraph("建设单位：某某公司")
    doc.add_paragraph("项目编号：P001")
    doc.add_heading("第一章 概述", level=1)
    data = extract_layout_from_docx(_docx_bytes(doc))
    assert data["cover_detected"] is True
    assert data["cover_template"]["logoPosition"] == "left"


def test_cover_detected_for_first_page_cover():
    doc = Document()
    s = doc.sections[0]
    s.page_width = Cm(21.0)
    s.page_height = Cm(29.7)
    title = doc.add_paragraph("某化工项目消防设计专篇")
    title.runs[0].font.size = Pt(22)
    doc.add_paragraph("建设单位：某某公司")
    doc.add_paragraph("项目编号：P001")
    doc.add_paragraph("2026-08-01")
    doc.add_heading("第一章 概述", level=1)
    data = extract_layout_from_docx(_docx_bytes(doc))
    assert data["cover_detected"] is True
    ct = data["cover_template"]
    assert ct["showTitle"] is True
    assert ct["showClient"] is True
    assert ct["showProjectNumber"] is True
    assert ct["showDate"] is True


def test_rejects_non_docx_bytes():
    with pytest.raises(ValueError):
        extract_layout_from_docx(b"this is definitely not a docx file")


def test_output_router_registers_import_layout():
    from app.extensions.output.routers import router

    paths = set()
    for route in router.routes:
        for method in getattr(route, "methods", None) or set():
            paths.add((route.path, method))
    assert ("/api/extensions/output/import-layout", "POST") in paths, "output import-layout route missing"


def test_validate_docx_upload_rejects_non_docx_filename():
    from app.extensions.output.layout_import import validate_docx_upload

    with pytest.raises(ValueError, match="仅支持 .docx 文件"):
        validate_docx_upload("report.pdf", b"PK\x03\x04" + b"\x00" * 100)


def test_validate_docx_upload_rejects_oversize():
    from app.extensions.output.layout_import import validate_docx_upload

    with pytest.raises(ValueError, match="不能超过 10MB"):
        validate_docx_upload("report.docx", b"PK\x03\x04" + b"\x00" * (10 * 1024 * 1024 + 1))


def test_validate_docx_upload_returns_extraction():
    from app.extensions.output.layout_import import validate_docx_upload

    doc = Document()
    s = doc.sections[0]
    s.page_width = Cm(21.0)
    s.page_height = Cm(29.7)
    data = _docx_bytes(doc)
    result = validate_docx_upload("report.docx", data)
    assert result["page_settings"]["paperSize"] == "A4"
