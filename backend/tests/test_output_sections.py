"""Tests for per-section page numbering + footer PAGE field."""
from docx import Document

from app.extensions.output.generator import _add_page_number_footer, _set_section_pagenum


def test_set_section_pagenum_upper_roman():
    doc = Document()
    _set_section_pagenum(doc.sections[0], fmt="upperRoman", start=1)
    xml = doc.sections[0]._sectPr.xml
    assert "w:pgNumType" in xml
    assert 'w:fmt="upperRoman"' in xml
    assert 'w:start="1"' in xml


def test_set_section_pagenum_decimal():
    doc = Document()
    _set_section_pagenum(doc.sections[0], fmt="decimal", start=1)
    xml = doc.sections[0]._sectPr.xml
    assert 'w:fmt="decimal"' in xml


def test_set_section_pagenum_idempotent():
    doc = Document()
    _set_section_pagenum(doc.sections[0], fmt="decimal", start=1)
    _set_section_pagenum(doc.sections[0], fmt="decimal", start=1)
    assert doc.sections[0]._sectPr.xml.count("w:pgNumType") == 1


def test_add_page_number_footer_inserts_page_field():
    doc = Document()
    section = doc.sections[0]
    section.footer.is_linked_to_previous = False
    _add_page_number_footer(section)
    xml = section.footer.paragraphs[0]._element.xml
    assert "PAGE" in xml
    assert 'fldCharType="begin"' in xml
    assert 'fldCharType="end"' in xml
