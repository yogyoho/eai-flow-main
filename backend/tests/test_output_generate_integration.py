"""End-to-end: generate_docx produces cover + TOC + body sections with numbering."""
from pathlib import Path

from docx import Document

from app.extensions.output.generator import generate_docx


TPL_WITH_COVER_TOC = {
    "page_settings": {"paperSize": "A4", "orientation": "portrait",
                      "marginTop": 2.54, "marginBottom": 2.54, "marginLeft": 3.17, "marginRight": 3.17},
    "body_styles": {"fontFamily": "宋体", "fontSize": 12, "lineHeight": 1.5,
                    "paragraphSpacing": 6, "firstLineIndent": 2},
    "heading_styles": [
        {"level": 1, "fontFamily": "黑体", "fontSize": 16, "fontWeight": 700, "color": "#000000", "numbering": "decimal"},
        {"level": 2, "fontFamily": "黑体", "fontSize": 14, "fontWeight": 700, "color": "#000000", "numbering": "decimal"},
    ],
    "cover_template": {"showLogo": False, "showTitle": True, "showClient": True,
                       "showDate": True, "showProjectNumber": True},
    "toc_settings": {"maxDepth": 2, "showPageNumbers": True, "leaderDots": True},
    "header_footer": {"headerText": "", "footerText": "", "showPageNumber": True, "showLogo": False},
}


def test_three_sections_cover_toc_body(tmp_path: Path):
    md = "# 总论\n\n正文段。\n## 子节\n\n更多正文。\n"
    out = tmp_path / "r.docx"
    generate_docx(md, TPL_WITH_COVER_TOC, out,
                  cover_fields={"title": "消防专篇", "client": "甲公司",
                                "date": "2026-07", "project_number": "P001"})
    doc = Document(str(out))
    assert len(doc.sections) == 3  # cover / toc / body


def test_cover_section_has_no_page_number(tmp_path: Path):
    out = tmp_path / "r.docx"
    generate_docx("# 总论\n", TPL_WITH_COVER_TOC, out,
                  cover_fields={"title": "T", "date": "2026-07"})
    doc = Document(str(out))
    assert "pgNumType" not in doc.sections[0]._sectPr.xml


def test_toc_section_is_upper_roman(tmp_path: Path):
    out = tmp_path / "r.docx"
    generate_docx("# 总论\n", TPL_WITH_COVER_TOC, out, cover_fields={"title": "T", "date": "2026-07"})
    doc = Document(str(out))
    assert 'w:fmt="upperRoman"' in doc.sections[1]._sectPr.xml


def test_body_section_is_decimal_restart_1(tmp_path: Path):
    out = tmp_path / "r.docx"
    generate_docx("# 总论\n", TPL_WITH_COVER_TOC, out, cover_fields={"title": "T", "date": "2026-07"})
    doc = Document(str(out))
    body_xml = doc.sections[2]._sectPr.xml
    assert 'w:fmt="decimal"' in body_xml
    assert 'w:start="1"' in body_xml


def test_toc_field_and_updatefields_present(tmp_path: Path):
    out = tmp_path / "r.docx"
    generate_docx("# 总论\n", TPL_WITH_COVER_TOC, out, cover_fields={"title": "T", "date": "2026-07"})
    doc = Document(str(out))
    assert 'TOC \\o "1-2"' in doc.element.body.xml
    assert "updateFields" in doc.settings.element.xml


def test_body_headings_carry_decimal_numbers(tmp_path: Path):
    md = "# 总论\n## 子节\n# 二章\n"
    out = tmp_path / "r.docx"
    generate_docx(md, TPL_WITH_COVER_TOC, out, cover_fields={"title": "T", "date": "2026-07"})
    doc = Document(str(out))
    heading_texts = [p.text for p in doc.paragraphs if p.style and p.style.name and p.style.name.startswith("Heading")]
    assert "1 总论" in heading_texts
    assert "1.1 子节" in heading_texts
    assert "2 二章" in heading_texts


def test_cover_title_rendered(tmp_path: Path):
    out = tmp_path / "r.docx"
    generate_docx("# 总论\n", TPL_WITH_COVER_TOC, out,
                  cover_fields={"title": "我的消防专篇", "date": "2026-07"})
    doc = Document(str(out))
    assert any("我的消防专篇" in p.text for p in doc.paragraphs)


def test_backward_compat_no_cover_no_toc(tmp_path: Path):
    """Template without cover_template/toc_settings → single section, no cover/TOC."""
    tpl = {
        "page_settings": {"paperSize": "A4"},
        "body_styles": {"fontFamily": "宋体", "fontSize": 12},
        "heading_styles": [{"level": 1, "fontFamily": "黑体", "numbering": "none"}],
    }
    out = tmp_path / "r.docx"
    generate_docx("# 总论\n正文\n", tpl, out)
    doc = Document(str(out))
    assert "TOC" not in doc.element.body.xml
    # no cover title rendered (no cover_template)
    assert not any("建设单位" in p.text for p in doc.paragraphs)
