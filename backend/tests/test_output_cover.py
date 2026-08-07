"""Tests for cover-page rendering + cover field resolution."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.extensions.output.generator import Block, _render_cover, _resolve_cover_fields


def _texts(doc):
    return [p.text for p in doc.paragraphs]


def test_renders_all_fields_when_all_shown():
    ct = {"showLogo": False, "showTitle": True, "showClient": True, "showDate": True, "showProjectNumber": True}
    cf = {"title": "消防专篇", "client": "甲公司", "date": "2026-07", "project_number": "P001"}
    doc = Document()
    _render_cover(doc, ct, cf)
    txt = _texts(doc)
    assert "消防专篇" in txt
    assert any("建设单位" in t and "甲公司" in t for t in txt)
    assert any("项目编号" in t and "P001" in t for t in txt)
    assert any("日期" in t and "2026-07" in t for t in txt)


def test_skips_line_when_value_missing():
    ct = {"showTitle": True, "showClient": True, "showDate": False, "showProjectNumber": True, "showLogo": False}
    cf = {"title": "T"}  # client/date/project_number 全缺
    doc = Document()
    _render_cover(doc, ct, cf)
    txt = _texts(doc)
    assert "T" in txt
    assert not any("建设单位" in t for t in txt), "no client value → no 建设单位 line"
    assert not any("项目编号" in t for t in txt)


def test_skips_line_when_toggle_false():
    ct = {"showTitle": True, "showClient": False, "showDate": False, "showProjectNumber": False, "showLogo": False}
    cf = {"title": "T", "client": "C", "date": "D", "project_number": "P"}
    doc = Document()
    _render_cover(doc, ct, cf)
    txt = _texts(doc)
    assert not any("建设单位" in t for t in txt)


def test_title_centered_heiti():
    ct = {"showTitle": True, "showClient": False, "showDate": False, "showProjectNumber": False, "showLogo": False}
    cf = {"title": "标题X"}
    doc = Document()
    _render_cover(doc, ct, cf)
    title_para = next(p for p in doc.paragraphs if p.text == "标题X")
    assert title_para.alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_no_cover_when_template_none():
    doc = Document()
    _render_cover(doc, None, {"title": "T"})
    # only spacer/empty paragraphs, no title rendered
    assert "T" not in _texts(doc)


def test_renders_logo_placeholder_when_shown():
    ct = {"showLogo": True, "showTitle": False, "showClient": False, "showDate": False, "showProjectNumber": False}
    doc = Document()
    _render_cover(doc, ct, {})
    assert any("LOGO" in t for t in _texts(doc))


# --- _resolve_cover_fields (Task 7) ---


def test_resolve_prefers_api_over_frontmatter_over_fallback():
    api = {"title": "API标题", "client": "API客户"}
    fm = {"title": "FM标题", "date": "2026-01"}
    blocks = [Block(kind="heading", level=1, text="H1兜底")]
    resolved = _resolve_cover_fields(api, fm, blocks)
    assert resolved["title"] == "API标题"  # api wins
    assert resolved["client"] == "API客户"  # api
    assert resolved["date"] == "2026-01"  # frontmatter (no api)


def test_resolve_title_falls_back_to_first_h1():
    api = {}
    fm = {}
    blocks = [Block(kind="paragraph", text="p"), Block(kind="heading", level=1, text="首个H1")]
    resolved = _resolve_cover_fields(api, fm, blocks)
    assert resolved["title"] == "首个H1"


def test_resolve_date_falls_back_to_today():
    import datetime

    api = {}
    fm = {}
    blocks = []
    resolved = _resolve_cover_fields(api, fm, blocks)
    assert resolved["date"] == datetime.date.today().isoformat()


def test_resolve_omits_missing_optional_fields():
    resolved = _resolve_cover_fields({}, {}, [])
    assert "client" not in resolved
    assert "project_number" not in resolved


def test_logo_position_alignment():
    """M2: logo paragraph alignment follows cover_template.logoPosition."""
    cases = [
        ("left", WD_ALIGN_PARAGRAPH.LEFT),
        ("center", WD_ALIGN_PARAGRAPH.CENTER),
        ("right", WD_ALIGN_PARAGRAPH.RIGHT),
        (None, WD_ALIGN_PARAGRAPH.CENTER),  # legacy default
    ]
    for pos, expected in cases:
        ct = {
            "showLogo": True,
            "showTitle": False,
            "showClient": False,
            "showDate": False,
            "showProjectNumber": False,
            "logoPosition": pos,
        }
        doc = Document()
        _render_cover(doc, ct, {})
        logo_para = next(p for p in doc.paragraphs if "[编制单位 LOGO]" in p.text)
        assert logo_para.alignment == expected, f"logoPosition={pos} → {expected}"


def test_logo_position_absent_key_defaults_center():
    """M2: a legacy cover_template WITHOUT a logoPosition key renders centered."""
    ct = {
        "showLogo": True,
        "showTitle": False,
        "showClient": False,
        "showDate": False,
        "showProjectNumber": False,
    }
    doc = Document()
    _render_cover(doc, ct, {})
    logo_para = next(p for p in doc.paragraphs if "[编制单位 LOGO]" in p.text)
    assert logo_para.alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_cover_slot_value_keys():
    """Cross-language drift guard: the generator's resolvable slot id set must
    equal the set mirrored by frontend cover-state.ts COVER_RESOLVABLE_SLOT_IDS
    (title/client/project_number/date/project_name/stage)."""
    from app.extensions.output.generator import COVER_SLOT_VALUE_KEYS

    assert set(COVER_SLOT_VALUE_KEYS) == {
        "title",
        "client",
        "project_number",
        "date",
        "project_name",
        "stage",
    }


def test_cover_render_failure_is_logged_not_silent(caplog):
    """M5: a failing cover render logs a warning and must not abort generation."""
    import logging
    from pathlib import Path

    from app.extensions.output.generator import generate_docx

    tpl = {
        "page_settings": {"paperSize": "A4", "orientation": "portrait", "marginTop": 2.54, "marginBottom": 2.54, "marginLeft": 3.17, "marginRight": 3.17},
        "body_styles": {"fontFamily": "宋体", "fontSize": 12, "lineHeight": 1.5, "paragraphSpacing": 0, "firstLineIndent": 2},
        "heading_styles": [],
        "cover_master": {"mode": "master", "xml": "<w:p>", "images": [], "slots": [], "sourceFile": "bad.docx", "boundary": "before_toc"},
        "cover_template": None,
        "toc_settings": None,
    }
    out = Path("_m5_cover_fail_test.docx")
    try:
        with caplog.at_level(logging.WARNING, logger="app.extensions.output.generator"):
            result = generate_docx("# 正文\n", tpl, out)
        assert "cover render failed" in caplog.text, "cover failure must be logged, not silent"
        assert result, "generation must not abort on cover failure"
        assert out.exists()
    finally:
        out.unlink(missing_ok=True)


def test_cover_preset_failure_logged_simple_branch(caplog):
    """M5: the generate_docx_simple cover path (docmgr export) also logs, not silent."""
    import logging
    from io import BytesIO

    from app.extensions.output.generator import generate_docx_simple

    buf = BytesIO()
    with caplog.at_level(logging.WARNING, logger="app.extensions.output.generator"):
        generate_docx_simple(
            "# 正文\n",
            buf,
            cover_preset={"elements": [{"type": "spacer", "lines": "abc"}]},
            cover_values={},
        )
    assert "cover render failed" in caplog.text
    assert buf.getbuffer().nbytes > 0  # generation did not abort
