"""三分支解析：docx→md（含 OLE/OMML 占位）、pdf 文字版、扫描判定异常、OCR 契约与统一分发。"""

import io

import pytest


def _build_docx() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("第1章 总论", level=1)
    doc.add_paragraph("本矿区位于云南省。")
    doc.add_heading("1.1 编制依据", level=2)
    doc.add_paragraph("依据DZ/T 0214-2020。")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "矿体编号"
    table.cell(0, 1).text = "品位%"
    table.cell(1, 0).text = "Ⅰ号"
    table.cell(1, 1).text = "0.85"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_docx_to_markdown_structure():
    from app.extensions.geo_samples.parsers import docx_to_markdown

    md = docx_to_markdown(_build_docx())
    assert md.startswith("## 第1章 总论")
    assert "### 1.1 编制依据" in md
    assert "| 矿体编号 | 品位% |" in md
    assert "| --- | --- |" in md
    assert "| Ⅰ号 | 0.85 |" in md


def test_docx_ole_formula_placeholder():
    """段落含 OLE 对象（公式）且无文本 → [公式:pN] 占位。"""
    from docx import Document
    from docx.oxml.ns import qn

    from app.extensions.geo_samples.parsers import docx_to_markdown

    doc = Document()
    p = doc.add_paragraph()
    p._p.append(p._p.makeelement(qn("w:object"), {}))
    buf = io.BytesIO()
    doc.save(buf)
    assert "[公式:p1]" in docx_to_markdown(buf.getvalue())


def test_docx_inline_and_omml_formula_placeholders():
    """I3：有文本的段落夹公式对象 → 文本保留 + [公式:pN] 后缀；纯 OMML(m:oMath) → 独立占位。"""
    from docx import Document
    from docx.oxml.ns import qn

    from app.extensions.geo_samples.parsers import docx_to_markdown

    doc = Document()
    p1 = doc.add_paragraph("比拟法估算资源量")
    p1._p.append(p1._p.makeelement(qn("w:object"), {}))  # 行内 OLE
    p2 = doc.add_paragraph()
    p2._p.append(p2._p.makeelement(qn("m:oMath"), {}))  # 纯 OMML 公式
    buf = io.BytesIO()
    doc.save(buf)
    md = docx_to_markdown(buf.getvalue())
    assert "比拟法估算资源量 [公式:p1]" in md  # 文本 + 占位后缀（计数器从 1 起）
    assert "[公式:p2]" in md  # 无文本 OMML → 独立占位行，计数递增


def test_pdf_scan_detection_raises():
    from app.extensions.geo_samples.parsers import ScannedPdfError, pdf_text_to_markdown

    # 1 页几乎无文本 → 判定扫描件
    fake_pdf = _make_pdf_with_text("仅几个字")
    with pytest.raises(ScannedPdfError):
        pdf_text_to_markdown(fake_pdf)


def _make_pdf_with_text(text: str) -> bytes:
    fitz = pytest.importorskip("fitz")
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text)
    return doc.tobytes()


def test_ocr_dispatch_contracts(monkeypatch):
    from app.extensions.geo_samples import parsers

    captured = {}

    class FakeResp:
        def raise_for_status(self):
            pass

        def json(self):
            return {
                "pages": [
                    {"page_no": 1, "text": "OCR 第一页"},
                    {"page_no": 2, "text": "OCR 第二页", "tables": [{"rows": [["钻孔", "深度"], ["ZK1", "100"]]}]},
                ]
            }

    class FakeClient:
        def __init__(self, *a, **k):
            pass

        async def __aenter__(self):
            return self

        async def __aexit__(self, *a):
            return False

        async def post(self, url, files=None, data=None):
            captured["url"] = url
            captured["files"] = files
            captured["data"] = data
            return FakeResp()

    monkeypatch.setattr(parsers.httpx, "AsyncClient", FakeClient)

    import asyncio

    md = asyncio.run(parsers.ocr_pdf_to_markdown(b"pdfbytes"))
    assert "http" in captured["url"] and captured["url"].endswith("/ocr")
    assert "OCR 第一页" in md and "OCR 第二页" in md
    # C1.b：显式放开整页文字 OCR 页数门控（默认只前 3 页）
    assert captured["data"] == {"text_pages": "999"}
    # C1.b：tables[].rows 拍平为 md 管道表（无表头分隔行），跟在该页文本之后
    assert "| 钻孔 | 深度 |" in md
    assert "| ZK1 | 100 |" in md
    assert md.index("OCR 第二页") < md.index("| 钻孔 | 深度 |")


def test_ocr_empty_result_guard(monkeypatch):
    """C1.c：OCR 全空（无文本无表格）→ ValueError，不静默返回空 md。"""
    from app.extensions.geo_samples import parsers

    class FakeResp:
        def raise_for_status(self):
            pass

        def json(self):
            return {"pages": [{"page_no": 1, "text": "  ", "tables": []}]}

    class FakeClient:
        def __init__(self, *a, **k):
            pass

        async def __aenter__(self):
            return self

        async def __aexit__(self, *a):
            return False

        async def post(self, url, files=None, data=None):
            return FakeResp()

    monkeypatch.setattr(parsers.httpx, "AsyncClient", FakeClient)

    import asyncio

    with pytest.raises(ValueError, match="OCR 返回空文本"):
        asyncio.run(parsers.ocr_pdf_to_markdown(b"pdfbytes"))


def test_ocr_retry_on_remote_protocol_error(monkeypatch):
    """M3：对端半途断流（RemoteProtocolError）重试一次后成功；退避不真睡。"""
    import asyncio

    from app.extensions.geo_samples import parsers

    calls = {"post": 0, "slept": []}

    class FakeResp:
        def raise_for_status(self):
            pass

        def json(self):
            return {"pages": [{"page_no": 1, "text": "重试后文本"}]}

    class FakeClient:
        def __init__(self, *a, **k):
            pass

        async def __aenter__(self):
            return self

        async def __aexit__(self, *a):
            return False

        async def post(self, url, files=None, data=None):
            calls["post"] += 1
            if calls["post"] == 1:
                raise parsers.httpx.RemoteProtocolError("peer closed connection")
            return FakeResp()

    async def fake_sleep(sec):
        calls["slept"].append(sec)

    monkeypatch.setattr(parsers.httpx, "AsyncClient", FakeClient)
    monkeypatch.setattr(asyncio, "sleep", fake_sleep)

    md = asyncio.run(parsers.ocr_pdf_to_markdown(b"pdfbytes"))
    assert "重试后文本" in md
    assert calls["post"] == 2
    assert calls["slept"] == [30.0]  # 首败退避 30s


def test_parse_document_dispatch(monkeypatch):
    """M1：统一分发——docx 大小写不敏感；不支持类型 ValueError；稀疏 pdf 落到 OCR 通道。"""
    import asyncio

    from app.extensions.geo_samples import parsers

    async def fake_ocr(data, base_url=None):
        return "OCR TEXT"

    monkeypatch.setattr(parsers, "ocr_pdf_to_markdown", fake_ocr)

    md, mode = asyncio.run(parsers.parse_document("报告.DOCX", _build_docx()))
    assert mode == "docx" and md.startswith("## 第1章 总论")

    with pytest.raises(ValueError, match="不支持的文件类型"):
        asyncio.run(parsers.parse_document("photo.png", b"x"))

    md, mode = asyncio.run(parsers.parse_document("scan.pdf", _make_pdf_with_text("短")))
    assert mode == "pdf_ocr" and md == "OCR TEXT"


def test_docx_localized_heading_styles():
    """中文 Word 样式名（标题 1/标题 2）须归一化进标题映射——R1 防线。"""
    from docx import Document

    from app.extensions.geo_samples.parsers import docx_to_markdown

    doc = Document()
    doc.add_paragraph("第1章 总论", style="Heading 1")
    p = doc.add_paragraph("1.1 编制依据")
    p.style = doc.styles["Heading 2"]
    # 保存前把样式改成真实中文 authored docx 的本地化 w:name 做 round-trip——
    # 令本测试只在别名表在场时通过，否则整篇降级正文、断言失败。
    doc.styles["Heading 1"].name = "标题 1"
    doc.styles["Heading 2"].name = "标题 2"
    buf = io.BytesIO()
    doc.save(buf)
    md = docx_to_markdown(buf.getvalue())
    assert "## 第1章 总论" in md and "### 1.1 编制依据" in md


def test_docx_style_alias_map():
    from app.extensions.geo_samples.parsers import _heading_level

    assert _heading_level("Heading 1") == 1
    assert _heading_level("Title") == 1
    assert _heading_level("标题 1") == 1
    assert _heading_level("标题 2") == 2
    assert _heading_level("标题 3") == 3
    assert _heading_level("标题 4") == 3
    assert _heading_level("标题 5") == 3
    assert _heading_level("Heading 4") == 3
    assert _heading_level("Normal") is None


def test_heading_level_whitespace_variants():
    """Minor 2：空白变体（双空格/首尾/全角/nbsp）须归一，不静默降级正文。"""
    from app.extensions.geo_samples.parsers import _heading_level

    assert _heading_level("标题  1") == 1  # 双空格
    assert _heading_level("  Heading 2  ") == 2  # 首尾空白
    assert _heading_level("标题　3") == 3  # 全角空格
    assert _heading_level("Heading 4") == 3  # 不换行空格


def test_heading_level_deep_levels():
    """Minor 3：Heading 6-9 与「3+ 一律 ####」契约一致，不静默降级正文。"""
    from app.extensions.geo_samples.parsers import _heading_level

    assert _heading_level("Heading 6") == 3
    assert _heading_level("Heading 9") == 3
    assert _heading_level("标题 6") == 3
    assert _heading_level("标题 9") == 3
