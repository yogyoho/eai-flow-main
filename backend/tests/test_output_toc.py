"""Tests for TOC field rendering + updateFields setting."""

from docx import Document

from app.extensions.output.generator import _render_toc, _set_update_fields


def test_toc_field_present_with_maxdepth():
    doc = Document()
    _render_toc(doc, {"maxDepth": 2})
    xml = doc.element.body.xml
    assert 'TOC \\o "1-2"' in xml
    assert "目录" in xml  # 目录 heading


def test_toc_uses_configured_maxdepth():
    doc = Document()
    _render_toc(doc, {"maxDepth": 3})
    assert 'TOC \\o "1-3"' in doc.element.body.xml


def test_toc_skipped_when_settings_none():
    doc = Document()
    _render_toc(doc, None)
    assert "TOC" not in doc.element.body.xml


def test_toc_skipped_when_maxdepth_zero():
    doc = Document()
    _render_toc(doc, {"maxDepth": 0})
    assert "TOC" not in doc.element.body.xml


def test_update_fields_written_to_settings():
    doc = Document()
    _set_update_fields(doc)
    assert "updateFields" in doc.settings.element.xml
    assert 'val="true"' in doc.settings.element.xml


def test_update_fields_idempotent():
    doc = Document()
    _set_update_fields(doc)
    _set_update_fields(doc)
    assert doc.settings.element.xml.count("updateFields") == 1
