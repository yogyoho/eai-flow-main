"""Regression tests for document-space export (Markdown + Word).

Root cause these guard against: the docmgr export endpoints and the
``generate_docx_simple`` helper were implemented on a divergent branch and
only ever captured in a git stash (c97665ec). They never landed on the
working branch, so the frontend's export menu hit non-existent routes and
both Markdown and Word export failed with 404. These tests lock the
endpoints + generator in place so they cannot silently disappear again.
"""

from io import BytesIO


def _docmgr_paths():
    """Return the set of (path, method) tuples registered on the docmgr router."""
    from app.extensions.docmgr.routers import router

    paths = set()
    for route in router.routes:
        methods = getattr(route, "methods", None) or set()
        for method in methods:
            paths.add((route.path, method))
    return paths


def test_docmgr_router_registers_get_export():
    """GET /documents/{doc_id}/export must be registered (Markdown / simple export)."""
    paths = _docmgr_paths()
    assert (
        "/api/extensions/docmgr/documents/{doc_id}/export",
        "GET",
    ) in paths, "docmgr GET export route is missing — Markdown export would 404"


def test_docmgr_router_registers_post_export():
    """POST /documents/{doc_id}/export must be registered (Word export with layout)."""
    paths = _docmgr_paths()
    assert (
        "/api/extensions/docmgr/documents/{doc_id}/export",
        "POST",
    ) in paths, "docmgr POST export route is missing — Word export would 404"


def test_generate_docx_simple_produces_valid_docx():
    """generate_docx_simple must emit a non-empty, valid OOXML (.docx) buffer."""
    from app.extensions.output.generator import generate_docx_simple

    md = "# 测试标题\n\n这是正文段落。\n\n## 二级标题\n\n- 列表项一\n- 列表项二\n"
    buf = BytesIO()
    generate_docx_simple(md, buf)

    data = buf.getvalue()
    assert len(data) > 100, "generated docx buffer is suspiciously small"
    # .docx is a ZIP archive — must start with the PK\x03\x04 magic bytes.
    assert data[:2] == b"PK", "generated buffer is not a valid .docx (ZIP) file"


def test_generate_docx_simple_applies_watermark_and_layout():
    """generate_docx_simple must accept layout_template + watermark without error."""
    from app.extensions.output.generator import generate_docx_simple

    md = "# 报告\n\n正文内容。"
    buf = BytesIO()
    generate_docx_simple(
        md,
        buf,
        template_data={
            "page_settings": {"paperSize": "A4", "orientation": "portrait"},
            "body_styles": {"fontFamily": "宋体", "fontSize": 12},
            "heading_styles": [{"level": 1, "fontSize": 22, "fontWeight": 700}],
            "header_footer": {"footerText": "页脚", "showPageNumber": True},
        },
        watermark="draft",
    )
    data = buf.getvalue()
    assert data[:2] == b"PK"


def test_generate_docx_simple_with_toc_emits_field_and_update_flag():
    """with toc_settings(maxDepth=3) → document.xml 含 TOC \\o "1-3";settings.xml 含 updateFields。"""
    import zipfile

    from app.extensions.output.generator import generate_docx_simple

    md = "# 一级标题\n\n正文。\n\n## 二级标题\n\n更多正文。\n"
    buf = BytesIO()
    generate_docx_simple(md, buf, toc_settings={"maxDepth": 3})
    buf.seek(0)
    with zipfile.ZipFile(buf) as z:
        document_xml = z.read("word/document.xml").decode("utf-8")
        settings_xml = z.read("word/settings.xml").decode("utf-8")
    assert 'TOC \\o "1-3"' in document_xml, "TOC field not injected"
    assert "w:updateFields" in settings_xml, "updateFields flag not set"


def test_generate_docx_simple_without_toc_is_unchanged():
    """默认(无 toc_settings)→ 不含 TOC 域、不含 updateFields。"""
    import zipfile

    from app.extensions.output.generator import generate_docx_simple

    md = "# 一级标题\n\n正文。\n"
    buf = BytesIO()
    generate_docx_simple(md, buf)  # 不传 toc_settings
    buf.seek(0)
    with zipfile.ZipFile(buf) as z:
        document_xml = z.read("word/document.xml").decode("utf-8")
        settings_xml = z.read("word/settings.xml").decode("utf-8")
    assert "TOC \\o" not in document_xml
    assert "updateFields" not in settings_xml


def test_export_requests_carry_toc_fields():
    """ExportRequest / ExportContentRequest 必须接受 with_toc / toc_depth,且默认 关/3。"""
    from app.extensions.docmgr.routers import ExportContentRequest, ExportRequest

    er = ExportRequest(with_toc=True, toc_depth=2)
    assert er.with_toc is True
    assert er.toc_depth == 2

    er_default = ExportRequest()
    assert er_default.with_toc is False
    assert er_default.toc_depth == 3

    ec = ExportContentRequest(with_toc=True, toc_depth=4)
    assert ec.with_toc is True
    assert ec.toc_depth == 4

    ec_default = ExportContentRequest()
    assert ec_default.with_toc is False
    assert ec_default.toc_depth == 3


def test_render_cover_preset_renders_fields():
    """_render_cover_preset renders title (text) + info lines for provided values."""
    from docx import Document

    from app.extensions.output.generator import _render_cover_preset

    preset = {
        "id": "t",
        "label": "T",
        "fields": [{"name": "title"}, {"name": "client"}],
        "elements": [
            {"type": "text", "field": "title", "align": "center", "font": "黑体", "size": 22, "bold": True},
            {"type": "info", "label": "建设单位", "field": "client"},
        ],
    }
    doc = Document()
    rendered = _render_cover_preset(doc, preset, {"title": "测试报告", "client": "测试单位"})
    assert rendered is True
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    import zipfile

    xml = zipfile.ZipFile(buf).read("word/document.xml").decode("utf-8")
    assert "测试报告" in xml
    assert "建设单位" in xml
    assert "测试单位" in xml


def test_render_cover_preset_skips_missing_value():
    """An info/text element whose field value is missing is skipped entirely (no empty label)."""
    from docx import Document

    from app.extensions.output.generator import _render_cover_preset

    preset = {
        "id": "t",
        "elements": [
            {"type": "text", "field": "title"},
            {"type": "info", "label": "建设单位", "field": "client"},
        ],
    }
    doc = Document()
    _render_cover_preset(doc, preset, {"title": "有标题"})  # client 缺值
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    import zipfile

    xml = zipfile.ZipFile(buf).read("word/document.xml").decode("utf-8")
    assert "有标题" in xml
    assert "建设单位" not in xml  # 缺值的 info 行整行不渲染


def test_render_cover_preset_none_is_noop():
    """preset=None → returns False, renders nothing."""
    from docx import Document

    from app.extensions.output.generator import _render_cover_preset

    doc = Document()
    assert _render_cover_preset(doc, None, {"title": "x"}) is False


def test_cover_presets_module_shape():
    """COVER_PRESETS has the fire_protection preset; get/public helpers work."""
    from app.extensions.output.cover_presets import (
        COVER_PRESETS,
        get_cover_preset,
        public_cover_presets,
    )

    ids = [p["id"] for p in COVER_PRESETS]
    assert "fire_protection" in ids

    fire = get_cover_preset("fire_protection")
    assert fire is not None and "elements" in fire

    assert get_cover_preset("does_not_exist") is None

    public = public_cover_presets()
    assert all("elements" not in p for p in public)  # trimmed view has no layout
    pub_fire = next(p for p in public if p["id"] == "fire_protection")
    assert pub_fire["label"] == "消防设计专篇"
    assert any(f["name"] == "title" for f in pub_fire["fields"])


def test_generate_docx_simple_with_cover_sections_and_pagenum():
    """With cover: 2 sections; cover (s0) has no PAGE footer; body (s1) restarts at page 1."""
    import zipfile

    from docx import Document
    from docx.oxml.ns import qn

    from app.extensions.output.cover_presets import get_cover_preset
    from app.extensions.output.generator import generate_docx_simple

    preset = get_cover_preset("fire_protection")
    md = "## 引言\n\n正文内容。\n"
    buf = BytesIO()
    generate_docx_simple(
        md,
        buf,
        cover_preset=preset,
        cover_values={"title": "封面标题", "client": "某单位", "date": "2026-07-22"},
    )
    buf.seek(0)
    doc = Document(buf)
    assert len(doc.sections) == 2, "cover + body must produce two sections"

    body = doc.sections[-1]
    pgnum = body._sectPr.find(qn("w:pgNumType"))
    assert pgnum is not None, "body section must have pgNumType"
    assert pgnum.get(qn("w:start")) == "1", "body page numbers must restart at 1"

    cover = doc.sections[0]
    assert "PAGE" not in cover.footer._element.xml, "cover page must have no page-number field"

    buf.seek(0)
    xml = zipfile.ZipFile(buf).read("word/document.xml").decode("utf-8")
    assert "封面标题" in xml  # cover rendered
    assert "引言" in xml  # body rendered


def test_generate_docx_simple_without_cover_still_single_section():
    """No cover → single section (no-cover path unchanged)."""
    from docx import Document

    from app.extensions.output.generator import generate_docx_simple

    buf = BytesIO()
    generate_docx_simple("## 引言\n\n正文。\n", buf)
    buf.seek(0)
    assert len(Document(buf).sections) == 1


def test_docmgr_router_registers_cover_presets():
    """GET /cover-presets must be registered."""
    paths = _docmgr_paths()
    assert ("/api/extensions/docmgr/cover-presets", "GET") in paths


def test_export_requests_carry_cover_fields():
    """ExportRequest / ExportContentRequest accept cover_preset_id / cover_values (default None)."""
    from app.extensions.docmgr.routers import ExportContentRequest, ExportRequest

    er = ExportRequest(cover_preset_id="fire_protection", cover_values={"title": "T"})
    assert er.cover_preset_id == "fire_protection"
    assert er.cover_values == {"title": "T"}
    er_default = ExportRequest()
    assert er_default.cover_preset_id is None
    assert er_default.cover_values is None

    ec = ExportContentRequest(cover_preset_id="fire_protection")
    assert ec.cover_preset_id == "fire_protection"
    assert ExportContentRequest().cover_preset_id is None
    assert ExportContentRequest().cover_values is None


def test_docmgr_router_registers_import_layout():
    """POST /import-layout must be registered — the ExportDocxDialog 导入排版 button depends on it."""
    from app.extensions.docmgr.routers import router

    paths = set()
    for route in router.routes:
        for method in getattr(route, "methods", None) or set():
            paths.add((route.path, method))
    assert (
        "/api/extensions/docmgr/import-layout",
        "POST",
    ) in paths, "docmgr import-layout route is missing — 导入排版 button 404s"


# ===========================================================================
# v4 T8: 多册合并导出(整单失败语义 + 册间分页) — bug-3109 WP-1.4
# ===========================================================================


class TestGenerateDocxMerged:
    """generate_docx_merged: 依序合并+册间分页; 整单失败指认册名; 超限拒绝。"""

    def _render(self, sections, **kwargs):
        from io import BytesIO

        from app.extensions.output.generator import generate_docx_merged

        buf = BytesIO()
        generate_docx_merged(sections, buf, **kwargs)
        return buf.getvalue()

    def test_merged_contains_sections_in_order_with_breaks(self):
        from docx import Document

        data = self._render([
            ("整体方案-01-投标函.md", "# 投标函\n\n致:招标人"),
            ("技术卷-01-总体设计.md", "# 总体设计\n\n分层架构说明。"),
        ])
        doc = Document(BytesIO(data))
        texts = [p.text for p in doc.paragraphs]
        joined = "\n".join(texts)
        assert "投标函" in joined and "总体设计" in joined
        idx_fun = next(i for i, t in enumerate(texts) if "投标函" in t)
        idx_tech = next(i for i, t in enumerate(texts) if "总体设计" in t)
        assert idx_fun < idx_tech, "册序=合并序"
        has_break = any('w:br w:type="page"' in p._element.xml for p in doc.paragraphs)
        assert has_break, "册间存在真实分页"

    def test_parse_failure_names_the_booklet(self):
        """整单失败语义: 任一册解析异常 → 错误指认册名, 不产部分文件。"""
        import pytest

        # 构造解析期异常: 传入不可迭代内容之外的最直接路径 = 预解析钩子
        # (parse_markdown 对任意 str 都健壮——这里用 monkeypatch 注入解析失败)
        import app.extensions.output.generator as gen

        original = gen.parse_markdown
        calls = []

        def fake_parse(md):
            calls.append(md)
            if md == "# bad":
                raise RuntimeError("boom")
            return original(md)

        gen.parse_markdown = fake_parse
        try:
            with pytest.raises(ValueError) as exc_info:
                gen.generate_docx_merged([("好册.md", "# ok"), ("坏册.md", "# bad")], BytesIO())
        finally:
            gen.parse_markdown = original
        assert "坏册.md" in str(exc_info.value) and "整单导出中止" in str(exc_info.value)
        assert len(calls) == 2, "预解析逐册归因(坏册之前的册已扫过)"

    def test_oversize_and_count_rejected(self):
        import pytest

        from app.extensions.output.generator import MAX_MERGE_SECTIONS, MAX_MERGE_TOTAL_CHARS, generate_docx_merged

        with pytest.raises(ValueError, match="超上限"):
            generate_docx_merged([(f"s{i}.md", "# x") for i in range(MAX_MERGE_SECTIONS + 1)], BytesIO())
        with pytest.raises(ValueError, match="内存预算"):
            generate_docx_merged([("huge.md", "字" * (MAX_MERGE_TOTAL_CHARS + 1))], BytesIO())

    def test_empty_sections_rejected(self):
        import pytest

        from app.extensions.output.generator import generate_docx_merged

        with pytest.raises(ValueError, match="为空"):
            generate_docx_merged([], BytesIO())
