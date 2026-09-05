#!/usr/bin/env python3
"""bid-proposal-writing 技能测试 fixture 生成器（设计 D4 测试计划，任务 T1）。

规格: docs/superpowers/specs/2026-08-16-bid-proposal-writing-skill-design.md

再生成方法（backend venv）:
    cd backend && uv run python tests/fixtures/bid_proposal/gen_fixtures.py

产物（全部确定性，无随机因子）:
    minimal_tender.docx      最小招标文件: 多级标题树 + ★强制条款句 + "实质性响应"条款
                             + 评分细则表(4行x4列, 含分值列) + 参数表(3行x3列)
    minimal_tender.pdf       手写最小合法 PDF 字节（backend venv 无 reportlab/fpdf 等生成库，
                             按 fixture 生成阶梯降级为字节级构造; ASCII 内容, Helvetica,
                             含 /Title Info 结构; pdfplumber 可解析出 3 个标题 + 若干正文行）
    clauses.json             义务清单样例（严格按设计文档条款 schema 字段表）
    structure.json           双卷结构镜像样例（注意 D7: fill_status 为派生字段, 不落盘）
    rubric.json              评分标尺样例（Σmax_score == total_score == 评分办法总分）
    rubric_bad_sum.json      故意 Σ 不一致负样本（供 T4 extract.py / T7 score_simulate.py
                             的 Σmax_score 校验测试消费）
    sections.json            ingest 产物样例（chunk_id/table_id/锚点/表行数）
    returned_word.md         团队回传稿样例（商务卷=标题链镜像; 技术卷=clause_id 埋锚;
                             故意含 1 个标题链匹配失败项 + 1 个重复 clause_id 项, 供 T7 D6 测试）
    entities_whitelist.json  实体白名单样例（确认门1 锁定形态）
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

FIXTURE_DIR = Path(__file__).resolve().parent

# ---------------------------------------------------------------------------
# 样例源文件: minimal_tender.docx
# ---------------------------------------------------------------------------


def build_tender_docx(out_dir: Path) -> None:
    """python-docx 生成最小招标文件（多级标题/正文/★强制句/实质性响应句/评分表/参数表）。"""
    from docx import Document

    doc = Document()
    doc.core_properties.title = "XX水泥厂二期智能控制系统采购项目招标文件(最小样例)"

    # --- 第一章 投标邀请 (H1 > H2 章节树) ---
    doc.add_heading("第一章 投标邀请", level=1)
    doc.add_paragraph("招标编号:EAI-BID-2026-017。项目名称:XX水泥厂二期智能控制系统采购项目。招标人:XX水泥有限责任公司。")
    doc.add_heading("一、项目概况", level=2)
    doc.add_paragraph("本项目采购智能控制系统一套,用于二期生产线技术改造,交货期为合同签订后90天。")
    doc.add_heading("二、投标人须知", level=2)
    doc.add_paragraph("★本项为强制条款:投标文件必须加盖投标人公章并由法定代表人或其授权代理人签字,任何偏离将导致废标。")
    doc.add_paragraph("投标人对招标文件第三章技术规范中的各项技术要求应作出实质性响应,未作实质性响应的其投标将被否决。")

    # --- 第三章 技术规范 (H1 > H2 > H3 章节树) ---
    doc.add_heading("第三章 技术规范", level=1)
    doc.add_heading("3.2 设备清单及技术参数", level=2)
    doc.add_heading("3.2.1 技术参数要求", level=3)
    doc.add_paragraph("设备防护等级不低于IP65,控制器采用西门子S7-1500系列(参数版本V2.3)。")
    doc.add_heading("3.2.2 技术参数表", level=3)

    # 参数表 T-001: 表头 + 2 数据行 = 3 行 x 3 列（合并语义从略, 见任务说明）
    param_table = doc.add_table(rows=3, cols=3)
    param_table.style = "Table Grid"
    for row, cells in enumerate(
        [
            ["参数名称", "单位", "要求值"],
            ["设备防护等级", "-", "≥IP65"],
            ["控制器型号", "-", "S7-1500(V2.3)"],
        ]
    ):
        for col, value in enumerate(cells):
            param_table.rows[row].cells[col].text = value

    # --- 第六章 评标办法 (评分细则表) ---
    doc.add_heading("第六章 评标办法", level=1)
    doc.add_paragraph("本项目采用综合评分法,总分100分。评分因素及分值分配见评分细则表。")
    doc.add_heading("6.1 评分细则", level=2)

    # 评分细则表 T-002: 表头 + 3 数据行 = 4 行 x 4 列, 含分值列
    score_table = doc.add_table(rows=4, cols=4)
    score_table.style = "Table Grid"
    for row, cells in enumerate(
        [
            ["序号", "评分项", "分值", "评分标准"],
            ["1", "投标报价", "60", "以有效投标中的最低价为基准价,投标报价等于基准价得满分,每高1%扣0.6分"],
            ["2", "技术方案先进性", "15", "优=12-15 良=8-11 一般=1-7 无=0"],
            ["3", "设备性能参数满足度", "25", "完全满足得25分,一般偏离每项扣5分,重大偏离得0分"],
        ]
    ):
        for col, value in enumerate(cells):
            score_table.rows[row].cells[col].text = value

    doc.save(str(out_dir / "minimal_tender.docx"))


# ---------------------------------------------------------------------------
# 样例源文件: minimal_tender.pdf（手写最小合法 PDF 字节）
# ---------------------------------------------------------------------------

_PDF_LINES: list[tuple[str, int]] = [
    # (文本, 字号) — 3 个标题 + 4 行正文, 单页
    ("1. Tender Notice", 16),
    ("This tender notice defines the general requirements for bidders.", 10),
    ("Bidders shall read all clauses carefully before submitting offers.", 10),
    ("2. Technical Specifications", 16),
    ("The equipment shall meet all parameters in the parameter table.", 10),
    ("3. Evaluation Method", 16),
    ("The scoring details are listed in the evaluation table of this document.", 10),
]


def _pdf_escape(text: str) -> str:
    """转义 PDF literal string 中的 () 与反斜杠。"""
    return text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")


def build_tender_pdf_bytes() -> bytes:
    """构造单页最小合法 PDF（Helvetica, ASCII 内容, 含 /Title Info）。

    无第三方生成库可用时的阶梯降级方案——xref 偏移逐对象精确计算,
    pdfminer/pdfplumber 可直接解析文本。
    """
    content_parts: list[str] = []
    y = 720
    for text, size in _PDF_LINES:
        content_parts.append(f"BT /F1 {size} Tf 72 {y} Td ({_pdf_escape(text)}) Tj ET")
        y -= 28
    stream = "\n".join(content_parts).encode("ascii")

    objects: dict[int, bytes] = {
        1: b"<< /Type /Catalog /Pages 2 0 R >>",
        2: b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        3: b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        4: b"<< /Length %d >>\nstream\n%s\nendstream" % (len(stream), stream),
        5: b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica /Encoding /WinAnsiEncoding >>",
        6: b"<< /Title (Minimal Tender Fixture) /Producer (gen_fixtures.py) >>",
    }

    out = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets: dict[int, int] = {}
    for num in sorted(objects):
        offsets[num] = len(out)
        out += f"{num} 0 obj\n".encode("ascii") + objects[num] + b"\nendobj\n"
    xref_pos = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode("ascii")
    out += b"0000000000 65535 f \n"
    for num in sorted(objects):
        out += f"{offsets[num]:010d} 00000 n \n".encode("ascii")
    out += f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R /Info 6 0 R >>\nstartxref\n{xref_pos}\n%%EOF\n".encode("ascii")
    return bytes(out)


# ---------------------------------------------------------------------------
# 样例状态 JSON（严格按设计文档三 schema 字段表 + sections/白名单）
# ---------------------------------------------------------------------------

# 义务清单 clauses.json —— 4 条: mandatory/scoring/normal 齐备;
# ZB-C-003 被 BY-C-004 (补遗条款) supersede; BY-C-004 来自补遗 PDF (page 锚点)。
CLAUSES = [
    {
        "clause_id": "ZB-C-001",
        "source_file": "minimal_tender.docx",
        "class": "mandatory",
        "category": "technical",
        "source_ref": {"page": None, "section": "3.2.1", "para": 1, "quote": "设备防护等级不低于IP65,控制器采用S7-1500系列"},
        "requirement": "设备防护等级不低于IP65,控制器采用西门子S7-1500系列(参数版本V2.3)——★强制条款,零偏离",
        "response_status": "compliant",
        "response_skeleton": {"points": ["防护等级IP65", "控制器S7-1500 V2.3"], "evidence_ref": None, "suggestion": None},
        "from_addendum": False,
        "superseded_by": None,
        "voided": False,
    },
    {
        "clause_id": "ZB-C-002",
        "source_file": "minimal_tender.docx",
        "class": "scoring",
        "category": "technical",
        "source_ref": {"page": None, "section": "6.1", "para": 1, "quote": "技术方案先进性 15 优=12-15 良=8-11 一般=1-7 无=0"},
        "requirement": "评分项:技术方案先进性,满分15分,按优/良/一般/无四档评分",
        "response_status": "draft",
        "response_skeleton": {"points": ["总体架构先进性", "与工艺场景契合度"], "evidence_ref": None, "suggestion": None},
        "from_addendum": False,
        "superseded_by": None,
        "voided": False,
    },
    {
        "clause_id": "ZB-C-003",
        "source_file": "minimal_tender.docx",
        "class": "normal",
        "category": "commercial",
        "source_ref": {"page": None, "section": "一、项目概况", "para": 1, "quote": "交货期为合同签订后90天"},
        "requirement": "交货期:合同签订后90天内交货(原条款,已被补遗文件-01修订)",
        "response_status": "unassigned",
        "response_skeleton": {"points": [], "evidence_ref": None, "suggestion": None},
        "from_addendum": False,
        "superseded_by": "BY-C-004",
        "voided": False,
    },
    {
        "clause_id": "BY-C-004",
        "source_file": "补遗文件-01.pdf",
        "class": "normal",
        "category": "commercial",
        "source_ref": {"page": 1, "section": "二", "para": 1, "quote": "交货期统一调整为合同签订后60天"},
        "requirement": "交货期:合同签订后60天内交货(补遗文件-01修订版)",
        "response_status": "pending_confirm",
        "response_skeleton": {"points": ["生产排期压缩至60天"], "evidence_ref": None, "suggestion": None},
        "from_addendum": True,
        "superseded_by": None,
        "voided": False,
    },
]

# 结构镜像 structure.json —— 双卷 + 全部 5 种槽位类型。
# D7: fill_status 为派生字段,渲染/重灌时现算,严禁落盘——此处刻意不含该键。
STRUCTURE = [
    {
        "node_id": "S-001",
        "volume": "commercial",
        "path": "投标文件格式",
        "slot_type": "group",
        "required_format": {"desc": "商务卷:严格按招标文件规定的格式逐项提供,不得自创结构", "table_spec": None, "template_text": None},
        "linked_clause_ids": [],
    },
    {
        "node_id": "S-002",
        "volume": "commercial",
        "path": "投标文件格式/一、投标函",
        "slot_type": "text",
        "required_format": {"desc": "按给定格式填报投标函并加盖投标人公章", "table_spec": None, "template_text": None},
        "linked_clause_ids": ["BY-C-004"],
    },
    {
        "node_id": "S-003",
        "volume": "commercial",
        "path": "投标文件格式/二、法定代表人身份证明",
        "slot_type": "image",
        "required_format": {"desc": "加盖公章的身份证正反面扫描件", "table_spec": None, "template_text": None},
        "linked_clause_ids": [],
    },
    {
        "node_id": "S-004",
        "volume": "commercial",
        "path": "投标文件格式/三、开标一览表",
        "slot_type": "table",
        "required_format": {
            "desc": "按下列列头复刻开标一览表格式",
            "template_text": None,
            "table_spec": {"columns": ["序号", "货物名称", "数量", "总价(元)"], "rows": 3},
        },
        "linked_clause_ids": ["BY-C-004"],
    },
    {
        "node_id": "S-005",
        "volume": "commercial",
        "path": "投标文件格式/四、投标文件签章与份数",
        "slot_type": "format_check",
        "required_format": {"desc": "正本壹份副本肆份,全部页面加盖公章——人核项,不做确定性判定", "table_spec": None, "template_text": None},
        "linked_clause_ids": [],
    },
    {
        "node_id": "S-006",
        "volume": "technical",
        "path": "技术部分",
        "slot_type": "group",
        "required_format": {"desc": "技术卷:格式章节未规定结构处按技术参数条款清单驱动组织", "table_spec": None, "template_text": None},
        "linked_clause_ids": [],
    },
    {
        "node_id": "S-007",
        "volume": "technical",
        "path": "技术部分/1 技术方案",
        "slot_type": "text",
        "required_format": {"desc": "技术方案总述,对应评分项'技术方案先进性'", "table_spec": None, "template_text": None},
        "linked_clause_ids": ["ZB-C-002"],
    },
    {
        "node_id": "S-008",
        "volume": "technical",
        "path": "技术部分/2 技术参数响应表",
        "slot_type": "table",
        "required_format": {
            "desc": "技术参数逐项响应表,条目标题内嵌 clause_id(D2 锚点契约)",
            "template_text": None,
            "table_spec": {"columns": ["序号", "招标要求", "响应情况", "满足状态"], "rows": 4},
        },
        "linked_clause_ids": ["ZB-C-001"],
    },
]

# 评分标尺 rubric.json —— price/subjective/objective 齐备, Σmax_score=100=评分办法总分。
# 与 minimal_tender.docx 评分细则表逐行一致(可作 T3 ingest→T4 extract 联动样例)。
RUBRIC = {
    "total_score": 100,
    "items": [
        {
            "rubric_id": "R-001",
            "item": "投标报价",
            "max_score": 60,
            "scoring_method": "以有效投标中的最低价为基准价,投标报价等于基准价得满分,每高1%扣0.6分",
            "score_type": "price",
            "linked_clause_ids": [],
            "source_ref": {"page": None, "section": "6.1", "para": 2, "quote": "投标报价 60 以有效投标中的最低价为基准价"},
        },
        {
            "rubric_id": "R-002",
            "item": "技术方案先进性",
            "max_score": 15,
            "scoring_method": "优=12-15 良=8-11 一般=1-7 无=0",
            "score_type": "subjective",
            "linked_clause_ids": ["ZB-C-002"],
            "source_ref": {"page": None, "section": "6.1", "para": 3, "quote": "技术方案先进性 15 优=12-15 良=8-11 一般=1-7 无=0"},
        },
        {
            "rubric_id": "R-003",
            "item": "设备性能参数满足度",
            "max_score": 25,
            "scoring_method": "完全满足得25分,一般偏离每项扣5分,重大偏离得0分",
            "score_type": "objective",
            "linked_clause_ids": ["ZB-C-001"],
            "source_ref": {"page": None, "section": "6.1", "para": 4, "quote": "设备性能参数满足度 25 完全满足得25分"},
        },
    ],
}

# 负样本: R-002 max_score 改为 12 → Σ=97 ≠ total_score=100, 供 Σmax_score 校验测试。
RUBRIC_BAD_SUM = {
    "total_score": 100,
    "items": [
        dict(RUBRIC["items"][0]),
        {**RUBRIC["items"][1], "max_score": 12},
        dict(RUBRIC["items"][2]),
    ],
}

# ingest 产物 sections.json —— chunk_id/table_id/锚点(docx=section+para, PDF=page+section)/表行数。
# 锚点与 clauses.json / rubric.json 的 source_ref 保持一致(extract.py 校验"锚点必须存在")。
SECTIONS = {
    "chunks": [
        {
            "chunk_id": "CH-001",
            "source_file": "minimal_tender.docx",
            "anchor": {"section": "3.2.1", "para": 1},
            "heading_path": ["第三章 技术规范", "3.2 设备清单及技术参数", "3.2.1 技术参数要求"],
            "n_paras": 1,
        },
        {
            "chunk_id": "CH-002",
            "source_file": "minimal_tender.docx",
            "anchor": {"section": "6.1", "para": 1},
            "heading_path": ["第六章 评标办法", "6.1 评分细则"],
            "n_paras": 1,
        },
        {
            "chunk_id": "CH-003",
            "source_file": "minimal_tender.pdf",
            "anchor": {"page": 1, "section": "2"},
            "heading_path": ["2. Technical Specifications"],
            "n_paras": 1,
        },
        {
            "chunk_id": "CH-004",
            "source_file": "补遗文件-01.pdf",
            "anchor": {"page": 1, "section": "二"},
            "heading_path": ["二、补遗内容"],
            "n_paras": 1,
        },
        {
            "chunk_id": "CH-005",
            "source_file": "minimal_tender.docx",
            "anchor": {"section": "一、项目概况", "para": 1},
            "heading_path": ["第一章 投标邀请", "一、项目概况"],
            "n_paras": 1,
        },
    ],
    "tables": [
        {
            "table_id": "T-001",
            "source_file": "minimal_tender.docx",
            "anchor": {"section": "3.2.2", "para": 1},
            "n_rows": 3,
            "n_cols": 3,
            "caption": "技术参数表",
        },
        {
            "table_id": "T-002",
            "source_file": "minimal_tender.docx",
            "anchor": {"section": "6.1", "para": 2},
            "n_rows": 4,
            "n_cols": 4,
            "caption": "评分细则表",
        },
    ],
}

# 回传稿 returned_word.md —— 商务卷标题链镜像 structure.json path; 技术卷条目嵌 clause_id;
# 故意留 1 个标题链匹配失败项("3 售后服务承诺"不在镜像中) + 1 个重复 clause_id 项(ZB-C-001 出现两次)。
RETURNED_WORD_MD = """# 投标文件格式

## 一、投标函

致:XX水泥厂二期智能控制系统采购项目招标人。
我方投标总价为人民币玖佰捌拾万元整(¥9,800,000.00),投标有效期90天。

## 二、法定代表人身份证明

(此处已插入加盖公章的法定代表人身份证正反面扫描件)

## 三、开标一览表

| 序号 | 货物名称 | 数量 | 总价(元) |
| --- | --- | --- | --- |
| 1 | 智能控制系统 | 1 套 | 9800000 |

## 四、投标文件签章与份数

正本壹份,副本肆份,全部页面已加盖投标人公章。

# 技术部分

## 1 技术方案

本技术方案针对XX水泥厂二期智能控制系统采购项目编制,围绕总体架构先进性与工艺场景契合度展开。

## 2 技术参数响应表

### 2.1 响应[ZB-C-001]

设备防护等级 IP65,控制器采用西门子 S7-1500 系列(参数版本 V2.3),完全响应★强制条款。

### 2.2 响应[ZB-C-002]

技术方案先进性:详见"1 技术方案"章节的架构说明与工艺契合分析。

### 2.3 响应[ZB-C-001]

(重复出现的 clause_id——D6 匹配器硬化用例,应进异常区而非静默取首个)

## 3 售后服务承诺

本标题链不在 structure.json 镜像中——D6 匹配失败用例,重灌应标 needs_human_verify。
"""

# 实体白名单 entities_whitelist.json —— 确认门1 锁定形态。
ENTITIES_WHITELIST = {
    "locked_at": "2026-08-16T12:00:00+08:00",
    "source": "确认门1 锁定(封面/投标人须知抽取+人工增删确认)",
    "entities": [
        {"type": "project", "value": "XX水泥厂二期智能控制系统采购项目"},
        {"type": "company", "value": "东智装备制造有限公司"},
        {"type": "spec_version", "value": "S7-1500 V2.3"},
        {"type": "person", "value": "王建国"},
    ],
}


def _write_json(out_dir: Path, name: str, data) -> None:
    text = json.dumps(data, ensure_ascii=False, indent=2) + "\n"
    (out_dir / name).write_text(text, encoding="utf-8", newline="\n")


def main() -> int:
    parser = argparse.ArgumentParser(description="Regenerate bid-proposal-writing test fixtures (deterministic)")
    parser.add_argument("--out", default=str(FIXTURE_DIR), help="output directory (default: alongside this script)")
    args = parser.parse_args()

    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)

    build_tender_docx(out_dir)
    (out_dir / "minimal_tender.pdf").write_bytes(build_tender_pdf_bytes())
    _write_json(out_dir, "clauses.json", CLAUSES)
    _write_json(out_dir, "structure.json", STRUCTURE)
    _write_json(out_dir, "rubric.json", RUBRIC)
    _write_json(out_dir, "rubric_bad_sum.json", RUBRIC_BAD_SUM)
    _write_json(out_dir, "sections.json", SECTIONS)
    _write_json(out_dir, "entities_whitelist.json", ENTITIES_WHITELIST)
    (out_dir / "returned_word.md").write_text(RETURNED_WORD_MD, encoding="utf-8", newline="\n")

    print(f"[gen_fixtures] wrote 9 fixtures to {out_dir}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
