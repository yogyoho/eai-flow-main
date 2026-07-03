"""Synthetic docx fixture for unit tests (avoids depending on the real sample)."""
from pathlib import Path
from docx import Document


def build_tiny_spec(path):
    """A minimal 设计说明书: TOC + 2 numbered sections + 1 table with a 表号 caption.

    Section 2.1 has a conflicting datum scenario mirroring §9.1 vs §9.2:
    a 给水 paragraph says 生活水 8L/s/DN150, a 消防 paragraph says 30L/s/DN200.
    """
    doc = Document()
    doc.add_paragraph("目  录")
    for line in ["1 概述\t1", "1.1 概况\t1", "2 给排水及消防\t2", "2.1 消防\t2"]:
        doc.add_paragraph(line)
    doc.add_paragraph("概述")
    doc.add_paragraph("1.1 概况")
    doc.add_paragraph("本项目为基地综合大队，占地面积23.8亩，特勤消防站。")
    doc.add_paragraph("2 给排水及消防")
    doc.add_paragraph("2.1 消防")
    # 给水 paragraph (the WRONG source for fire water — must not be picked)
    doc.add_paragraph("生活水系统设计室外消火栓水量30L/s，生活用水量8L/s，管径DN150，自东侧市政管网引入。")
    # 消防 paragraph (the AUTHORITATIVE source for fire water)
    doc.add_paragraph("表2.1-1 消防水量")
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "项目"
    t.rows[0].cells[1].text = "水量"
    t.rows[1].cells[0].text = "室外消火栓"
    t.rows[1].cells[1].text = "30L/s"
    doc.add_paragraph("本项目设计室外消火栓水量30L/s（108m³/h），生活用水量10L/s（36m³/h），引入两根管径DN200的管线。")
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return str(path)
