#!/usr/bin/env python3
"""
Markdown to Word Document Converter

This script converts Markdown content to Microsoft Word (.docx) documents
using the python-docx library.
"""

import argparse
import os
import re
import sys
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx is not installed. Install it with: pip install python-docx")
    sys.exit(1)


def parse_markdown_to_document(markdown_text: str, doc: Document, title: Optional[str] = None, author: Optional[str] = None):
    """Parse Markdown text and add content to the Word document."""

    # Set document metadata
    if title:
        doc.core_properties.title = title
    if author:
        doc.core_properties.author = author

    lines = markdown_text.split('\n')
    i = 0
    in_code_block = False
    code_content = []
    code_language = ""

    # Track list states
    bullet_list_started = False
    numbered_list_started = False

    # Table processing
    in_table = False
    table_data = []

    while i < len(lines):
        line = lines[i]

        # Handle code blocks
        if line.strip().startswith('```'):
            if not in_code_block:
                # Start of code block
                in_code_block = True
                code_language = line.strip()[3:].strip()
                code_content = []
            else:
                # End of code block - add as preformatted text
                in_code_block = False
                if code_content:
                    # Add code block as a paragraph with monospace font
                    p = doc.add_paragraph()
                    p.style = 'Normal'
                    for line in code_content:
                        run = p.add_run(line + '\n')
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        # Handle horizontal rules / page breaks
        if re.match(r'^---+$', line.strip()) or re.match(r'^\*\*\*+$', line.strip()):
            doc.add_page_break()
            i += 1
            continue

        # Handle tables
        if '|' in line:
            # Check if it's a table row
            if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                # This is a table separator line, skip it
                i += 1
                continue

            # Parse table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                table_data.append(cells)
                in_table = True
            i += 1
            continue
        elif in_table and table_data:
            # End of table, create it
            if len(table_data) > 0:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Light Grid Accent 1'

                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_text in enumerate(row_data):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = cell_text
                        # Bold header row
                        if row_idx == 0:
                            cell.paragraphs[0].runs[0].bold = True
                table_data = []
            in_table = False
            continue

        # Handle headings
        heading_match = re.match(r'^(#{1,9})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()

            # Close any open lists
            bullet_list_started = False
            numbered_list_started = False

            # Add heading
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Handle bullet lists
        bullet_match = re.match(r'^[\-\*]\s+(.+)$', line)
        if bullet_match:
            text = bullet_match.group(1).strip()

            # Remove checkboxes
            text = re.sub(r'^\[\s*[xX]?\s*\]\s*', '', text)

            doc.add_paragraph(text, style='List Bullet')
            bullet_list_started = True
            numbered_list_started = False
            i += 1
            continue

        # Handle numbered lists
        numbered_match = re.match(r'^(\d+)\.\s+(.+)$', line)
        if numbered_match:
            text = numbered_match.group(2).strip()

            # Remove checkboxes
            text = re.sub(r'^\[\s*[xX]?\s*\]\s*', '', text)

            doc.add_paragraph(text, style='List Number')
            numbered_list_started = True
            bullet_list_started = False
            i += 1
            continue

        # Close lists if we hit a non-list line
        if line.strip() and not line.startswith(' '):
            bullet_list_started = False
            numbered_list_started = False

        # Handle empty lines
        if not line.strip():
            i += 1
            continue

        # Handle regular paragraphs with inline formatting
        paragraph = doc.add_paragraph()

        # Process inline formatting
        process_inline_formatting(line, paragraph)

        i += 1

    # Close any remaining table
    if in_table and table_data:
        if len(table_data) > 0:
            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            table.style = 'Light Grid Accent 1'

            for row_idx, row_data in enumerate(table_data):
                for col_idx, cell_text in enumerate(row_data):
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = cell_text
                    if row_idx == 0:
                        cell.paragraphs[0].runs[0].bold = True


def process_inline_formatting(text: str, paragraph):
    """Process inline Markdown formatting (bold, italic, etc.)."""

    # Patterns for inline formatting
    patterns = [
        (r'\*\*\*(.+?)\*\*\*', 'bold_italic'),  # ***bold italic***
        (r'___(.+?)___', 'bold_italic'),          # ___bold italic___
        (r'\*\*(.+?)\*\*', 'bold'),               # **bold**
        (r'__(.+?)__', 'bold'),                   # __bold__
        (r'\*(.+?)\*', 'italic'),                 # *italic*
        (r'_(.+?)_', 'italic'),                   # _italic_
        (r'~~(.+?)~~', 'strikethrough'),          # ~~strikethrough~~~
    ]

    # Track positions
    pos = 0
    matches = []

    for pattern, fmt_type in patterns:
        for match in re.finditer(pattern, text):
            matches.append((match.start(), match.end(), match.group(1), fmt_type))

    # Sort by position
    matches.sort(key=lambda x: x[0])

    # Build paragraph with formatting
    for start, end, content, fmt_type in matches:
        # Add text before match
        if pos < start:
            paragraph.add_run(text[pos:start])

        # Add formatted text
        run = paragraph.add_run(content)
        if fmt_type == 'bold':
            run.bold = True
        elif fmt_type == 'italic':
            run.font.italic = True
        elif fmt_type == 'bold_italic':
            run.bold = True
            run.font.italic = True
        elif fmt_type == 'strikethrough':
            run.fontstrike = True

        pos = end

    # Add remaining text
    if pos < len(text):
        paragraph.add_run(text[pos:])


def main():
    parser = argparse.ArgumentParser(description='Convert Markdown to Word Document')
    parser.add_argument('--markdown', type=str, help='Path to Markdown input file')
    parser.add_argument('--content', type=str, help='Inline Markdown content')
    parser.add_argument('--output', type=str, required=True, help='Path to output Word document (.docx)')
    parser.add_argument('--title', type=str, help='Document title')
    parser.add_argument('--author', type=str, help='Author name')

    args = parser.parse_args()

    # Validate output format
    output_path = Path(args.output)
    if output_path.suffix.lower() != '.docx':
        print("Error: Output must be a .docx file")
        sys.exit(1)

    # Get markdown content
    if args.markdown:
        if not os.path.exists(args.markdown):
            print(f"Error: Input file not found: {args.markdown}")
            sys.exit(1)
        with open(args.markdown, 'r', encoding='utf-8') as f:
            markdown_text = f.read()
    elif args.content:
        markdown_text = args.content
    else:
        print("Error: Either --markdown or --content must be provided")
        sys.exit(1)

    # Create output directory if needed
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Create Word document
    doc = Document()

    # Parse and convert
    parse_markdown_to_document(markdown_text, doc, args.title, args.author)

    # Save
    doc.save(args.output)
    print(f"Successfully created: {args.output}")


if __name__ == '__main__':
    main()
