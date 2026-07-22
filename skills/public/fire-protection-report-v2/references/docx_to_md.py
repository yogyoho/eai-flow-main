#!/usr/bin/env python3
"""将 .docx 转成 Markdown（供消防设计专篇技能读取上传的设计说明书）。

为什么存在：项目的 `uploads.auto_convert_documents=false`（host 端解析有安全风险，
刻意关闭），且 gateway 容器里没装 markitdown / 可能没装 python-docx。所以 agent 拿
不到现成的 .md，必须自己在 sandbox 里转一次。本脚本就是那条"转一次"的可靠路径——
优先用 python-docx（结构更好），装不上就回退到标准库 zipfile+xml（永远可用）。

用法（在 sandbox bash 里，用虚拟路径）：
    python /mnt/skills/public/fire-protection-report-v2/references/docx_to_md.py \
        "/mnt/user-data/uploads/xxx.docx" \
        "/mnt/user-data/workspace/xxx.md"

然后 read_file 那个 .md 即可。只跑一次，不要重复转、不要 base64。
"""
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _convert_with_docx(src: Path) -> tuple[str, int, int]:
    """python-docx 路径：保留段落 + 表格结构。"""
    from docx import Document  # type: ignore

    doc = Document(str(src))
    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if "heading 1" in style or style == "title":
            lines.append(f"# {text}")
        elif "heading 2" in style:
            lines.append(f"## {text}")
        elif "heading 3" in style:
            lines.append(f"### {text}")
        else:
            lines.append(text)
    n_tables = 0
    for table in doc.tables:
        n_tables += 1
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        header = rows[0]
        lines.append("")
        lines.append("| " + " | ".join(header) + " |")
        lines.append("|" + "|".join(["---"] * len(header)) + "|")
        for row in rows[1:]:
            # pad to header width
            cells = row + [""] * (len(header) - len(row))
            lines.append("| " + " | ".join(cells) + " |")
        lines.append("")
    return "\n".join(lines), len(doc.paragraphs), n_tables


def _convert_with_stdlib(src: Path) -> tuple[str, int, int]:
    """标准库回退：zipfile + 解析 word/document.xml，按段落重组 <w:t> 文本。"""
    with zipfile.ZipFile(str(src)) as z:
        xml = z.read("word/document.xml")
    root = ET.fromstring(xml)
    lines: list[str] = []
    n_para = 0
    n_tables = 0
    for p in root.iter(f"{{{W_NS}}}p"):
        texts = [t.text for t in p.iter(f"{{{W_NS}}}t") if t.text]
        line = "".join(texts).strip()
        if line:
            lines.append(line)
            n_para += 1
    # 表格：粗略抽取每个 <w:tbl> 的单元格文本
    for tbl in root.iter(f"{{{W_NS}}}tbl"):
        n_tables += 1
        rows: list[list[str]] = []
        for tr in tbl.iter(f"{{{W_NS}}}tr"):
            cells = []
            for tc in tr.iter(f"{{{W_NS}}}tc"):
                cell_text = "".join(t.text or "" for t in tc.iter(f"{{{W_NS}}}t")).strip()
                cells.append(cell_text)
            if cells:
                rows.append(cells)
        if not rows:
            continue
        width = max(len(r) for r in rows)
        lines.append("")
        lines.append("| " + " | ".join(rows[0] + [""] * (width - len(rows[0]))) + " |")
        lines.append("|" + "|".join(["---"] * width) + "|")
        for r in rows[1:]:
            cells = r + [""] * (width - len(r))
            lines.append("| " + " | ".join(cells) + " |")
        lines.append("")
    return "\n".join(lines), n_para, n_tables


def main() -> int:
    if len(sys.argv) != 3:
        print(__doc__)
        return 2
    src = Path(sys.argv[1])
    dst = Path(sys.argv[2])
    if not src.exists():
        print(f"ERROR: input not found: {src}", file=sys.stderr)
        return 1
    try:
        try:
            text, n_para, n_tbl = _convert_with_docx(src)
            backend = "python-docx"
        except ImportError:
            text, n_para, n_tbl = _convert_with_stdlib(src)
            backend = "stdlib(zipfile+xml)"
    except Exception as e:
        print(f"ERROR: conversion failed: {e}", file=sys.stderr)
        return 1
    dst.parent.mkdir(parents=True, exist_ok=True)
    dst.write_text(text, encoding="utf-8")
    print(f"OK [{backend}] -> {dst} | {n_para} paragraphs, {n_tbl} tables, {len(text)} chars")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
